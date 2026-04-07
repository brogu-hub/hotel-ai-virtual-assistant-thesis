#!/usr/bin/env python3
"""
Patch thesis_final.docx:
1. Add NVIDIA AI Blueprint section to CH2 (in Thai)
2. Replace CH5 with expanded implementation (in Thai)
3. Attach PNG figures at placeholder positions

Writes directly in Thai — no LLM translation.
"""
import os
import sys
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final.docx")
OUT = Path("thesis/thesis_final_v2.docx")
FIGURES_DIR = Path("thesis/figures")


def make_para(doc, text, style_name, bold=False, italic=False, font_name=None,
              font_size=None, color=None, alignment=None):
    """Create a w:p XML element."""
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        if italic:
            rPr.append(OxmlElement('w:i'))
        if font_name:
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), font_name)
            rf.set(qn('w:hAnsi'), font_name)
            rf.set(qn('w:cs'), font_name)
            rPr.append(rf)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(sz)
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p


def make_code_para(doc, text):
    """Create a code-style paragraph."""
    return make_para(doc, text, "TU_Paragraph_Normal", font_name="Consolas", font_size=10)


def insert_after(body, ref, nodes):
    """Insert nodes after ref element."""
    parent = ref.getparent()
    idx = list(parent).index(ref)
    for j, n in enumerate(nodes):
        parent.insert(idx + 1 + j, n)


def insert_image_paragraph(doc, body, ref_element, image_path, width_inches=5.5):
    """Insert an image paragraph after ref_element."""
    # We need to use the Document API for images, then move the element
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    # Move the paragraph element to the correct position
    elem = p._element
    parent = elem.getparent()
    parent.remove(elem)
    ref_parent = ref_element.getparent()
    idx = list(ref_parent).index(ref_element)
    ref_parent.insert(idx + 1, elem)
    return elem


def main():
    print("=" * 60)
    print("Patching thesis_final.docx")
    print("=" * 60)

    # Copy original
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    body = doc.element.body

    # =========================================================================
    # 1. Add NVIDIA AI Blueprint to CH2 (before section 2.3)
    # =========================================================================
    print("\n1. Adding NVIDIA AI Blueprint to CH2...")

    # Find "2.3 Agentic AI" heading
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "2.3" in p.text and "Agentic" in p.text:
            target_idx = i
            break

    if target_idx:
        ref = doc.paragraphs[target_idx]._element
        nvidia_nodes = [
            make_para(doc, "2.2.4 NVIDIA AI Blueprint สำหรับ Virtual Assistant",
                      "TU_Sub-heading 1", bold=True),
            make_para(doc, "โครงการนี้ถูกพัฒนาต่อยอดจาก NVIDIA AI Blueprint for AI Virtual Assistant "
                      "ซึ่งเป็น reference architecture ที่ NVIDIA เผยแพร่สำหรับการสร้างระบบ AI virtual assistant "
                      "ในระดับ production โดย blueprint ดั้งเดิมใช้ NVIDIA NIM endpoints สำหรับ LLM inference, "
                      "NVIDIA NeMo Guardrails สำหรับ safety filtering, และ Milvus สำหรับ vector storage",
                      "TU_Paragraph_Normal"),
            make_para(doc, "สถาปัตยกรรมหลักของ blueprint ประกอบด้วย:",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• LangGraph StateGraph สำหรับ multi-agent orchestration พร้อม conditional routing",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• RAG pipeline ที่ใช้ structured data (Vanna.AI สำหรับ NL-to-SQL) และ "
                      "unstructured data (document embeddings)",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• NeMo Guardrails สำหรับ input/output safety checking ด้วย Colang flows",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• API Gateway สำหรับ request routing และ load balancing",
                      "TU_Paragraph_Normal"),
            make_para(doc, "ในวิทยานิพนธ์นี้ได้ปรับเปลี่ยน blueprint ดังกล่าวอย่างมีนัยสำคัญ โดย:",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• แทนที่ NVIDIA NIM ด้วย OpenRouter API (Qwen3 Max) และ local Ollama "
                      "(Qwen3.5 Opus 9B) เพื่อลดค่าใช้จ่ายและเพิ่มความเป็นส่วนตัวของข้อมูล",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• แทนที่ Milvus ด้วย Qdrant สำหรับ vector storage เพื่อความง่ายในการ deploy "
                      "ผ่าน Docker container เดียว",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• เพิ่มระบบ JWT authentication, audit logging, chat scaling primitives, "
                      "และ admin dashboard ซึ่งไม่มีใน blueprint ดั้งเดิม",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• สร้าง hotel-specific tools 15 ตัว (จอง, เช็คอิน, ค้นหาความรู้, คำนวณราคา) "
                      "แทนที่ retail/e-commerce tools ของ blueprint เดิม",
                      "TU_Paragraph_Normal"),
            make_para(doc, "• พัฒนา frontend ด้วย Next.js 15 + Ant Design 5 (blueprint เดิมมีเฉพาะ backend)",
                      "TU_Paragraph_Normal"),
            make_para(doc, "การปรับเปลี่ยนเหล่านี้ทำให้ระบบเหมาะสมกับ domain โรงแรมในประเทศไทย "
                      "โดยรองรับการสนทนาสองภาษา (ไทย/อังกฤษ) และสามารถ deploy ได้ทั้งแบบ local "
                      "(ใช้ GPU ของโรงแรมเอง) และ cloud (ใช้ OpenRouter API)",
                      "TU_Paragraph_Normal"),
        ]
        # Insert before the 2.3 heading (so it appears after 2.2.3)
        parent = ref.getparent()
        idx = list(parent).index(ref)
        for j, n in enumerate(nvidia_nodes):
            parent.insert(idx + j, n)
        print(f"  Inserted {len(nvidia_nodes)} paragraphs before P{target_idx}")
    else:
        print("  [WARN] Could not find 2.3 Agentic AI heading")

    # =========================================================================
    # 2. Attach PNG figures at placeholder positions
    # =========================================================================
    print("\n2. Attaching PNG figures...")

    figure_map = {
        "1.1": "Fig_1.1_System_Context.png",
        "2.1": "Fig_2.1_Chatbot_Taxonomy.png",
        "2.2": "Fig_2.2_LangGraph_Concept.png",
        "2.3": "Fig_2.3_RAG_Pipeline.png",
        "3.1": "Fig_3.1_Framework_Comparison.png",
        "4.1": "Fig_4.1_System_Architecture.png",
        "4.2": "Fig_4.2_LangGraph_State_Machine.png",
        "4.3": "Fig_4.3_ER_Diagram.png",
        "4.4": "Fig_4.4_RAG_Pipeline_Detail.png",
        "4.5": "Fig_4.5_JWT_Auth_Flow.png",
        "4.6": "Fig_4.6_Access_Control_Matrix.png",
        "4.7": "Fig_4.7_Docker_Topology.png",
        "5.1": "Fig_5.1_Runtime_Model_Switch.png",
        "5.2": "Fig_5.2_Prompt_Structure.png",
        "5.3": "Fig_5.3_PII_Redaction.png",
        "5.4": "Fig_5.4_Scaling_Pipeline.png",
        "6.1": "Fig_6.1_Accuracy_Comparison.png",
        "6.2": "Fig_6.2_Test_Coverage_Pie.png",
        "6.3": "Fig_6.3_Before_After_Optimization.png",
        "6.4": "Fig_6.4_Kappa_Matrix.png",
        "7.1": "Fig_7.1_Cost_Accuracy_Latency.png",
        "7.2": "Fig_7.2_Orchestrator_Comparison.png",
    }

    attached = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "[Figure" in text or "Figure" in text and "Insert" in text:
            # Find which figure this is
            m = re.search(r'Figure\s+(\d+\.\d+)', text)
            if m:
                fig_id = m.group(1)
                fig_file = figure_map.get(fig_id)
                if fig_file:
                    fig_path = FIGURES_DIR / fig_file
                    if fig_path.exists():
                        try:
                            last_elem = insert_image_paragraph(doc, body, p._element,
                                                                fig_path, width_inches=5.0)
                            attached += 1
                            print(f"  [OK] Figure {fig_id} → {fig_file}")
                        except Exception as e:
                            print(f"  [FAIL] Figure {fig_id}: {e}")

    print(f"  Total figures attached: {attached}")

    # =========================================================================
    # 3. Save
    # =========================================================================
    doc.save(str(OUT))
    sz = os.path.getsize(str(OUT)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUT} ({sz:.0f} KB)")
    print("=" * 60)


if __name__ == "__main__":
    main()
