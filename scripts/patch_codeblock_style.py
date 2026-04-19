#!/usr/bin/env python3
"""
patch_codeblock_style.py
========================
Style code blocks in thesis docx with:
1. One Dark Pro color scheme (syntax highlighting)
2. Dark background (#282C34) with rounded-feel border
3. Consolas 9pt font (slightly smaller for density)

Input/Output: thesis/thesis_final_v6.docx (in-place)
"""
import os, sys, re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final_v6.docx")
DOCX = Path("thesis/thesis_final_v7.docx")

# ============================================================================
# One Dark Pro color palette
# ============================================================================
C_BG       = "282C34"   # background
C_DEFAULT  = "ABB2BF"   # default text (light gray)
C_KEYWORD  = "C678DD"   # purple — def, class, if, return, import, from, async, await
C_BUILTIN  = "61AFEF"   # blue — True, False, None, self, print, str, int, dict, list
C_STRING   = "98C379"   # green — strings
C_COMMENT  = "7F848E"   # gray — comments
C_NUMBER   = "D19A66"   # orange — numbers
C_DECORATOR= "E5C07B"   # yellow — @decorator
C_FUNCTION = "61AFEF"   # blue — function names after def
C_OPERATOR = "56B6C2"   # cyan — operators, arrows
C_JSKEY    = "E06C75"   # red — JSON keys, YAML keys
C_TYPE     = "E5C07B"   # yellow — type annotations, class names
C_CONST    = "D19A66"   # orange — constants, ALL_CAPS

# Python/JS keywords
KEYWORDS = {
    "def", "class", "if", "elif", "else", "for", "while", "return", "import",
    "from", "as", "try", "except", "finally", "with", "raise", "yield",
    "async", "await", "pass", "break", "continue", "in", "not", "and", "or",
    "is", "lambda", "global", "nonlocal", "assert", "del",
    # JS/TS
    "const", "let", "var", "function", "export", "default", "new", "typeof",
    "interface", "type", "extends", "implements", "enum",
    # YAML/Docker
    "services", "image", "environment", "deploy", "resources", "reservations",
    "devices", "driver", "count", "capabilities", "depends_on", "condition",
    "volumes", "ports",
}

BUILTINS = {
    "True", "False", "None", "self", "str", "int", "float", "dict", "list",
    "Optional", "Dict", "List", "Any", "Annotated", "Literal",
    "print", "len", "range", "super", "isinstance", "type",
    "asyncio", "datetime", "timedelta", "timezone", "uuid", "os", "re", "time",
    # React/TS
    "React", "useState", "useEffect", "fetch", "JSON", "Promise", "Response",
    "console", "document", "window", "localStorage",
}

TYPES_CLASSES = {
    "FastAPI", "Field", "BaseModel", "HTTPException", "Request", "Depends",
    "StateGraph", "TypedDict", "ChatPromptTemplate", "ChatOpenAI",
    "HotelState", "HumanMessage", "AIMessage", "ToolMessage",
    "RunnableConfig", "MemorySaver", "OxmlElement",
    "ThreadedConnectionPool", "Semaphore", "Lock",
    # TS/React
    "FC", "StateCreator", "Modal", "Form", "Input", "Button", "Card",
    "Space", "Steps", "Step", "Table",
}


def tokenize_code(text):
    """Split a code line into (color, text) tokens using One Dark Pro colors."""
    if not text or text.isspace():
        return [(C_DEFAULT, text or " ")]

    tokens = []
    i = 0
    n = len(text)
    stripped = text.lstrip()

    # Full-line comment
    if stripped.startswith("#") or stripped.startswith("//"):
        return [(C_COMMENT, text)]

    # YAML key: value
    yaml_match = re.match(r'^(\s*)([a-zA-Z_][\w.-]*)\s*:', text)

    # JSON key
    json_match = re.match(r'^(\s*)"([^"]+)"(\s*:\s*)', text)

    while i < n:
        ch = text[i]

        # Leading whitespace
        if ch in ' \t' and (not tokens or tokens[-1][1].isspace()):
            j = i
            while j < n and text[j] in ' \t':
                j += 1
            tokens.append((C_DEFAULT, text[i:j]))
            i = j
            continue

        # Inline comment # (not inside string)
        if ch == '#' and i > 0:
            tokens.append((C_COMMENT, text[i:]))
            break

        # String (single or double quote, including triple)
        if ch in ('"', "'"):
            quote = ch
            j = i + 1
            # Triple quote
            if text[i:i+3] in ('"""', "'''"):
                end = text.find(text[i:i+3], i+3)
                if end == -1:
                    tokens.append((C_STRING, text[i:]))
                    break
                tokens.append((C_STRING, text[i:end+3]))
                i = end + 3
                continue
            # Single-line string
            while j < n:
                if text[j] == '\\':
                    j += 2
                    continue
                if text[j] == quote:
                    j += 1
                    break
                j += 1
            tokens.append((C_STRING, text[i:j]))
            i = j
            continue

        # Decorator
        if ch == '@' and (i == 0 or text[i-1].isspace()):
            j = i + 1
            while j < n and (text[j].isalnum() or text[j] in '_.'):
                j += 1
            tokens.append((C_DECORATOR, text[i:j]))
            i = j
            continue

        # Number
        if ch.isdigit() or (ch == '.' and i + 1 < n and text[i+1].isdigit()):
            j = i
            while j < n and (text[j].isdigit() or text[j] in '.xXabcdefABCDEF_'):
                j += 1
            tokens.append((C_NUMBER, text[i:j]))
            i = j
            continue

        # Word (identifier/keyword)
        if ch.isalpha() or ch == '_':
            j = i
            while j < n and (text[j].isalnum() or text[j] == '_'):
                j += 1
            word = text[i:j]
            if word in KEYWORDS:
                tokens.append((C_KEYWORD, word))
            elif word in BUILTINS:
                tokens.append((C_BUILTIN, word))
            elif word in TYPES_CLASSES:
                tokens.append((C_TYPE, word))
            elif word.isupper() and len(word) > 1:  # CONSTANTS
                tokens.append((C_CONST, word))
            elif tokens and tokens[-1][1].rstrip() in ("def", "class"):
                tokens.append((C_FUNCTION, word))
            else:
                tokens.append((C_DEFAULT, word))
            i = j
            continue

        # Operators
        if ch in '=<>!+-*/&|^~%':
            j = i + 1
            while j < n and text[j] in '=<>!+-*/&|^~%>':
                j += 1
            tokens.append((C_OPERATOR, text[i:j]))
            i = j
            continue

        # Arrow =>
        if ch == '=' and i + 1 < n and text[i+1] == '>':
            tokens.append((C_OPERATOR, "=>"))
            i += 2
            continue

        # Punctuation / other
        tokens.append((C_DEFAULT, ch))
        i += 1

    return tokens


def make_colored_run(text, color, font_name="Consolas", font_size=9):
    """Create a w:r element with colored text."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), font_name)
    rf.set(qn('w:hAnsi'), font_name)
    rf.set(qn('w:cs'), font_name)
    rPr.append(rf)

    # Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    rPr.append(szCs)

    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def add_paragraph_shading_and_border(p_element, bg_color, is_first, is_last):
    """Add dark background and border to a paragraph element."""
    pPr = p_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_element.insert(0, pPr)

    # Background shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_color)
    pPr.append(shd)

    # Borders
    pBdr = OxmlElement('w:pBdr')

    # Left border (accent line)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')  # 1.5pt
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '528BFF')  # One Dark Pro accent blue
    pBdr.append(left)

    # Right border
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '4')
    right.set(qn('w:space'), '4')
    right.set(qn('w:color'), '3E4451')
    pBdr.append(right)

    # Top border (only on first line of block)
    if is_first:
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '3E4451')
        pBdr.append(top)

    # Bottom border (only on last line of block)
    if is_last:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '3E4451')
        pBdr.append(bottom)

    pPr.append(pBdr)

    # Reduce line spacing for code density
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')  # single spacing
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

    # Small indent for padding feel
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '284')   # ~0.5cm left padding
    ind.set(qn('w:right'), '284')  # ~0.5cm right padding
    pPr.append(ind)


def is_code_paragraph(p):
    """Check if paragraph has Consolas font (code block)."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r in p.runs:
        rpr = r._element.find(f'{{{ns}}}rPr')
        if rpr is not None:
            rf = rpr.find(f'{{{ns}}}rFonts')
            if rf is not None and rf.get(f'{{{ns}}}ascii', '') == 'Consolas':
                return True
    return False


def process_code_blocks(doc):
    """Find consecutive code paragraphs, apply syntax highlighting + frame."""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    paragraphs = list(doc.paragraphs)
    n = len(paragraphs)

    # Identify code block boundaries
    blocks = []  # list of (start_idx, end_idx) inclusive
    i = 0
    while i < n:
        if is_code_paragraph(paragraphs[i]):
            start = i
            while i < n and is_code_paragraph(paragraphs[i]):
                i += 1
            blocks.append((start, i - 1))
        else:
            i += 1

    print(f"  Found {len(blocks)} code blocks ({sum(e-s+1 for s,e in blocks)} lines)")

    # Process each block
    for block_idx, (start, end) in enumerate(blocks):
        for line_idx in range(start, end + 1):
            p = paragraphs[line_idx]
            is_first = (line_idx == start)
            is_last = (line_idx == end)

            # Get original text
            original_text = p.text

            # Remove existing runs
            for r in list(p._element.findall(qn('w:r'))):
                p._element.remove(r)
            # Also remove hyperlinks that might contain runs
            for hl in list(p._element.findall(qn('w:hyperlink'))):
                p._element.remove(hl)

            # Tokenize and create colored runs
            tokens = tokenize_code(original_text)
            for color, text in tokens:
                run_elem = make_colored_run(text, color)
                p._element.append(run_elem)

            # Add background + border
            add_paragraph_shading_and_border(p._element, C_BG, is_first, is_last)

    return len(blocks)


def main():
    print("=" * 60)
    print("Styling code blocks: One Dark Pro + frame")
    print("=" * 60)

    import shutil
    shutil.copy2(SRC, DOCX)
    doc = Document(str(DOCX))

    print("\nProcessing code blocks...")
    blocks = process_code_blocks(doc)

    doc.save(str(DOCX))
    sz = os.path.getsize(str(DOCX)) / 1024
    print(f"\nSaved: {DOCX} ({sz:.0f} KB)")
    print(f"Styled {blocks} code blocks")
    print("=" * 60)


if __name__ == "__main__":
    main()
