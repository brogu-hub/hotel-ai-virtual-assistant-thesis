#!/usr/bin/env python3
"""
patch_thai_translation.py
=========================
Inject Thai translations for Abstract, CH5, CH6, CH7, CH8 into thesis docx.

Input:  thesis/thesis_final_v2.docx
Output: thesis/thesis_final_v3.docx

Content is written directly in Thai with English technical terms preserved.
Code blocks remain in English. Tables are created as Word tables.
"""

import os
import sys
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final_v2.docx")
OUT = Path("thesis/thesis_final_v3.docx")
FIGURES_DIR = Path("thesis/figures")


# ============================================================================
# Helper functions
# ============================================================================

def make_para(doc, text, style_name, bold=False, italic=False,
              font_name=None, font_size=None, alignment=None):
    """Create a paragraph XML element with the given style."""
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
            rPr.append(OxmlElement('w:bCs'))
        if italic:
            rPr.append(OxmlElement('w:i'))
            rPr.append(OxmlElement('w:iCs'))
        if font_name:
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), font_name)
            rf.set(qn('w:hAnsi'), font_name)
            rf.set(qn('w:cs'), font_name)
            rPr.append(rf)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size * 2)))
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(sz)
            rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p


def make_code_para(doc, text):
    return make_para(doc, text, "TU_Paragraph_Normal",
                     font_name="Consolas", font_size=10)


def build_nodes(doc, specs):
    """Build paragraph XML nodes from specification list.
    Each spec: (kind, text)
      h1 → TU_Sub-heading 1, h2 → TU_Sub-heading 2, h3 → TU_Sub-heading 3
      p → TU_Paragraph_Normal, pb → bold, pi → italic
      code → multi-line Consolas, figure → italic centered
    """
    nodes = []
    for spec in specs:
        kind = spec[0]
        text = spec[1] if len(spec) > 1 else ""
        if kind == "h1":
            nodes.append(make_para(doc, text, "TU_Sub-heading 1", bold=True))
        elif kind == "h2":
            nodes.append(make_para(doc, text, "TU_Sub-heading 2", bold=True))
        elif kind == "h3":
            nodes.append(make_para(doc, text, "TU_Sub-heading 3", bold=True))
        elif kind == "p":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal"))
        elif kind == "pb":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", bold=True))
        elif kind == "pi":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", italic=True))
        elif kind == "code":
            for line in text.split("\n"):
                nodes.append(make_code_para(doc, line if line else " "))
        elif kind == "figure":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", italic=True,
                                   alignment="center"))
    return nodes


def insert_after_element(body, ref_elem, nodes):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for j, node in enumerate(nodes):
        parent.insert(idx + 1 + j, node)


def remove_between_elements(body, start_elem, end_elem):
    """Remove all body children between start_elem and end_elem (exclusive both)."""
    children = list(body)
    start_pos = children.index(start_elem)
    end_pos = children.index(end_elem) if end_elem is not None else len(children)
    to_remove = children[start_pos + 1:end_pos]
    for elem in to_remove:
        body.remove(elem)
    return len(to_remove)


def create_word_table(doc, headers, rows, ref_element):
    """Create a Word table and position it after ref_element."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols, style='Table Grid')
    hdr = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                table.rows[i + 1].cells[j].text = str(cell_text)
    tbl = table._tbl
    tbl_parent = tbl.getparent()
    tbl_parent.remove(tbl)
    ref_parent = ref_element.getparent()
    idx = list(ref_parent).index(ref_element)
    ref_parent.insert(idx + 1, tbl)
    return tbl


def insert_image_paragraph(doc, body, ref_element, image_path, width_inches=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    elem = p._element
    parent = elem.getparent()
    parent.remove(elem)
    ref_parent = ref_element.getparent()
    idx = list(ref_parent).index(ref_element)
    ref_parent.insert(idx + 1, elem)
    return elem


# ============================================================================
# Chapter boundary detection
# ============================================================================

def find_chapter_indices(doc):
    """Scan paragraphs to find chapter heading indices."""
    chapters = {}

    # Map heading text to chapter number
    HEADING_MAP = {
        "introduction": 1,
        "literature review": 2,
        "methodology": 3,
        "system design": 4,
        "implementation": 5,
        "testing and evaluation": 6,
        "discussion": 7,
        "conclusion": 8,
    }

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""

        # Heading 1 = chapter titles (English names)
        if style == "Heading 1":
            lower = text.lower()
            for heading_text, ch_num in HEADING_MAP.items():
                if heading_text in lower:
                    if ch_num not in chapters:
                        chapters[ch_num] = i
                    break

        # TU_Chapter = front matter sections
        if style == "TU_Chapter":
            if "บทคัดย่อ" in text:
                chapters["abstract_th"] = i
            elif text.strip() == "ABSTRACT":
                chapters["abstract_en"] = i
            elif "ภาคผนวก" in text or "Appendix" in text:
                if "appendix" not in chapters:
                    chapters["appendix"] = i

        # Also check for Thai chapter headings
        ch_match = re.search(r'บทที่\s*(\d+)', text)
        if ch_match:
            ch_num = int(ch_match.group(1))
            if ch_num not in chapters:
                chapters[ch_num] = i

    return chapters


# ============================================================================
# TABLE DEFINITIONS
# ============================================================================

TABLES = {
    "TABLE_5_1": (
        ["ระดับ", "เทคโนโลยี", "เวอร์ชัน", "วัตถุประสงค์"],
        [
            ["Backend framework", "FastAPI", "0.115", "Async REST API พร้อม auto OpenAPI docs"],
            ["Agent framework", "LangGraph", "0.2.32", "Multi-agent state machine"],
            ["LLM interface", "LangChain-OpenAI", "0.1+", "Unified LLM API สำหรับ Ollama/OpenRouter"],
            ["Database", "PostgreSQL", "16", "Relational storage (rooms, bookings, users)"],
            ["Vector store", "Qdrant", "latest", "Embedding similarity search"],
            ["Cache", "Redis", "7.0", "Session state"],
            ["Local LLM", "Ollama", "latest", "GPU inference สำหรับ Qwen3.5 Opus 9B"],
            ["Cloud LLM", "OpenRouter API", "—", "Qwen3 Max (cloud fallback)"],
            ["Frontend", "Next.js", "15.5", "React App Router พร้อม RSC"],
            ["UI library", "Ant Design", "5.29", "Enterprise component library"],
            ["State mgmt", "Zustand", "5.0", "Lightweight global state"],
            ["Data fetching", "SWR", "2.4", "stale-while-revalidate hooks"],
            ["Language", "TypeScript", "5.x", "Static typing"],
            ["Auth", "PyJWT + bcrypt", "2.12/5.0", "JWT tokens + password hashing"],
            ["Container", "Docker Compose", "—", "5-service orchestration"],
        ],
    ),
    "TABLE_5_2": (
        ["เอกสาร", "เนื้อหา"],
        [
            ["dining.md", "เวลาร้านอาหาร เมนู room service"],
            ["spa.md", "ทรีตเมนต์สปา ราคา การจอง"],
            ["facilities.md", "สระว่ายน้ำ ฟิตเนส business center"],
            ["policies.md", "นโยบายยกเลิก สัตว์เลี้ยง สูบบุหรี่ เด็ก"],
            ["faq.md", "เวลา check-in/out WiFi ฝากกระเป๋า"],
            ["transport.md", "รถรับส่งสนามบิน แท็กซี่ BTS"],
            ["rooms.md", "ประเภทห้อง สิ่งอำนวยความสะดวก วิว"],
            ["services.md", "Concierge ซักรีด ปลุก"],
            ["events.md", "ห้องประชุม แพ็คเกจงานแต่งงาน"],
            ["loyalty.md", "โปรแกรมสะสมแต้ม ระดับสมาชิก"],
        ],
    ),
    "TABLE_5_3": (
        ["Tool", "วัตถุประสงค์"],
        [
            ["check_room_availability", "ค้นหาห้องว่างตามช่วงวัน"],
            ["create_reservation", "สร้างการจองพร้อม dynamic pricing"],
            ["confirm_reservation", "ยืนยันการจองที่ pending"],
            ["update_reservation", "แก้ไขวันที่ จำนวนผู้เข้าพัก หรือห้อง"],
            ["cancel_reservation", "ยกเลิกพร้อมเหตุผล"],
            ["check_in_guest", "บันทึก check-in"],
            ["check_out_guest", "บันทึก check-out"],
            ["get_reservation_details", "ค้นหาด้วย confirmation number"],
            ["get_guest_reservations", "ค้นหาด้วย email"],
            ["calculate_dynamic_price", "คำนวณราคา early bird / last-minute"],
            ["check_upsell_opportunity", "แนะนำ upgrade ห้องหลัง booking"],
            ["generate_payment_link", "สร้าง mock payment URL"],
            ["search_hotel_knowledge", "ค้นหา RAG พร้อม knowledge cache"],
            ["get_hotel_services", "แสดงบริการที่มี"],
            ["create_service_request", "ขอผ้าเช็ดตัว room service ฯลฯ"],
        ],
    ),
    "TABLE_5_4": (
        ["Config", "Per-request Latency", "VRAM", "GPU Offload"],
        [
            ["NUM_PARALLEL=4", "30-50s", "9.9 GB", "100% GPU"],
            ["NUM_PARALLEL=2 (เลือกใช้)", "4-9s", "9.9 GB", "100% GPU"],
            ["NUM_PARALLEL=2 + Q8_0 KV", ">120s", "7.9 GB", "37% GPU / 63% CPU"],
        ],
    ),
    "TABLE_6_1": (
        ["หมวดหมู่", "จำนวน", "ขอบเขต"],
        [
            ["Knowledge (K01-K08)", "8", "อาหารเช้า WiFi สระว่ายน้ำ สัตว์เลี้ยง นโยบายยกเลิก สปา check-in/out รถรับส่ง"],
            ["Booking (B01-B06)", "6", "ห้องว่าง ราคา ยกเลิก วันที่ไทย full booking ค้นหาการจอง"],
            ["Greeting (G01-G04)", "4", "ทักทายอังกฤษ/ไทย ขอบคุณ off-topic"],
            ["Language (L01-L03)", "3", "ตอบเฉพาะอังกฤษ เฉพาะไทย ตำแหน่งฟิตเนส"],
            ["Edge Cases (E01-E04)", "4", "ขอบริการ จองวันหยุด จองหลายห้อง ข้อความว่าง"],
        ],
    ),
    "TABLE_6_2": (
        ["ตัวชี้วัด", "Qwen3.5 Opus 9B (Local)", "Qwen3 Max (Cloud)"],
        [
            ["ความแม่นยำโดยรวม", "23/25 (92%)", "25/25 (100%)"],
            ["Keyword accuracy", "81%", "89%"],
            ["Language accuracy", "100%", "100%"],
            ["Errors / timeouts", "0", "0"],
        ],
    ),
    "TABLE_6_3": (
        ["หมวดหมู่", "Local 9B", "Cloud", "Agreement"],
        [
            ["Knowledge (8)", "8/8 (100%)", "8/8 (100%)", "Perfect"],
            ["Booking (6)", "6/6 (100%)", "6/6 (100%)", "Perfect"],
            ["Greeting (4)", "3/4 (75%)", "4/4 (100%)", "3/4"],
            ["Language (3)", "3/3 (100%)", "3/3 (100%)", "Perfect"],
            ["Edge Cases (4)", "3/4 (75%)", "4/4 (100%)", "3/4"],
        ],
    ),
    "TABLE_6_4": (
        ["ตัวชี้วัด", "Local 9B", "Cloud"],
        [
            ["Average", "9,879 ms", "8,921 ms"],
            ["Median (p50)", "9,049 ms", "6,703 ms"],
            ["p95", "18,360 ms", "37,955 ms"],
        ],
    ),
    "TABLE_6_5": (
        ["ชุดทดสอบ", "จำนวน", "ผ่าน", "ขอบเขต"],
        [
            ["Auth Baseline", "72", "72", "Login, register, JWT, role separation"],
            ["Auth Hardening", "38", "38", "Rate limiting, lockout, blocklist, password change"],
            ["Audit + DB Scaling", "46", "46", "Audit CRUD, filters, DB pool, user cache"],
            ["Chat Scaling", "37", "37", "LLM semaphore, session locks, rate limit, SSE cap"],
        ],
    ),
    "TABLE_6_6": (
        ["การปรับปรุง", "ก่อน", "หลัง", "ผลกระทบ"],
        [
            ["ลบ Reranker", "18s/chat", "5s/chat", "เร็วขึ้น 3.6x"],
            ["ตัดทอน Prompt", "5,500 chars", "2,800 chars", "-50% tokens"],
            ["NUM_PARALLEL 4→2", "30-50s/chat", "4-9s/chat", "Full GPU throughput"],
            ["Flash attention", "—", "เปิดใช้", "เร่ง attention"],
            ["Q8_0 KV cache", "ทดสอบ", "ลบออก", "CPU offload ช้า 10x"],
            ["Knowledge cache", "ไม่มี", "500 entries 5min TTL", "~1ms vs ~500ms"],
            ["DB connection pool", "New conn/req", "Pool min=2 max=20", "ลดต้นทุน conn"],
        ],
    ),
    "TABLE_6_7": (
        ["Config", "VRAM", "GPU %", "Per-Request Latency"],
        [
            ["NUM_PARALLEL=4", "9.9 GB", "100%", "30-50s"],
            ["NUM_PARALLEL=2 (เลือกใช้)", "9.9 GB", "100%", "4-9s"],
            ["NUM_PARALLEL=2 + Q8_0 KV", "7.9 GB", "37% GPU / 63% CPU", ">120s"],
        ],
    ),
    "TABLE_6_8": (
        ["Benchmark", "ผลลัพธ์"],
        [
            ["30 concurrent GET /auth/me", "รวม 0.06s, p95 20ms"],
            ["50 sequential GET /admin/audit", "50/50 success"],
            ["20 concurrent GET /admin/audit", "รวม 0.16s, ทุก request 200"],
            ["2 concurrent /chat (within GPU slots)", "14.5s wall time"],
            ["4 concurrent /chat (2 queued)", "19.7s wall time, 0 failures"],
        ],
    ),
    "TABLE_7_1": (
        ["วัตถุประสงค์", "สถานะ", "หลักฐาน"],
        [
            ["(ก) Multi-agent LangGraph", "สำเร็จ", "4 sub-agents, conditional routing, tool loops, checkpointed memory"],
            ["(ข) RAG ฐานความรู้โรงแรม", "สำเร็จ", "8/8 knowledge tests ผ่านทั้ง local และ cloud"],
            ["(ค) Fullstack ระดับ production", "สำเร็จ", "193/193 tests, JWT, audit, scaling, Docker"],
            ["(ง) เปรียบเทียบ local vs cloud", "สำเร็จ", "25-case eval: 92% local, 100% cloud"],
        ],
    ),
    "TABLE_7_2": (
        ["มิติ", "Local (Qwen3.5 Opus 9B)", "Cloud (Qwen3 Max)"],
        [
            ["ต้นทุนต่อ query", "$0 (amortized GPU)", "~$0.001-0.01"],
            ["ต้นทุน/เดือน (1000 q/day)", "$0 operational", "~$30-100"],
            ["ความแม่นยำ", "92% (23/25)", "100% (25/25)"],
            ["p50 latency", "9.0s", "6.7s"],
            ["p95 latency", "18.4s (คาดเดาได้)", "38.0s (ผันผวน)"],
            ["ความเป็นส่วนตัว", "On-premise", "Third-party API"],
            ["Max concurrent", "2 (GPU-bound)", "ไม่จำกัด"],
            ["Offline", "ได้เต็มที่", "ต้องใช้ internet"],
        ],
    ),
    "TABLE_7_3": (
        ["เกณฑ์", "LangGraph", "CrewAI", "AutoGen"],
        [
            ["State mgmt", "TypedDict + checkpointer", "Manual state passing", "Conversation-based"],
            ["Persistence", "AsyncPostgresSaver", "ไม่มี", "ไม่มี"],
            ["Tool calling", "ToolNode + error handling", "Tool decorator", "Function calling"],
            ["Debugging", "Time-travel replay", "จำกัด", "จำกัด"],
            ["Visualization", "Built-in draw_mermaid", "ไม่มี", "ไม่มี"],
            ["Maturity", "Production-grade", "กำลังเติบโต", "เน้นวิจัย"],
        ],
    ),
}


# ============================================================================
# CONTENT: Abstract (บทคัดย่อ)
# ============================================================================

def get_abstract_specs():
    return [
        ("p", "วิทยานิพนธ์นี้นำเสนอการออกแบบ พัฒนา และประเมินผลระบบ AI Virtual Assistant "
              "สำหรับธุรกิจโรงแรม โดยใช้กรณีศึกษาโรงแรม The Grand Horizon Hotel "
              "ระบบถูกพัฒนาต่อยอดจาก NVIDIA AI Blueprint for AI Virtual Assistant "
              "ซึ่งเป็น reference architecture สำหรับการสร้างระบบ AI assistant ระดับ production "
              "โดยปรับเปลี่ยนให้เหมาะสมกับ domain โรงแรมในประเทศไทยที่รองรับ "
              "การสนทนาสองภาษาทั้งไทยและอังกฤษ"),
        ("p", "สถาปัตยกรรมหลักของระบบใช้ LangGraph state machine สำหรับ multi-agent orchestration "
              "ประกอบด้วย 4 sub-agents ได้แก่ การจองห้องพัก (booking) บริการโรงแรม (service) "
              "ฐานความรู้ (knowledge) และบทสนทนาทั่วไป (general conversation) พร้อมด้วย "
              "hotel-specific tools จำนวน 15 ตัวที่เชื่อมต่อกับฐานข้อมูล PostgreSQL "
              "โดยใช้ Retrieval-Augmented Generation (RAG) pipeline ผ่าน Qdrant vector store"),
        ("p", "ระบบ backend พัฒนาด้วย FastAPI พร้อมระบบรักษาความปลอดภัยครบวงจร ประกอบด้วย "
              "JWT authentication ที่รองรับ role-based access control แบ่งเป็น user และ admin, "
              "audit logging จำนวน 26 ประเภท, PII redaction สำหรับปกปิดข้อมูลส่วนบุคคล "
              "และ scaling primitives 5 ตัวสำหรับรองรับผู้ใช้งานพร้อมกัน "
              "ส่วน frontend พัฒนาด้วย Next.js 15 (App Router) ร่วมกับ Ant Design 5, "
              "Zustand สำหรับ state management และ SWR สำหรับ data fetching "
              "ระบบทั้งหมด deploy ผ่าน Docker Compose จำนวน 5 services"),
        ("p", "ผลการประเมินด้วย golden dataset จำนวน 25 test cases พบว่า local model "
              "(Qwen3.5 Opus 9B ทำงานบน Ollama ด้วย RTX 5080) มีความแม่นยำ 92% (23/25) "
              "ขณะที่ cloud model (Qwen3 Max บน OpenRouter) มีความแม่นยำ 100% (25/25) "
              "ทั้งสองโมเดลมีความแม่นยำ 100% ในหมวดความรู้โรงแรมและการจอง "
              "แต่ local model มี latency ที่คงที่และคาดเดาได้ดีกว่า "
              "(p95: 18.4 วินาที เทียบกับ 38.0 วินาที) "
              "การทดสอบโครงสร้างพื้นฐาน 193 tests ผ่านทั้งหมด"),
        ("p", "การเพิ่มประสิทธิภาพระบบประกอบด้วย การลบ reranker ทำให้ latency ของ warm chat "
              "ลดลง 3.6 เท่า (จาก 18 วินาทีเหลือ 5 วินาที) การตัดทอน system prompt "
              "จาก 5,500 เหลือ 2,800 ตัวอักษร และการปรับแต่ง Ollama GPU "
              "(NUM_PARALLEL=2, FLASH_ATTENTION=1) ทำให้ระบบรองรับ "
              "5 concurrent users ได้ภายใน 3 วินาที"),
        ("pb", "คำสำคัญ: AI Virtual Assistant, LangGraph, Retrieval-Augmented Generation, "
               "Hotel Chatbot, Multi-Agent System, Ollama, FastAPI, Next.js"),
    ]


# ============================================================================
# CONTENT: Chapter 5 — Implementation (การพัฒนาระบบ)
# ============================================================================

def get_ch5_specs():
    return [
        # --- 5.1 ---
        ("h1", "5.1 สภาพแวดล้อมการพัฒนาและเครื่องมือ"),
        ("p", "ตารางที่ 5.1 แสดงเทคโนโลยีหลักที่ใช้ในการพัฒนาระบบ แบ่งตามระดับชั้นของสถาปัตยกรรม"),
        ("p", "TABLE_5_1_PLACEHOLDER"),

        # --- 5.2 ---
        ("h1", "5.2 ลำดับขั้นตอนการพัฒนาระบบ"),
        ("p", "ระบบถูกพัฒนาแบบ incremental ตามลำดับ commit history ใน git repository "
              "โดยแต่ละ phase สร้าง working increment ที่ deploy ผ่าน Docker Compose "
              "และผ่านการทดสอบก่อนดำเนินการ phase ถัดไป"),

        # Phase 1
        ("h2", "5.2.1 Phase 1: พื้นฐานระบบ — FastAPI Server และ Database"),
        ("p", "Phase แรกเป็นการวางโครงสร้างพื้นฐานของระบบ ประกอบด้วย FastAPI application server "
              "พร้อม OpenAPI documentation, PostgreSQL database schema จำนวน 10 ตาราง "
              "และ REST endpoints พื้นฐานสำหรับ rooms, bookings และ session history"),
        ("code", "# src/hotel_guardrails/server.py — FastAPI app setup\n"
                 "app = FastAPI(\n"
                 "    title=\"The Grand Horizon Hotel Concierge API\",\n"
                 "    version=\"1.0.0\",\n"
                 "    docs_url=\"/docs\",\n"
                 "    redoc_url=\"/redoc\",\n"
                 ")"),
        ("p", "ไฟล์หลักที่สร้างในขั้นตอนนี้ ได้แก่:"),
        ("p", "• server.py — FastAPI application พร้อม middleware, CORS และ error handling"),
        ("p", "• database.py — PostgreSQL connection และ CRUD operations"),
        ("p", "• models.py — Pydantic request/response models"),
        ("p", "• init-hotel.sql — Database schema จำนวน 10 ตาราง"),

        # Phase 2
        ("h2", "5.2.2 Phase 2: RAG Pipeline — Knowledge Base และ Vector Search"),
        ("p", "ระบบ RAG pipeline ถูกสร้างขึ้นเพื่อให้ chatbot สามารถตอบคำถามเกี่ยวกับข้อมูลโรงแรม "
              "เช่น เวลาอาหารเช้า รหัส WiFi บริการสปา เป็นต้น โดยอาศัยฐานความรู้ที่มีโครงสร้าง"),
        ("code", "# src/retrievers/hotel_knowledge/chains.py\n"
                 "class HotelKnowledgeRetriever:\n"
                 "    def __init__(self):\n"
                 "        self.embeddings = get_openrouter_embeddings()  # qwen3-embedding-8b, 4096 dims\n"
                 "        self.vectorstore = create_qdrant_vectorstore(\n"
                 "            self.embeddings, collection_name=\"hotel_knowledge\"\n"
                 "        )\n"
                 "        self.top_k_retrieval = 30\n"
                 "        self.reranker = get_reranker(top_n=5)  # ถูกลบในภายหลัง\n"
                 "\n"
                 "    def document_search(self, content, num_docs=3):\n"
                 "        retriever = self.vectorstore.as_retriever(\n"
                 "            search_kwargs={\"k\": self.top_k_retrieval}\n"
                 "        )\n"
                 "        docs = retriever.invoke(content)\n"
                 "        return [{\"source\": doc.metadata.get(\"source\"),\n"
                 "                 \"content\": doc.page_content} for doc in docs[:num_docs]]"),
        ("p", "ฐานความรู้ประกอบด้วยเอกสาร markdown จำนวน 10 ฉบับเฉพาะด้านโรงแรม ดังตารางที่ 5.2"),
        ("p", "TABLE_5_2_PLACEHOLDER"),

        # Phase 3
        ("h2", "5.2.3 Phase 3: ระบบ LangGraph Multi-Agent"),
        ("p", "แกนหลักของระบบ AI agent ถูกพัฒนาเป็น LangGraph state machine ที่ประกอบด้วย "
              "4 specialized sub-agents ขั้นตอนนี้เป็นส่วนที่ซับซ้อนที่สุดในการพัฒนา "
              "เนื่องจากต้องมีการ debug tool-call routing, state management "
              "และ prompt engineering อย่างต่อเนื่อง"),
        ("code", "# src/hotel_guardrails/hotel_langgraph.py — Graph construction\n"
                 "def build_hotel_graph(checkpointer=None):\n"
                 "    builder = StateGraph(HotelState)\n"
                 "\n"
                 "    # Nodes\n"
                 "    builder.add_node(\"primary_assistant\", HotelAssistant(primary_prompt, primary_tools))\n"
                 "    builder.add_node(\"enter_booking\", create_entry_node(\"Booking Assistant\"))\n"
                 "    builder.add_node(\"enter_service\", create_entry_node(\"Service Assistant\"))\n"
                 "    builder.add_node(\"enter_knowledge\", create_entry_node(\"Knowledge Assistant\"))\n"
                 "    builder.add_node(\"hotel_booking\", handle_booking)\n"
                 "    builder.add_node(\"hotel_service\", handle_service)\n"
                 "    builder.add_node(\"hotel_knowledge\", handle_knowledge)\n"
                 "    builder.add_node(\"other_talk\", handle_other_talk)\n"
                 "\n"
                 "    # Tool nodes (cyclic loops)\n"
                 "    builder.add_node(\"booking_tools\", create_tool_node_with_fallback(booking_tools))\n"
                 "    builder.add_node(\"service_tools\", create_tool_node_with_fallback(service_tools))\n"
                 "\n"
                 "    # Edges\n"
                 "    builder.add_edge(START, \"primary_assistant\")\n"
                 "    builder.add_conditional_edges(\"primary_assistant\", route_primary_assistant)\n"
                 "    builder.add_edge(\"enter_booking\", \"hotel_booking\")\n"
                 "    builder.add_conditional_edges(\"hotel_booking\", route_booking)\n"
                 "    builder.add_edge(\"booking_tools\", \"hotel_booking\")\n"
                 "\n"
                 "    return builder.compile(checkpointer=checkpointer)"),
        ("p", "ระบบมี hotel-specific tools จำนวน 15 ตัว ดังตารางที่ 5.3"),
        ("p", "TABLE_5_3_PLACEHOLDER"),

        # Phase 4
        ("h2", "5.2.4 Phase 4: Docker Stack และ Local LLM"),
        ("p", "Docker Compose stack ที่ประกอบด้วย 5 services ถูกสร้างขึ้น "
              "เพื่อทดแทน cloud dependencies ด้วย local alternatives"),
        ("code", "# deploy/compose/docker-compose.hotel.yaml\n"
                 "services:\n"
                 "  hotel-ollama:    # GPU LLM inference\n"
                 "    image: ollama/ollama:latest\n"
                 "    environment:\n"
                 "      OLLAMA_NUM_PARALLEL: 2\n"
                 "      OLLAMA_FLASH_ATTENTION: 1\n"
                 "    deploy:\n"
                 "      resources:\n"
                 "        reservations:\n"
                 "          devices:\n"
                 "            - driver: nvidia\n"
                 "              count: all\n"
                 "              capabilities: [gpu]\n"
                 "\n"
                 "  hotel-db:        # PostgreSQL 16\n"
                 "  hotel-redis:     # Session cache\n"
                 "  hotel-qdrant:    # Vector store\n"
                 "  hotel-api:       # FastAPI + LangGraph\n"
                 "    depends_on:\n"
                 "      hotel-ollama: { condition: service_healthy }\n"
                 "      hotel-db: { condition: service_healthy }"),
        ("p", "ระบบ runtime model switching ถูกพัฒนาผ่าน thread-safe singleton pattern:"),
        ("code", "# src/hotel_guardrails/config.py\n"
                 "class RuntimeLLMConfig:\n"
                 "    \"\"\"Thread-safe singleton for runtime LLM configuration.\"\"\"\n"
                 "    _instance = None\n"
                 "    _lock = threading.Lock()\n"
                 "\n"
                 "    def update(self, backend=None, model=None, temperature=None, max_tokens=None):\n"
                 "        with self._lock:\n"
                 "            changes = {}\n"
                 "            if backend:\n"
                 "                self.backend = LLMBackend(backend.lower())\n"
                 "                changes[\"backend\"] = backend\n"
                 "            if model:\n"
                 "                self.active_model = model\n"
                 "                changes[\"model\"] = model\n"
                 "            return changes"),

        # Phase 5
        ("h2", "5.2.5 Phase 5: ระบบ Authentication และ Security"),
        ("p", "ระบบ JWT authentication พร้อม role-based access control ถูกเพิ่มเข้ามา "
              "เพื่อแยก guest endpoints (สาธารณะ) และ admin endpoints (ต้องมีสิทธิ์) ออกจากกัน"),
        ("code", "# src/hotel_guardrails/auth.py\n"
                 "def hash_password(password: str) -> str:\n"
                 "    return bcrypt.hashpw(password.encode(\"utf-8\"),\n"
                 "                         bcrypt.gensalt(rounds=12)).decode(\"utf-8\")\n"
                 "\n"
                 "def create_access_token(data: Dict[str, Any]) -> str:\n"
                 "    to_encode = data.copy()\n"
                 "    now = datetime.now(timezone.utc)\n"
                 "    to_encode.update({\n"
                 "        \"iat\": now,\n"
                 "        \"exp\": now + timedelta(hours=JWT_EXPIRE_HOURS),\n"
                 "        \"jti\": uuid.uuid4().hex,\n"
                 "    })\n"
                 "    return jwt.encode(to_encode, JWT_SECRET, algorithm=\"HS256\")\n"
                 "\n"
                 "async def get_current_user(credentials=Depends(bearer_scheme)):\n"
                 "    payload = decode_access_token(credentials.credentials)\n"
                 "    if token_blocklist.contains(payload.get(\"jti\")):\n"
                 "        raise HTTPException(401, \"Token revoked\")\n"
                 "    user = await db.get_user_by_username(payload[\"sub\"])\n"
                 "    if int(payload[\"iat\"]) < int(user[\"password_changed_at\"].timestamp()):\n"
                 "        raise HTTPException(401, \"Token invalidated by password change\")\n"
                 "    return user"),
        ("p", "ระบบ production hardening เพิ่ม 3 ชั้นการป้องกันที่ login endpoint:"),
        ("code", "# src/hotel_guardrails/server.py — Login rate limiting\n"
                 "@app.post(\"/auth/login\")\n"
                 "async def auth_login(request: UserLoginRequest, http_request: Request):\n"
                 "    client_ip = get_client_ip(http_request)\n"
                 "\n"
                 "    # Layer 1: Per-IP rate limit (100/min)\n"
                 "    allowed, retry_after = login_rate_limiter_ip.check_and_record(client_ip)\n"
                 "    if not allowed:\n"
                 "        raise HTTPException(429, headers={\"Retry-After\": str(retry_after)})\n"
                 "\n"
                 "    # Layer 2: Per-username rate limit (5/min)\n"
                 "    allowed_u, retry_u = login_rate_limiter_user.check_and_record(username)\n"
                 "    if not allowed_u:\n"
                 "        raise HTTPException(429, headers={\"Retry-After\": str(retry_u)})\n"
                 "\n"
                 "    # Layer 3: Account lockout (5 failures -> 15-min lock)\n"
                 "    lockout = check_account_lockout(user)\n"
                 "    if lockout:\n"
                 "        raise HTTPException(423, headers={\"Retry-After\": str(lockout)})"),

        # Phase 6
        ("h2", "5.2.6 Phase 6: Audit Logging และ Database Scaling"),
        ("p", "ทุกการกระทำของ admin เหตุการณ์ authentication และการดำเนินการที่เกี่ยวข้อง "
              "กับข้อมูลส่วนบุคคล ถูกบันทึกในระบบ audit log:"),
        ("code", "# src/hotel_guardrails/audit.py\n"
                 "class AuditActions:\n"
                 "    LOGIN_SUCCESS = \"auth.login.success\"\n"
                 "    LOGIN_FAILED = \"auth.login.failed\"\n"
                 "    SESSION_VIEWED = \"admin.session.viewed\"\n"
                 "    ROOM_STATUS_CHANGED = \"admin.room.status_changed\"\n"
                 "    BOOKING_STATUS_CHANGED = \"admin.booking.status_changed\"\n"
                 "    CHAT_OVERRIDE = \"admin.chat.override\"\n"
                 "    # ... รวม 26 ประเภท"),
        ("p", "การเชื่อมต่อฐานข้อมูลถูกปรับปรุงจากการสร้าง connection ใหม่ทุก request "
              "เป็น ThreadedConnectionPool:"),
        ("code", "# src/hotel_guardrails/database.py\n"
                 "def get_db_pool():\n"
                 "    global _db_pool\n"
                 "    if _db_pool is None:\n"
                 "        with _db_pool_lock:\n"
                 "            if _db_pool is None:\n"
                 "                _db_pool = pg_pool.ThreadedConnectionPool(\n"
                 "                    minconn=int(os.getenv(\"DB_POOL_MIN\", \"2\")),\n"
                 "                    maxconn=int(os.getenv(\"DB_POOL_MAX\", \"20\")),\n"
                 "                    dsn=os.getenv(\"DATABASE_URL\"),\n"
                 "                )\n"
                 "    return _db_pool"),

        # Phase 7
        ("h2", "5.2.7 Phase 7: Chat Scaling Primitives"),
        ("p", "Concurrency primitives จำนวน 5 ตัวถูกเพิ่มเข้ามาเพื่อรองรับผู้ใช้หลายคนพร้อมกัน:"),
        ("code", "# src/hotel_guardrails/chat_scaling.py\n"
                 "class LLMConcurrencyLimiter:\n"
                 "    \"\"\"Async semaphore with queue timeout for LLM calls.\"\"\"\n"
                 "    def __init__(self, max_concurrent: int, queue_timeout: float):\n"
                 "        self._semaphore = asyncio.Semaphore(max_concurrent)\n"
                 "        self.queue_timeout = queue_timeout\n"
                 "\n"
                 "    async def acquire(self):\n"
                 "        try:\n"
                 "            await asyncio.wait_for(\n"
                 "                self._semaphore.acquire(),\n"
                 "                timeout=self.queue_timeout,\n"
                 "            )\n"
                 "        except asyncio.TimeoutError:\n"
                 "            raise LLMQueueTimeout(\"LLM queue saturated — 503\")\n"
                 "\n"
                 "class SessionLockManager:\n"
                 "    \"\"\"Per-session asyncio.Lock with bounded LRU eviction.\"\"\"\n"
                 "    def get(self, session_id: str) -> asyncio.Lock:\n"
                 "        with self._lock:\n"
                 "            if session_id not in self._locks:\n"
                 "                self._locks[session_id] = asyncio.Lock()\n"
                 "            self._locks.move_to_end(session_id)\n"
                 "            return self._locks[session_id]\n"
                 "\n"
                 "class KnowledgeCache:\n"
                 "    \"\"\"TTL + LRU cache for RAG query results.\"\"\"\n"
                 "    def get(self, query: str):\n"
                 "        k = \" \".join(query.lower().split())\n"
                 "        entry = self._cache.get(k)\n"
                 "        if entry and time.time() - entry[0] <= self.ttl:\n"
                 "            self._hits += 1\n"
                 "            return entry[1]\n"
                 "        self._misses += 1\n"
                 "        return None"),
        ("figure", "[Figure 5.4: Scaling component pipeline — POST /chat → chat_rate_limiter (429) "
                   "→ session_lock → llm_limiter (503) → LangGraph → knowledge_cache → Ollama]"),

        # Phase 8
        ("h2", "5.2.8 Phase 8: การเพิ่มประสิทธิภาพ"),
        ("h3", "8a. การลบ Reranker"),
        ("p", "การ profiling พบว่า CrossEncoder reranker ทำให้ FastAPI event loop ถูก block:"),
        ("code", "# src/retrievers/hotel_knowledge/chains.py — ก่อนปรับปรุง\n"
                 "RERANKER_BACKEND = os.getenv(\"RERANKER_BACKEND\", \"qwen\")  # ~1-2s CPU/query\n"
                 "\n"
                 "# หลังปรับปรุง\n"
                 "RERANKER_BACKEND = os.getenv(\"RERANKER_BACKEND\", \"none\")  # ข้าม reranker\n"
                 "# Vector search alone ผ่าน 8/8 knowledge tests"),
        ("p", "ผลกระทบ: warm chat latency ลดลงจาก 18 วินาทีเหลือ 5 วินาที (เร็วขึ้น 3.6 เท่า)"),

        ("h3", "8b. การตัดทอน Prompt"),
        ("p", "System prompt ถูกลดขนาดจาก 5,500 ตัวอักษรเหลือ 2,800 ตัวอักษร "
              "โดยการลบส่วนที่ซ้ำซ้อนในการแปลภาษาไทย:"),
        ("code", "# src/agent/hotel_prompt.yaml — ก่อนปรับปรุง (5,500 chars)\n"
                 "# ทุกคำสั่งถูก duplicate เป็นภาษาไทย:\n"
                 "# \"DO NOT answer from memory.\"\n"
                 "# \"ห้ามตอบคำถามข้อมูลโรงแรมจากความจำ\"\n"
                 "\n"
                 "# หลังปรับปรุง (2,800 chars) — ลบ Thai duplicate\n"
                 "main_prompt: |\n"
                 "  You are a professional bilingual hotel assistant...\n"
                 "  ## Tools (ALWAYS use tools — never answer from memory)\n"
                 "  - `search_hotel_knowledge` → hotel info\n"
                 "  - `check_room_availability` → room types, pricing"),

        ("h3", "8c. การปรับแต่ง Ollama GPU"),
        ("p", "การ benchmark บน RTX 5080 (16 GB VRAM) เพื่อหาค่า parallelism ที่เหมาะสม:"),
        ("p", "TABLE_5_4_PLACEHOLDER"),
        ("p", "NUM_PARALLEL=2 ร่วมกับ FLASH_ATTENTION=1 ถูกเลือกเป็นค่าที่เหมาะสมที่สุด"),

        ("h3", "8d. การปิด Thinking Mode สำหรับ Local Model"),
        ("p", "Qwen3.5 สามารถแสดง <think> tags ได้โดยธรรมชาติ "
              "การเปิด explicit thinking mode เพิ่ม overhead โดยไม่จำเป็น:"),
        ("code", "# src/hotel_guardrails/config.py — ก่อนปรับปรุง\n"
                 "\"presets\": {\"temperature\": 0.3, \"max_tokens\": 4096, \"thinking\": True}\n"
                 "\n"
                 "# หลังปรับปรุง — native thinking, ไม่มี explicit overhead\n"
                 "\"presets\": {\"temperature\": 0.3, \"max_tokens\": 2048, \"thinking\": False}"),

        # --- 5.3 ---
        ("h1", "5.3 Prompt Engineering"),
        ("h2", "5.3.1 Primary Router Prompt"),
        ("p", "Router prompt ถูกปรับปรุงซ้ำหลายรอบเพื่อแก้ไขปัญหาการ routing ที่ผิดพลาด 2 กรณี "
              "ได้แก่ การ route คำสั่งยกเลิกการจอง และการจำแนกระหว่าง service กับ knowledge:"),
        ("code", "# src/hotel_guardrails/hotel_langgraph.py — Router prompt\n"
                 "## Your Role\n"
                 "Route every guest message to exactly ONE specialist:\n"
                 "1. **ToHotelBooking** — reservations, availability, check-in/out\n"
                 "   Examples: \"Is there a room?\", \"ยกเลิกการจอง\"\n"
                 "2. **ToHotelService** — room service, amenities, housekeeping\n"
                 "   Examples: \"I need extra towels\", \"จองสปา\"\n"
                 "3. **ToHotelKnowledge** — hotel info, facilities, dining\n"
                 "   Examples: \"What time is breakfast?\", \"รหัส WiFi\"\n"
                 "4. **HandleOtherTalk** — greetings, thanks, off-topic\n"
                 "   Examples: \"Hello\", \"ขอบคุณ\"\n"
                 "\n"
                 "IMPORTANT: \"cancel my booking\" → ToHotelBooking (NOT HandleOtherTalk)"),

        ("h2", "5.3.2 Knowledge Context Injection"),
        ("p", "ตำแหน่งของ RAG context ในการส่งข้อมูลให้ LLM มีความสำคัญมากสำหรับ 9B model "
              "การวาง knowledge context ก่อนคำถามผู้ใช้ทำให้โมเดลสรุปเนื้อหา context "
              "แทนที่จะตอบคำถาม:"),
        ("code", "# ผิด: knowledge context บดบังคำถาม\n"
                 "rag_prompt = ChatPromptTemplate.from_messages([\n"
                 "    (\"system\", main_prompt),\n"
                 "    (\"system\", f\"Context: {knowledge_result}\"),  # โมเดลสรุปตรงนี้\n"
                 "    (\"human\", last_user_message),\n"
                 "])\n"
                 "\n"
                 "# ถูกต้อง: คำถามผู้ใช้มาก่อน\n"
                 "rag_prompt = ChatPromptTemplate.from_messages([\n"
                 "    (\"system\", main_prompt),\n"
                 "    (\"human\", last_user_message),  # โมเดลโฟกัสตรงนี้\n"
                 "    (\"system\", f\"Use this hotel information:\\n{knowledge_result}\"),\n"
                 "])"),

        # --- 5.4 ---
        ("h1", "5.4 PII Redaction"),
        ("p", "ระบบ PII redaction ใช้ regular expressions เพื่อตรวจจับและปกปิดข้อมูลส่วนบุคคล "
              "ก่อนส่งให้ LLM ประมวลผล:"),
        ("code", "# src/hotel_guardrails/pii_redactor.py\n"
                 "PII_PATTERNS = {\n"
                 "    \"CREDIT_CARD\": re.compile(r\"\\b(?:\\d{4}[-\\s]?){3}\\d{4}\\b\"),\n"
                 "    \"THAI_NATIONAL_ID\": re.compile(r\"\\b\\d{1}-\\d{4}-\\d{5}-\\d{2}-\\d{1}\\b\"),\n"
                 "    \"PASSPORT\": re.compile(r\"\\b[A-Z]{1,2}\\d{6,9}\\b\"),\n"
                 "    \"PHONE_TH\": re.compile(r\"\\b0[689]\\d[-\\s]?\\d{3}[-\\s]?\\d{4}\\b\"),\n"
                 "    \"EMAIL\": re.compile(r\"\\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Za-z]{2,}\\b\"),\n"
                 "}\n"
                 "\n"
                 "def redact_pii(text, preserve_email=False):\n"
                 "    redacted = text\n"
                 "    for label, pattern in PII_PATTERNS.items():\n"
                 "        if label == \"EMAIL\" and preserve_email:\n"
                 "            continue\n"
                 "        redacted = pattern.sub(f\"[{label}]\", redacted)\n"
                 "    return redacted"),
        ("figure", "[Figure 5.3: PII redaction flow — ผู้ใช้พิมพ์ 'My card is 4111-1111-1111-1111' "
                   "→ regex ตรวจจับ CREDIT_CARD → LLM เห็น 'My card is [CREDIT_CARD]']"),

        # --- 5.5 ---
        ("h1", "5.5 ระบบ Human Escalation"),
        ("p", "ระบบ escalation ตรวจจับสถานการณ์ที่ต้องส่งต่อให้พนักงานจริง "
              "โดยใช้ 3 กลไก ได้แก่ sentiment detection, repetition detection และ high-value detection:"),
        ("code", "# src/hotel_guardrails/escalation.py\n"
                 "class EscalationMonitor:\n"
                 "    FRUSTRATION_EN = [\"terrible\", \"worst\", \"unacceptable\"]\n"
                 "    FRUSTRATION_TH = [\"แย่มาก\", \"ขอพูดกับผู้จัดการ\", \"ร้องเรียน\"]\n"
                 "\n"
                 "    def check_sentiment(self, message):\n"
                 "        for word in self.FRUSTRATION_EN + self.FRUSTRATION_TH:\n"
                 "            if word.lower() in message.lower():\n"
                 "                return True, f\"Frustration: '{word}'\"\n"
                 "        return False, None\n"
                 "\n"
                 "    def check_repetition(self, session_id, message):\n"
                 "        \"\"\"3+ messages with >70% similarity = guest stuck.\"\"\"\n"
                 "        ...\n"
                 "\n"
                 "    def check_high_value(self, context):\n"
                 "        \"\"\"Auto-escalate for bookings > 50,000 THB or Penthouse.\"\"\"\n"
                 "        ..."),

        # --- 5.6 ---
        ("h1", "5.6 การพัฒนา Frontend"),
        ("p", "Frontend ถูกพัฒนาด้วย Next.js 15 (App Router), Ant Design 5 สำหรับ UI components, "
              "Zustand สำหรับ global state management และ SWR สำหรับ server data fetching"),

        ("h2", "5.6.1 Chat Interface"),
        ("figure", "[Figure 5.5: Chat interface — แสดง hotel chatbot พร้อม message bubbles "
                   "คำตอบภาษาไทย/อังกฤษ และ typing indicator ระหว่าง SSE streaming]"),
        ("p", "Chat interface ใช้ Server-Sent Events (SSE) สำหรับ real-time token streaming:"),
        ("code", "// src/services/hotelAssistant/index.ts — SSE streaming client\n"
                 "export async function streamChat(message: string, sessionId: string,\n"
                 "                                  onToken: (token: string) => void) {\n"
                 "  const response = await fetch('/api/hotel/chat/stream', {\n"
                 "    method: 'POST',\n"
                 "    headers: { 'Content-Type': 'application/json' },\n"
                 "    body: JSON.stringify({ message, session_id: sessionId }),\n"
                 "  });\n"
                 "  const reader = response.body?.getReader();\n"
                 "  const decoder = new TextDecoder();\n"
                 "  while (true) {\n"
                 "    const { done, value } = await reader!.read();\n"
                 "    if (done) break;\n"
                 "    const chunk = decoder.decode(value);\n"
                 "    const lines = chunk.split('\\n').filter(l => l.startsWith('data: '));\n"
                 "    for (const line of lines) {\n"
                 "      const data = JSON.parse(line.slice(6));\n"
                 "      if (data.content) onToken(data.content);\n"
                 "      if (data.done) return data.session_id;\n"
                 "    }\n"
                 "  }\n"
                 "}"),

        ("h2", "5.6.2 Room Catalog และ Booking"),
        ("figure", "[Figure 5.6: Room catalog page — แสดงการ์ดประเภทห้องพร้อมรูปภาพ ราคา "
                   "สถานะห้องว่าง และปุ่ม Book Now]"),
        ("figure", "[Figure 5.7: Booking wizard — แบบฟอร์มหลายขั้นตอนพร้อม date picker "
                   "เลือกประเภทห้อง dynamic pricing และสรุปการจอง]"),
        ("code", "// src/app/hotel/bookings/features/BookingWizard.tsx\n"
                 "const BookingWizard: React.FC = () => {\n"
                 "  const [step, setStep] = useState<'dates'|'room'|'details'|'confirm'>('dates');\n"
                 "  return (\n"
                 "    <Steps current={['dates','room','details','confirm'].indexOf(step)}>\n"
                 "      <Step title=\"Dates\" />\n"
                 "      <Step title=\"Room Type\" />\n"
                 "      <Step title=\"Guest Details\" />\n"
                 "      <Step title=\"Confirm & Pay\" />\n"
                 "    </Steps>\n"
                 "  );\n"
                 "};"),

        ("h2", "5.6.3 Admin Dashboard"),
        ("figure", "[Figure 5.8: Admin dashboard — room occupancy chart, check-ins/check-outs, revenue summary]"),
        ("figure", "[Figure 5.9: Admin session monitor — active chat sessions พร้อม Takeover button]"),
        ("figure", "[Figure 5.10: Admin chat viewer — conversation history พร้อม admin reply input]"),
        ("code", "// src/app/hotel/admin/sessions/[sessionId]/page.tsx\n"
                 "export default function SessionDetailPage({ params }) {\n"
                 "  const { data } = useSWR(\n"
                 "    `/api/hotel/admin/sessions/${params.sessionId}/messages`,\n"
                 "    fetcher, { refreshInterval: 3000 }\n"
                 "  );\n"
                 "  return (\n"
                 "    <Card title={`Session: ${params.sessionId}`}>\n"
                 "      <MessageList messages={data?.messages} />\n"
                 "      <AdminReplyInput sessionId={params.sessionId} />\n"
                 "      <Space>\n"
                 "        <Button danger onClick={handleTakeover}>Take Over</Button>\n"
                 "        <Button onClick={handleRelease}>Release to Bot</Button>\n"
                 "      </Space>\n"
                 "    </Card>\n"
                 "  );\n"
                 "}"),

        ("h2", "5.6.4 Settings และ Model Switcher"),
        ("figure", "[Figure 5.11: Settings page — LLM backend, model name, temperature, available models]"),
        ("code", "// src/app/hotel/settings/page.tsx — Runtime model switching\n"
                 "const handleSwitch = async (backend: string, model: string) => {\n"
                 "  await fetch('/api/hotel/settings/llm', {\n"
                 "    method: 'PUT',\n"
                 "    headers: { 'Authorization': `Bearer ${token}` },\n"
                 "    body: JSON.stringify({ backend, model }),\n"
                 "  });\n"
                 "  mutate('/api/hotel/settings/llm');\n"
                 "  message.success(`Switched to ${model}`);\n"
                 "};"),

        ("h2", "5.6.5 Authentication UI"),
        ("figure", "[Figure 5.12: Login/Register modal — username, password, Login/Register tabs]"),
        ("code", "// src/app/hotel/features/AuthModal.tsx\n"
                 "const AuthModal: React.FC = ({ open, onClose }) => {\n"
                 "  const login = useHotelStore(s => s.login);\n"
                 "  const handleLogin = async (values) => {\n"
                 "    const res = await fetch('/api/hotel/auth/login', {\n"
                 "      method: 'POST',\n"
                 "      body: JSON.stringify(values),\n"
                 "    });\n"
                 "    if (res.ok) {\n"
                 "      const { access_token, user } = await res.json();\n"
                 "      login(access_token, user);\n"
                 "      onClose();\n"
                 "    }\n"
                 "  };\n"
                 "  return (\n"
                 "    <Modal open={open} onCancel={onClose} title=\"Login\">\n"
                 "      <Form onFinish={handleLogin}>\n"
                 "        <Form.Item name=\"username\"><Input placeholder=\"Username\" /></Form.Item>\n"
                 "        <Form.Item name=\"password\"><Input.Password /></Form.Item>\n"
                 "        <Button type=\"primary\" htmlType=\"submit\">Login</Button>\n"
                 "      </Form>\n"
                 "    </Modal>\n"
                 "  );\n"
                 "};"),

        ("h2", "5.6.6 State Management ด้วย Zustand"),
        ("code", "// src/store/hotel/slices/auth.ts\n"
                 "export const createAuthSlice: StateCreator<AuthSlice> = (set) => ({\n"
                 "  token: typeof window !== 'undefined' ? localStorage.getItem('auth_token') : null,\n"
                 "  user: null,\n"
                 "  isAdmin: false,\n"
                 "  login: (token, user) => {\n"
                 "    localStorage.setItem('auth_token', token);\n"
                 "    set({ token, user, isAdmin: user.role === 'admin' });\n"
                 "  },\n"
                 "  logout: () => {\n"
                 "    localStorage.removeItem('auth_token');\n"
                 "    set({ token: null, user: null, isAdmin: false });\n"
                 "  },\n"
                 "});"),

        ("h2", "5.6.7 API Proxy (หลีกเลี่ยง CORS)"),
        ("p", "ทุก backend call ผ่าน Next.js API route ที่ทำหน้าที่ proxy ไปยัง FastAPI server:"),
        ("code", "// src/app/api/hotel/[...path]/route.ts — Catch-all proxy\n"
                 "export async function POST(request, { params }) {\n"
                 "  const path = params.path.join('/');\n"
                 "  const body = await request.json();\n"
                 "  const headers = { 'Content-Type': 'application/json' };\n"
                 "  const auth = request.headers.get('Authorization');\n"
                 "  if (auth) headers['Authorization'] = auth;\n"
                 "  const res = await fetch(`${BACKEND_URL}/${path}`, {\n"
                 "    method: 'POST', headers, body: JSON.stringify(body),\n"
                 "  });\n"
                 "  return new Response(res.body, { status: res.status });\n"
                 "}"),
        ("p", "วิธีนี้ขจัดปัญหา CORS ทั้งหมด เนื่องจาก browser สื่อสารกับ Next.js origin เท่านั้น "
              "ส่วน Next.js ทำ server-to-server call ไปยัง FastAPI"),

        # --- 5.7 ---
        ("h1", "5.7 LLM Integration — Runtime Switching"),
        ("figure", "[Figure 5.1: Runtime model switching — Admin เรียก PUT /settings/llm "
                   "→ RuntimeLLMConfig singleton อัปเดตแบบ thread-safe "
                   "→ /chat call ถัดไปใช้ backend ใหม่]"),
        ("code", "# src/hotel_guardrails/hotel_langgraph.py — get_llm()\n"
                 "def get_llm(temperature=0.3, max_tokens=2048, streaming=False):\n"
                 "    runtime_config = get_runtime_llm_config()\n"
                 "\n"
                 "    if runtime_config.backend == LLMBackend.OLLAMA:\n"
                 "        return ChatOpenAI(\n"
                 "            model=runtime_config.ollama_model,\n"
                 "            openai_api_base=runtime_config.ollama_base_url,\n"
                 "            temperature=temperature,\n"
                 "            max_tokens=max_tokens,\n"
                 "        )\n"
                 "    else:  # OpenRouter cloud\n"
                 "        runtime_config.rate_limiter.wait_and_acquire()\n"
                 "        return ChatOpenAI(\n"
                 "            model=runtime_config.openrouter_model,\n"
                 "            openai_api_base=runtime_config.openrouter_base_url,\n"
                 "            temperature=temperature,\n"
                 "            max_tokens=max_tokens,\n"
                 "        )"),
        ("figure", "[Figure 5.2: Prompt template structure — main_prompt ถูกรวมกับ booking_flow "
                   "หรือ service_prompt ขึ้นอยู่กับ sub-agent ที่ถูก route]"),
    ]


# ============================================================================
# CONTENT: Chapter 6 — Testing and Evaluation (การทดสอบและประเมินผล)
# ============================================================================

def get_ch6_specs():
    return [
        ("h1", "6.1 วิธีการประเมิน"),
        ("h2", "6.1.1 การออกแบบการทดสอบ"),
        ("p", "กลยุทธ์การประเมินใช้สองแนวทางที่เสริมกัน:"),
        ("p", "1. การประเมินโมเดล (Model evaluation) — test cases จำนวน 25 กรณี "
              "ครอบคลุมงานใน domain โรงแรม โดยให้คะแนนความแม่นยำเทียบกับพฤติกรรมที่คาดหวัง"),
        ("p", "2. การทดสอบโครงสร้างพื้นฐาน (Infrastructure testing) — assertions จำนวน 193 ข้อ "
              "ที่ครอบคลุม auth, security, scaling และ API correctness"),

        ("h2", "6.1.2 Golden Dataset (25 Test Cases)"),
        ("p", "Golden dataset ครอบคลุม 5 หมวดหมู่ที่แสดงถึงขอบเขตการใช้งาน chatbot โรงแรมทั้งหมด:"),
        ("p", "TABLE_6_1_PLACEHOLDER"),
        ("p", "แต่ละ test case ประกอบด้วย:"),
        ("p", "• ข้อความ input (ภาษาไทยหรืออังกฤษ)"),
        ("p", "• Expected keywords ที่ต้องปรากฏในคำตอบ"),
        ("p", "• คำอธิบาย expected behavior"),
        ("p", "• การตรวจสอบภาษา (ไม่บังคับ)"),

        ("h2", "6.1.3 เกณฑ์การให้คะแนน"),
        ("p", "คำตอบถือว่าผ่านเมื่อเข้าเงื่อนไข 3 ข้อดังนี้:"),
        ("p", "1. Keyword score ≥ 50% — อย่างน้อยครึ่งหนึ่งของ expected keywords ปรากฏในคำตอบ"),
        ("p", "2. ภาษาถูกต้อง — หากกำหนดการตรวจสอบภาษา คำตอบต้องอยู่ในภาษาที่ระบุ"),
        ("p", "3. มีคำตอบ — คำตอบไม่ว่างเปล่าและไม่มี error เกิดขึ้น"),

        ("h1", "6.2 ผลการประเมินโมเดล"),
        ("h2", "6.2.1 ความแม่นยำโดยรวม"),
        ("p", "TABLE_6_2_PLACEHOLDER"),
        ("figure", "[Figure 6.1: Model accuracy comparison bar chart — Local 92% vs Cloud 100%]"),

        ("h2", "6.2.2 การวิเคราะห์รายหมวด"),
        ("p", "TABLE_6_3_PLACEHOLDER"),
        ("figure", "[Figure 6.2: Per-category accuracy heatmap — เขียว = 100%, เหลือง = 75%]"),

        ("h2", "6.2.3 การวิเคราะห์ Latency"),
        ("p", "TABLE_6_4_PLACEHOLDER"),
        ("figure", "[Figure 6.3: Latency distribution box plot — Local 9B มีการกระจายที่แน่นกว่า "
                   "p95 ต่ำกว่า Cloud มี median ต่ำกว่าแต่ p95 สูงกว่า]"),
        ("p", "Local model มี latency ที่สม่ำเสมอกว่า (p95 ต่ำกว่า) แม้ค่าเฉลี่ยจะใกล้เคียงกัน "
              "ซึ่งเป็นข้อได้เปรียบสำคัญสำหรับ user experience "
              "เนื่องจาก worst-case response time คาดเดาได้ดีกว่า"),

        ("h2", "6.2.4 Cohen's Kappa Inter-Model Agreement"),
        ("pb", "κ = 0.000"),
        ("p", "การตีความ: κ = 0.000 ไม่ได้หมายความว่าสองโมเดลไม่เห็นด้วยกัน "
              "แต่บ่งชี้ว่า observed agreement (92%) ใกล้เคียงกับค่าที่คาดหวังโดยบังเอิญ "
              "เนื่องจากทั้งสองโมเดลมีความแม่นยำสูงมาก ทำให้ตัวหาร (1 - p_e) เข้าใกล้ศูนย์"),
        ("p", "ในทางปฏิบัติ: ทั้งสองโมเดลให้ผลการตัดสินตรงกัน 23 จาก 25 test cases (agreement 92%)"),
        ("figure", "[Figure 6.4: Cohen's Kappa 2x2 confusion matrix — Both Pass: 23, "
                   "Local Fail/Cloud Pass: 2, Both Fail: 0]"),

        ("h2", "6.2.5 การวิเคราะห์ความล้มเหลว"),
        ("pb", "G03 (Thank you response)"),
        ("p", "Local 9B model ตอบอย่างสุภาพแต่ไม่มี expected keywords เพียงพอ "
              "(thank, welcome, help, pleasure) คำตอบเหมาะสมแต่ได้คะแนนต่ำกว่าเกณฑ์ "
              "keyword 50% กรณีนี้เป็น scoring artifact ไม่ใช่ข้อจำกัดของ capability"),
        ("pb", "E03 (Multi-room group booking)"),
        ("p", "Local model ไม่สามารถ route คำขอ \"I want to book 3 rooms for 10 people\" "
              "ไปยัง booking handler ได้ Cloud model จำแนก booking intent ได้ถูกต้อง "
              "กรณีนี้แสดงถึง routing capability gap ที่แท้จริงของ 9B model"),

        ("h1", "6.3 ผลการทดสอบโครงสร้างพื้นฐาน (193/193)"),
        ("h2", "6.3.1 สรุปชุดทดสอบ"),
        ("p", "TABLE_6_5_PLACEHOLDER"),
        ("figure", "[Figure 6.2: Infrastructure test coverage pie chart — Auth 37%, "
                   "Hardening 20%, Audit+Scaling 24%, Chat Scaling 19%]"),

        ("h2", "6.3.2 การตรวจสอบที่สำคัญ"),
        ("p", "• Access control: ทุก admin/dashboard endpoint ถูกทดสอบด้วย 3 สถานการณ์ "
              "(ไม่มี token → 401, user token → 403, admin token → 200)"),
        ("p", "• Rate limiting: Per-IP และ per-username login limits ส่ง 429 พร้อม Retry-After header"),
        ("p", "• Token revocation: Token ที่ logged-out ถูกปฏิเสธในการใช้งานครั้งถัดไป"),
        ("p", "• Password-change invalidation: การเปลี่ยนรหัสผ่านทำให้ token ทั้งหมดก่อนหน้าใช้ไม่ได้ "
              "(persistent ผ่าน password_changed_at ในฐานข้อมูล ทำงานได้แม้ server restart)"),
        ("p", "• Concurrent chat: 5 parallel /chat requests ไปยัง session ต่างกัน "
              "เสร็จสิ้นภายใน 3 วินาที (ไม่ถูก serialized)"),

        ("h1", "6.4 ผลการเพิ่มประสิทธิภาพ"),
        ("h2", "6.4.1 Before/After Benchmarks"),
        ("p", "TABLE_6_6_PLACEHOLDER"),
        ("figure", "[Figure 6.3: Before/after optimization chart — warm chat latency 18s→5s, "
                   "concurrent 5-session 90s→3s]"),

        ("h2", "6.4.2 การปรับแต่ง Ollama GPU (RTX 5080, 16 GB)"),
        ("p", "TABLE_6_7_PLACEHOLDER"),
        ("p", "ข้อค้นพบสำคัญ: OLLAMA_NUM_PARALLEL แบ่ง throughput คงที่ของ GPU "
              "ให้กับ active sequences ทุกตัว การลดจาก 4 เป็น 2 ทำให้ per-request speed "
              "เพิ่มขึ้นเป็นสองเท่า ซึ่งเป็น trade-off ที่ดีกว่าสำหรับ interactive hotel chat"),

        ("h2", "6.4.3 Scaling Metrics Under Load"),
        ("p", "TABLE_6_8_PLACEHOLDER"),
    ]


# ============================================================================
# CONTENT: Chapter 7 — Discussion (อภิปรายผล)
# ============================================================================

def get_ch7_specs():
    return [
        ("h1", "7.1 การบรรลุวัตถุประสงค์"),
        ("p", "TABLE_7_1_PLACEHOLDER"),

        ("h1", "7.2 ข้อแลกเปลี่ยนระหว่าง Local กับ Cloud Model"),
        ("h2", "7.2.1 สามเหลี่ยม Cost-Accuracy-Latency"),
        ("figure", "[Figure 7.1: Cost-accuracy-latency triangle — Local 9B อยู่ที่ "
                   "(ต้นทุนต่ำ, 92%, medium latency) Cloud อยู่ที่ (ต้นทุนปานกลาง, 100%, variable latency)]"),
        ("p", "TABLE_7_2_PLACEHOLDER"),

        ("h2", "7.2.2 ผลกระทบด้านความเป็นส่วนตัว"),
        ("p", "สำหรับโรงแรมที่ประมวลผลข้อมูลส่วนบุคคลของแขก (ชื่อ หมายเลขหนังสือเดินทาง "
              "บัตรเครดิต) การรัน LLM แบบ local หมายความว่า PII ไม่เคยออกจากเครือข่ายของโรงแรม "
              "PII redactor เป็นชั้นป้องกันเชิงลึก (defense-in-depth) แต่ local model ขจัดความเสี่ยงนี้ "
              "โดยสิ้นเชิง ซึ่งสอดคล้องกับพระราชบัญญัติคุ้มครองข้อมูลส่วนบุคคล (PDPA) ของประเทศไทย "
              "และ General Data Protection Regulation (GDPR) ของสหภาพยุโรป"),

        ("h2", "7.2.3 จุดอ่อนของ 9B Model"),
        ("p", "ความล้มเหลว 2 กรณีของ local model เผยให้เห็นข้อจำกัดของ 9B model:"),
        ("p", "1. การวิเคราะห์ intent ที่ซับซ้อนแบบ multi-entity — คำขอ \"3 rooms for 10 people\" "
              "ต้องการความสามารถในการเข้าใจปริมาณ จัดกลุ่มเป็นหลาย booking และ route ได้ถูกต้อง "
              "ซึ่ง 9B model ขาดความลึกในการ reasoning"),
        ("p", "2. ความหลากหลายของคำตอบ — สำหรับ \"thank you\" โมเดล 9B ตอบอย่างสุภาพ "
              "แต่ใช้คำศัพท์ต่างจากที่คาดหวัง ซึ่งเป็นข้อจำกัดของระบบ scoring มากกว่าตัวโมเดล"),
        ("pb", "คำแนะนำ: ใช้ local 9B model สำหรับ 92% ของ routine queries "
               "และ fallback ไปยัง cloud สำหรับ complex multi-step operations"),

        ("h1", "7.3 ประสิทธิภาพของ RAG"),
        ("h2", "7.3.1 คุณภาพการค้นหา"),
        ("p", "RAG pipeline มีความแม่นยำ 100% ในทั้ง 8 knowledge test cases สำหรับทั้งสองโมเดล:"),
        ("p", "• ฐานความรู้ของโรงแรมมีโครงสร้างที่ดี (markdown 10 ฉบับที่มี headings ชัดเจน)"),
        ("p", "• Embedding model (qwen3-embedding-8b, 4096 dimensions) รองรับสองภาษาได้ดี"),
        ("p", "• ขนาด chunk ที่คำนวณอัตโนมัติ (80% ของ token limit) ป้องกันการสูญเสียข้อมูล"),

        ("h2", "7.3.2 ผลกระทบจากการลบ Reranker"),
        ("p", "การปิดใช้งาน CrossEncoder reranker ไม่มีผลกระทบต่อความแม่นยำ "
              "แต่ลด latency ลง 3.6 เท่า ผลลัพธ์นี้ขัดกับคำแนะนำทั่วไปในเอกสารวิชาการ "
              "แต่อธิบายได้จากลักษณะเฉพาะของ domain:"),
        ("p", "• ฐานความรู้มีเพียง 10 เอกสาร (ไม่ใช่หลายพัน)"),
        ("p", "• เอกสารแต่ละฉบับมีหัวข้อที่แตกต่างชัดเจน (spa ≠ dining ≠ policies)"),
        ("p", "• Embedding model ได้รับการ optimize สำหรับสองภาษาอยู่แล้ว"),
        ("p", "สำหรับฐานความรู้ที่ใหญ่กว่าหรือมีโครงสร้างน้อยกว่า reranking น่าจะมีความจำเป็น"),

        ("h2", "7.3.3 ประสิทธิภาพของ Knowledge Cache"),
        ("p", "Knowledge cache (500 entries, 5-minute TTL) มี hit rate 76% "
              "ในระหว่างการทดสอบต่อเนื่อง หมายความว่า 3 ใน 4 knowledge queries "
              "ถูกให้บริการจาก memory (~1ms) แทนที่จะผ่าน Qdrant (~500ms) "
              "TTL 5 นาทีทำให้การอัปเดตฐานความรู้เผยแพร่ได้อย่างรวดเร็ว"),

        ("h1", "7.4 การตัดสินใจด้านสถาปัตยกรรมและข้อแลกเปลี่ยน"),
        ("h2", "7.4.1 LangGraph เทียบกับ Orchestrator ทางเลือก"),
        ("figure", "[Figure 7.2: Orchestrator comparison matrix]"),
        ("p", "TABLE_7_3_PLACEHOLDER"),
        ("p", "LangGraph ถูกเลือกเนื่องจาก checkpointed state persistence "
              "(conversation memory รอดจาก server restarts) และ time-travel debugging "
              "(admin สามารถย้อนกลับและ replay conversations จาก checkpoint ใดก็ได้)"),

        ("h2", "7.4.2 In-Memory vs Redis Scaling Primitives"),
        ("p", "Scaling primitives ทั้ง 5 ตัว เป็นแบบ in-memory และ per-process "
              "ซึ่งเหมาะสมกับ single-worker Docker deployment แต่มีข้อจำกัด:"),
        ("p", "• Primitives จะ reset เมื่อ server restart"),
        ("p", "• หลาย workers จะไม่แชร์ rate limit counters"),
        ("p", "• Token blocklist entries จะหายเมื่อ restart"),
        ("p", "ข้อยกเว้นคือ password-change invalidation ที่ persist ผ่าน "
              "users.password_changed_at column ในฐานข้อมูล "
              "ทำให้ทำงานได้ข้ามการ restart และหลาย workers"),
        ("p", "สำหรับ production horizontal scaling ควรเปลี่ยนเป็น Redis-backed equivalents"),

        ("h1", "7.5 ข้อจำกัด"),
        ("p", "1. ไม่มีการเชื่อมต่อระบบชำระเงินจริง — payment links เป็นแบบจำลอง "
              "ใน production จำเป็นต้องใช้ Stripe หรือ PromptPay"),
        ("p", "2. ฐานความรู้ภาษาเดียว — เอกสารเป็นแบบสองภาษาภายในแต่ละไฟล์ "
              "ระบบไม่รองรับการเพิ่มภาษาที่สามโดยไม่ปรับโครงสร้าง"),
        ("p", "3. Single-worker deployment — in-memory scaling primitives "
              "ไม่ประสานงานข้ามหลาย workers"),
        ("p", "4. ไม่มี voice interface — เฉพาะข้อความ ต้องใช้ Vapi หรือ Twilio"),
        ("p", "5. รองรับโรงแรมเดียว — ต้องการ tenant isolation สำหรับ multi-property"),
        ("p", "6. ต้องใช้ GPU สำหรับ local model — 9B model ต้องการ NVIDIA GPU สมัยใหม่ "
              "(ทดสอบบน RTX 5080) CPU-only deploy ไม่เหมาะสม"),
    ]


# ============================================================================
# CONTENT: Chapter 8 — Conclusion (สรุปผลการวิจัย)
# ============================================================================

def get_ch8_specs():
    return [
        ("h1", "8.1 สรุปผลงานวิจัย"),
        ("p", "วิทยานิพนธ์นี้นำเสนอการออกแบบ พัฒนา และประเมินผลระบบ AI virtual assistant "
              "เต็มรูปแบบสำหรับโรงแรม The Grand Horizon Hotel ซึ่งเป็นโรงแรมหรู 5 ดาวในประเทศไทย "
              "ระบบแสดงให้เห็นว่าสถาปัตยกรรม multi-agent ที่ใช้ LLM สมัยใหม่สามารถจัดการ "
              "งานโรงแรมจริงได้อย่างมีประสิทธิภาพ ตั้งแต่การตอบคำถามแขกเกี่ยวกับเวลาอาหารเช้า "
              "ไปจนถึงการสร้างการจองที่เชื่อมต่อกับฐานข้อมูล โดยรักษาระดับความปลอดภัย "
              "และ scalability ระดับ production"),
        ("pb", "ผลงานหลัก:"),
        ("p", "1. สถาปัตยกรรม multi-agent ด้วย LangGraph ที่ประกอบด้วย 4 specialized sub-agents "
              "(booking, service, knowledge, general conversation) ทำหน้าที่ route คำขอแขก "
              "ตาม intent classification โดย LangGraph state machine มี checkpointed conversation memory, "
              "tool-calling loops และความสามารถ time-travel debugging"),
        ("p", "2. RAG pipeline สำหรับฐานความรู้โรงแรมใช้ Qdrant vector embeddings "
              "มีความแม่นยำ 100% ในการตอบคำถามข้อมูลโรงแรม (8/8 test cases) ทั้ง local และ cloud "
              "ข้อค้นพบที่ว่า reranking ไม่จำเป็นสำหรับฐานความรู้ขนาดเล็กที่มีโครงสร้างดี "
              "ขัดกับคำแนะนำทั่วไปของ RAG และให้แนวทางปฏิบัติสำหรับการ deploy ในลักษณะเดียวกัน"),
        ("p", "3. ระบบ fullstack ระดับ production ประกอบด้วย JWT authentication พร้อม bcrypt hashing, "
              "rate limiting, account lockout, token blocklist และ audit logging (193/193 tests ผ่าน); "
              "scaling primitives 5 ตัวที่ทำให้รองรับผู้ใช้พร้อมกันบน GPU เดียว; "
              "runtime switching ระหว่าง local Ollama และ cloud OpenRouter โดยไม่ต้อง restart server; "
              "Next.js 15 frontend พร้อม Ant Design admin dashboard และ SSE chat streaming"),
        ("p", "4. การเปรียบเทียบเชิงประจักษ์ระหว่าง local vs cloud LLM แสดงให้เห็นว่า "
              "9B model ที่รัน local มีความแม่นยำ 92% (23/25) เทียบกับ 100% ของ cloud model "
              "โดย local model จัดการงานโรงแรมทั่วไปได้ทั้งหมด (knowledge: 100%, booking: 100%, "
              "bilingual: 100%) ค่าเฉลี่ย latency ใกล้เคียงกัน (9.5s vs 8.6s) "
              "โดย local model มี tail latency ที่คาดเดาได้ดีกว่า (p95: 18s vs 38s)"),
        ("p", "5. กรอบการทดสอบและประเมินผลที่สมบูรณ์ ประกอบด้วย 25 domain-specific test cases "
              "พร้อม keyword scoring, language detection และ Cohen's Kappa inter-model agreement "
              "ที่สามารถ reproduce ได้สำหรับงานวิจัย hotel chatbot ในอนาคต"),

        ("h1", "8.2 งานในอนาคต"),
        ("h2", "8.2.1 Voice Integration"),
        ("p", "การผสมผสาน speech-to-text (เช่น Whisper) และ text-to-speech (เช่น Vapi, Twilio) "
              "จะขยายขอบเขตของ assistant ไปยังการโทรศัพท์และอุปกรณ์เสียงในห้องพัก "
              "pipeline ที่ใช้ข้อความปัจจุบันสามารถเป็น backbone สำหรับ language understanding ได้"),

        ("h2", "8.2.2 Payment Gateway Integration"),
        ("p", "การแทนที่ mock payment links ด้วยระบบชำระเงินจริง "
              "(Stripe สำหรับแขกต่างชาติ, PromptPay สำหรับแขกไทย) "
              "จะทำให้กระบวนการจองครบถ้วนตั้งแต่ต้นจนจบ"),

        ("h2", "8.2.3 MCP (Model Context Protocol) Integration"),
        ("p", "Model Context Protocol ของ Anthropic จะช่วยให้ AI assistant เชื่อมต่อโดยตรง "
              "กับระบบ PMS ของโรงแรม (Opera, Cloudbeds) ในฐานะ MCP servers "
              "ลดความจำเป็นในการสร้าง custom API integration สำหรับแต่ละ vendor"),

        ("h2", "8.2.4 Domain Fine-tuning"),
        ("p", "การ fine-tune 9B model ด้วยข้อมูลสนทนาเฉพาะ domain โรงแรม "
              "(booking dialogues, service requests, FAQ pairs) อาจปิดช่องว่างความแม่นยำ 8% "
              "กับ cloud model โดยใช้ LoRA (Low-Rank Adaptation) บน GPU เดียว"),

        ("h2", "8.2.5 Multi-Property Support"),
        ("p", "การเพิ่ม tenant isolation ในฐานข้อมูล (hotel_id foreign key ในทุกตาราง) "
              "และ Qdrant collection naming แยกตามโรงแรม "
              "จะขยายระบบให้รองรับ hotel chains ที่จัดการหลาย property จาก deployment เดียว"),

        ("h2", "8.2.6 Redis-Backed Scaling"),
        ("p", "การย้าย scaling primitives 5 ตัวจาก in-memory เป็น Redis-backed equivalents "
              "จะช่วยให้ horizontal scaling ข้ามหลาย workers และ containers "
              "รองรับ concurrent users มากขึ้นโดยไม่ต้องเปลี่ยน application logic"),

        ("h2", "8.2.7 Automated Quality Monitoring"),
        ("p", "การ implement continuous evaluation — รัน 25-case test suite อัตโนมัติตามกำหนดเวลา "
              "และแจ้งเตือนเมื่อความแม่นยำลดลง — จะให้ quality assurance สำหรับ production "
              "สคริปต์ eval_model_comparison.py เป็นพื้นฐานสำหรับ monitoring นี้"),
    ]


# ============================================================================
# Figure map for re-attachment
# ============================================================================

FIGURE_MAP = {
    "1.1": "Fig_1.1_System_Context.png",
    "2.1": "Fig_2.1_Chatbot_Taxonomy.png",
    "2.2": "Fig_2.2_LangGraph_Concept.png",
    "2.3": "Fig_2.3_RAG_Pipeline.png",
    "3.1": "Fig_3.1_Framework_Comparison.png",
    "4.1": "Fig_4.1_System_Architecture.png",
    "4.2": "Fig_4.2_LangGraph_State_Machine.png",
    "4.3": "Fig_4.3_ER_Diagram.png",
    "4.4": "Fig_4.4_RAG_Pipeline_Detail.png",
    "4.5": "Fig_4.5_JWT_Auth_Flow.png",
    "4.6": "Fig_4.6_Access_Control_Matrix.png",
    "4.7": "Fig_4.7_Docker_Topology.png",
    "5.1": "Fig_5.1_Runtime_Model_Switch.png",
    "5.2": "Fig_5.2_Prompt_Structure.png",
    "5.3": "Fig_5.3_PII_Redaction.png",
    "5.4": "Fig_5.4_Scaling_Pipeline.png",
    "6.1": "Fig_6.1_Accuracy_Comparison.png",
    "6.2": "Fig_6.2_Test_Coverage_Pie.png",
    "6.3": "Fig_6.3_Before_After_Optimization.png",
    "6.4": "Fig_6.4_Kappa_Matrix.png",
    "7.1": "Fig_7.1_Cost_Accuracy_Latency.png",
    "7.2": "Fig_7.2_Orchestrator_Comparison.png",
}


# ============================================================================
# Main
# ============================================================================

def main():
    print("=" * 60)
    print("Patching thesis with Thai translations")
    print("  Abstract + CH5 + CH6 + CH7 + CH8")
    print("=" * 60)

    if not SRC.exists():
        print(f"ERROR: {SRC} not found")
        sys.exit(1)

    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    body = doc.element.body

    # ---- 1. Find chapter boundaries ----
    print("\n1. Scanning document for chapter boundaries...")
    chapters = find_chapter_indices(doc)
    for ch, idx in sorted(chapters.items(), key=lambda x: str(x[0])):
        text_preview = doc.paragraphs[idx].text[:60]
        print(f"  {ch}: P{idx} — {text_preview}")

    if not chapters:
        print("ERROR: No chapter headings found!")
        sys.exit(1)

    # Build ordered list of integer chapters
    int_chs = sorted([(k, v) for k, v in chapters.items() if isinstance(k, int)],
                     key=lambda x: x[1])

    def get_next_elem(ch_num):
        """Get the element of the next chapter heading after ch_num."""
        found = False
        for k, v in int_chs:
            if found:
                return doc.paragraphs[v]._element
            if k == ch_num:
                found = True
        if "appendix" in chapters:
            return doc.paragraphs[chapters["appendix"]]._element
        return None

    # ---- 2. Replace chapters in REVERSE order (8, 7, 6, 5) ----
    for ch_num, specs_fn, label in [
        (8, get_ch8_specs, "CH8 Conclusion"),
        (7, get_ch7_specs, "CH7 Discussion"),
        (6, get_ch6_specs, "CH6 Testing & Evaluation"),
        (5, get_ch5_specs, "CH5 Implementation"),
    ]:
        if ch_num not in chapters:
            print(f"\n  [SKIP] {label} — heading not found")
            continue

        start_elem = doc.paragraphs[chapters[ch_num]]._element
        end_elem = get_next_elem(ch_num)

        print(f"\n2.{ch_num}. Replacing {label}...")
        if end_elem is not None:
            removed = remove_between_elements(body, start_elem, end_elem)
            print(f"  Removed {removed} old elements")
        else:
            # Last chapter — remove everything after it up to end of body
            children = list(body)
            start_pos = children.index(start_elem)
            # Keep last element if it's sectPr
            end_pos = len(children)
            if children[-1].tag.endswith('}sectPr'):
                end_pos = len(children) - 1
            to_remove = children[start_pos + 1:end_pos]
            for elem in to_remove:
                body.remove(elem)
            removed = len(to_remove)
            print(f"  Removed {removed} old elements (last chapter)")

        nodes = build_nodes(doc, specs_fn())
        insert_after_element(body, start_elem, nodes)
        print(f"  Inserted {len(nodes)} Thai paragraphs")

    # ---- 3. Replace Thai Abstract ----
    print("\n3. Replacing Thai Abstract (บทคัดย่อ)...")
    if "abstract_th" in chapters:
        abstract_elem = doc.paragraphs[chapters["abstract_th"]]._element
        # Find the next TU_Chapter heading (English abstract or acknowledgment)
        next_key = "abstract_en"
        if next_key not in chapters:
            # Fall back to finding the next TU_Chapter after abstract
            next_key = None
            for i in range(chapters["abstract_th"] + 1, len(doc.paragraphs)):
                if doc.paragraphs[i].style and doc.paragraphs[i].style.name == "TU_Chapter":
                    next_key = i
                    break
        if next_key is not None:
            if isinstance(next_key, str):
                end_elem = doc.paragraphs[chapters[next_key]]._element
            else:
                end_elem = doc.paragraphs[next_key]._element
            removed = remove_between_elements(body, abstract_elem, end_elem)
            print(f"  Removed {removed} old abstract elements")
        nodes = build_nodes(doc, get_abstract_specs())
        insert_after_element(body, abstract_elem, nodes)
        print(f"  Inserted {len(nodes)} Thai abstract paragraphs")
    else:
        print("  [SKIP] Thai abstract heading not found")

    # ---- 4. Insert tables at placeholders ----
    print("\n4. Inserting Word tables at placeholders...")
    table_count = 0
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for table_id, (headers, rows) in TABLES.items():
            placeholder = f"{table_id}_PLACEHOLDER"
            if text == placeholder:
                try:
                    create_word_table(doc, headers, rows, p._element)
                    body.remove(p._element)
                    table_count += 1
                    print(f"  [OK] {table_id}")
                except Exception as e:
                    print(f"  [FAIL] {table_id}: {e}")
                break
    print(f"  Total tables: {table_count}")

    # ---- 5. Attach PNG figures ----
    print("\n5. Attaching PNG figures...")
    attached = 0
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if "[Figure" in text or ("Figure" in text and ("Insert" in text or ":" in text)):
            m = re.search(r'Figure\s+(\d+\.\d+)', text)
            if m:
                fig_id = m.group(1)
                fig_file = FIGURE_MAP.get(fig_id)
                if fig_file:
                    fig_path = FIGURES_DIR / fig_file
                    if fig_path.exists():
                        try:
                            insert_image_paragraph(doc, body, p._element,
                                                   fig_path, width_inches=5.0)
                            attached += 1
                            print(f"  [OK] Figure {fig_id}")
                        except Exception as e:
                            print(f"  [FAIL] Figure {fig_id}: {e}")
    print(f"  Total figures attached: {attached}")

    # ---- 6. Save ----
    doc.save(str(OUT))
    sz = os.path.getsize(str(OUT)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUT} ({sz:.0f} KB)")
    print("=" * 60)


if __name__ == "__main__":
    main()
