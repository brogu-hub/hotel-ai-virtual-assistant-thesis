#!/usr/bin/env python3
"""
patch_v5_enhancements.py
========================
1. Fix fonts: strip run-level font overrides on non-code paragraphs so TH Sarabun New inherits
2. Add CH5: use case diagrams, sequence diagrams, Ollama/OpenRouter API samples, tool calling
3. Add Python 3.10 to CH2 Literature Review
4. Update Figure 6.1 caption for per-category chart

Input:  thesis/thesis_final_v4.docx
Output: thesis/thesis_final_v5.docx
"""
import os, sys, re, shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final_v4.docx")
OUT = Path("thesis/thesis_final_v5.docx")
FIGURES_DIR = Path("thesis/figures")

# ============================================================================
# 1. FONT FIX — strip all run-level rFonts except on Consolas (code) runs
# ============================================================================
def fix_fonts(doc):
    """Remove run-level font overrides on non-code paragraphs.
    TU_ styles already define TH Sarabun New; our injected runs override it."""
    fixed = 0
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        for run in p.runs:
            rpr = run._element.find(qn('w:rPr'))
            if rpr is None:
                continue
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                continue
            # Keep Consolas (code blocks) — remove everything else
            ascii_font = rfonts.get(qn('w:ascii'), '')
            if ascii_font == 'Consolas':
                continue
            # Remove the rFonts element so style font (TH Sarabun New) takes over
            rpr.remove(rfonts)
            fixed += 1
    return fixed


# ============================================================================
# Helpers (same as prior scripts, but make_para does NOT set fonts on non-code)
# ============================================================================
def make_para(doc, text, style_name, bold=False, italic=False,
              font_name=None, font_size=None, alignment=None):
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
            rPr.append(OxmlElement('w:bCs'))
        if italic:
            rPr.append(OxmlElement('w:i'))
            rPr.append(OxmlElement('w:iCs'))
        # Only set font for code (Consolas)
        if font_name:
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), font_name)
            rf.set(qn('w:hAnsi'), font_name)
            rf.set(qn('w:cs'), font_name)
            rPr.append(rf)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size * 2)))
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(sz)
            rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p

def make_code_para(doc, text):
    return make_para(doc, text, "TU_Paragraph_Normal", font_name="Consolas", font_size=10)

def build_nodes(doc, specs):
    nodes = []
    for spec in specs:
        kind, text = spec[0], spec[1] if len(spec) > 1 else ""
        if kind == "h1":
            nodes.append(make_para(doc, text, "TU_Sub-heading 1", bold=True))
        elif kind == "h2":
            nodes.append(make_para(doc, text, "TU_Sub-heading 2", bold=True))
        elif kind == "h3":
            nodes.append(make_para(doc, text, "TU_Sub-heading 3", bold=True))
        elif kind == "p":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal"))
        elif kind == "pb":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", bold=True))
        elif kind == "pi":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", italic=True))
        elif kind == "code":
            for line in text.split("\n"):
                nodes.append(make_code_para(doc, line if line else " "))
        elif kind == "figure":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", italic=True, alignment="center"))
    return nodes

def insert_after_element(body, ref_elem, nodes):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for j, node in enumerate(nodes):
        parent.insert(idx + 1 + j, node)

def find_chapter_indices(doc):
    chapters = {}
    HEADING_MAP = {
        "introduction": 1, "literature review": 2, "methodology": 3,
        "system design": 4, "implementation": 5, "testing and evaluation": 6,
        "discussion": 7, "conclusion": 8,
    }
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            lower = text.lower()
            for ht, cn in HEADING_MAP.items():
                if ht in lower and cn not in chapters:
                    chapters[cn] = i
                    break
    return chapters

def find_paragraph_containing(doc, text_fragment, start=0, end=None):
    """Find paragraph index containing text_fragment."""
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, end):
        if text_fragment in doc.paragraphs[i].text:
            return i
    return None

def insert_image_paragraph(doc, body, ref_element, image_path, width_inches=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    elem = p._element
    parent = elem.getparent()
    parent.remove(elem)
    ref_parent = ref_element.getparent()
    idx = list(ref_parent).index(ref_element)
    ref_parent.insert(idx + 1, elem)
    return elem


# ============================================================================
# 2. CH5 ADDITIONS — inserted AFTER specific anchor paragraphs
# ============================================================================

# --- After "ระบบมี hotel-specific tools จำนวน 15 ตัว ดังตารางที่ 5.3" ---
TOOL_CALLING_SAMPLE = [
    ("p", ""),
    ("pb", "ตัวอย่างการเรียกใช้ tool: check_room_availability"),
    ("p", "เมื่อแขกถามว่า \"มีห้องว่างวันที่ 15-17 เมษายนไหม\" LangGraph agent จะ route "
          "ไปยัง booking sub-agent ซึ่ง LLM จะ emit tool call ในรูปแบบ JSON:"),
    ("code", "# LLM emits structured tool call:\n"
             "{\n"
             "  \"name\": \"check_room_availability\",\n"
             "  \"args\": {\n"
             "    \"check_in_date\": \"2026-04-15\",\n"
             "    \"check_out_date\": \"2026-04-17\",\n"
             "    \"room_type\": null\n"
             "  }\n"
             "}"),
    ("p", "LangGraph ToolNode ดักจับ tool call นี้และ execute Python function:"),
    ("code", "# src/agent/hotel_tools.py — check_room_availability\n"
             "@tool\n"
             "def check_room_availability(check_in_date: str, check_out_date: str,\n"
             "                            room_type: Optional[str] = None) -> str:\n"
             "    query = \"\"\"\n"
             "        SELECT r.room_number, rt.name, rt.base_price\n"
             "        FROM rooms r\n"
             "        JOIN room_types rt ON r.room_type_id = rt.room_type_id\n"
             "        WHERE r.status = 'available'\n"
             "        AND r.room_id NOT IN (\n"
             "            SELECT room_id FROM reservations\n"
             "            WHERE status NOT IN ('cancelled', 'no_show')\n"
             "            AND check_in_date < %s AND check_out_date > %s\n"
             "        )\n"
             "        ORDER BY rt.base_price LIMIT 10\n"
             "    \"\"\"\n"
             "    with get_db_connection() as conn:\n"
             "        with conn.cursor() as cur:\n"
             "            cur.execute(query, (check_out_date, check_in_date))\n"
             "            rooms = cur.fetchall()\n"
             "    # Format bilingual response\n"
             "    return format_rooms(rooms)"),
    ("p", "ผลลัพธ์ของ tool ถูกส่งกลับเป็น ToolMessage ใน LangGraph state "
          "ซึ่ง booking sub-agent จะรับไปสร้างคำตอบให้แขก "
          "หาก LLM ต้องการเรียก tool เพิ่ม (เช่น create_reservation) "
          "จะเกิด cyclic loop กลับไปที่ booking node อีกรอบ"),
]

# --- After Phase 4 Docker section, before Phase 5 ---
OLLAMA_OPENROUTER_SAMPLE = [
    ("p", ""),
    ("pb", "การเรียก LLM: Ollama (Local) vs OpenRouter (Cloud)"),
    ("p", "ทั้ง Ollama และ OpenRouter ใช้ OpenAI-compatible API format เหมือนกัน "
          "ทำให้สลับได้ทันทีผ่าน RuntimeLLMConfig singleton:"),
    ("p", ""),
    ("pb", "Ollama Local API Call (port 11435):"),
    ("code", "# HTTP Request → Ollama (local GPU)\n"
             "POST http://localhost:11435/v1/chat/completions\n"
             "Content-Type: application/json\n"
             "\n"
             "{\n"
             "  \"model\": \"fredrezones55/qwen3.5-opus:9b\",\n"
             "  \"messages\": [\n"
             "    {\"role\": \"system\", \"content\": \"You are a hotel assistant...\"},\n"
             "    {\"role\": \"user\", \"content\": \"What time is breakfast?\"}\n"
             "  ],\n"
             "  \"temperature\": 0.3,\n"
             "  \"max_tokens\": 2048,\n"
             "  \"tools\": [{\"type\": \"function\", \"function\": {\n"
             "    \"name\": \"search_hotel_knowledge\",\n"
             "    \"parameters\": {\"type\": \"object\", \"properties\": {\n"
             "      \"query\": {\"type\": \"string\"}\n"
             "    }}\n"
             "  }}]\n"
             "}"),
    ("code", "# Response (tool call)\n"
             "{\n"
             "  \"choices\": [{\n"
             "    \"message\": {\n"
             "      \"role\": \"assistant\",\n"
             "      \"tool_calls\": [{\n"
             "        \"function\": {\n"
             "          \"name\": \"search_hotel_knowledge\",\n"
             "          \"arguments\": \"{\\\"query\\\": \\\"breakfast hours\\\"}\"\n"
             "        }\n"
             "      }]\n"
             "    }\n"
             "  }]\n"
             "}"),
    ("p", ""),
    ("pb", "OpenRouter Cloud API Call:"),
    ("code", "# HTTP Request → OpenRouter (cloud)\n"
             "POST https://openrouter.ai/api/v1/chat/completions\n"
             "Authorization: Bearer sk-or-v1-xxx\n"
             "HTTP-Referer: https://grand-horizon-hotel.com\n"
             "X-Title: Grand Horizon Concierge\n"
             "\n"
             "{\n"
             "  \"model\": \"qwen/qwen3-max\",\n"
             "  \"messages\": [\n"
             "    {\"role\": \"system\", \"content\": \"You are a hotel assistant...\"},\n"
             "    {\"role\": \"user\", \"content\": \"มีห้องว่างไหม\"}\n"
             "  ],\n"
             "  \"temperature\": 0.3,\n"
             "  \"max_tokens\": 4096,\n"
             "  \"reasoning\": {\"effort\": \"high\"}\n"
             "}"),
    ("p", "ความแตกต่างหลักคือ: Ollama ไม่ต้อง API key, ทำงาน offline, "
          "latency คงที่ (4-9s) ส่วน OpenRouter ต้อง API key, "
          "มี rate limiter (20 req/min) ป้องกัน 429, latency ผันผวน (2-38s)"),
]

# --- Use Case Diagrams (ASCII) ---
USE_CASE_DIAGRAMS = [
    ("p", ""),
    ("h2", "5.2.9 Use Case Diagrams"),
    ("p", "แผนภาพ use case แสดงการใช้งานหลักของระบบ แบ่งตามผู้ใช้ 2 ประเภท: "
          "Guest (แขก) และ Admin (พนักงานโรงแรม)"),
    ("p", ""),
    ("pb", "Use Case 1: Guest — การสนทนากับ Chatbot"),
    ("code", "                    +-----------------------------------+\n"
             "                    |   Hotel AI Virtual Assistant      |\n"
             "                    |                                   |\n"
             "   +-------+        |  (UC1) ถามข้อมูลโรงแรม           |\n"
             "   | Guest | -----> |  (UC2) จองห้องพัก                 |\n"
             "   +-------+        |  (UC3) แก้ไข/ยกเลิกการจอง        |\n"
             "                    |  (UC4) ขอบริการ (ผ้าเช็ดตัว ฯลฯ) |\n"
             "                    |  (UC5) Check-in / Check-out       |\n"
             "                    |  (UC6) สนทนาทั่วไป                |\n"
             "                    +-----------------------------------+"),
    ("p", ""),
    ("pb", "Use Case 2: Admin — การจัดการระบบ"),
    ("code", "                    +-----------------------------------+\n"
             "                    |   Hotel Admin Dashboard           |\n"
             "                    |                                   |\n"
             "   +-------+        |  (UC7)  ดู live chat sessions     |\n"
             "   | Admin | -----> |  (UC8)  Take over / release chat  |\n"
             "   +-------+        |  (UC9)  จัดการสถานะห้อง           |\n"
             "                    |  (UC10) จัดการ booking             |\n"
             "                    |  (UC11) ดู audit logs              |\n"
             "                    |  (UC12) สลับ LLM model            |\n"
             "                    +-----------------------------------+"),
    ("p", ""),
    ("pb", "Use Case 3: ขั้นตอนการจองห้องพัก (UC2 Detail)"),
    ("code", "  Guest                        System                     Database\n"
             "    |                            |                           |\n"
             "    |  \"I want to book a room\"   |                           |\n"
             "    |--------------------------->|                           |\n"
             "    |                            | route → booking agent     |\n"
             "    |  \"What dates?\"             |                           |\n"
             "    |<---------------------------|                           |\n"
             "    |                            |                           |\n"
             "    |  \"Apr 15-17, Deluxe\"       |                           |\n"
             "    |--------------------------->|                           |\n"
             "    |                            | check_room_availability() |\n"
             "    |                            |-------------------------->|\n"
             "    |                            |   [10 Deluxe rooms]       |\n"
             "    |                            |<--------------------------|\n"
             "    |  \"Available at 4,500/night  |                           |\n"
             "    |   Your email?\"             |                           |\n"
             "    |<---------------------------|                           |\n"
             "    |                            |                           |\n"
             "    |  \"my@email.com\"            |                           |\n"
             "    |--------------------------->|                           |\n"
             "    |                            | create_reservation()      |\n"
             "    |                            |-------------------------->|\n"
             "    |                            |   HTL2604150001           |\n"
             "    |                            |<--------------------------|\n"
             "    |  \"Booked! Confirm?\"        |                           |\n"
             "    |<---------------------------|                           |\n"
             "    |                            |                           |\n"
             "    |  \"Yes\"                     |                           |\n"
             "    |--------------------------->|                           |\n"
             "    |                            | confirm_reservation()     |\n"
             "    |                            |-------------------------->|\n"
             "    |  \"Confirmed! HTL2604150001\"|                           |\n"
             "    |<---------------------------|                           |"),
    ("p", "การจองเป็นกรณีที่ LangGraph agent loop ทำงานหลายรอบ — booking sub-agent เรียก tool "
          "ตรวจสอบห้องว่าง, สร้างการจอง, ยืนยันการจอง โดยวน loop กลับมาที่ booking node "
          "ทุกครั้งที่ต้องเรียก tool เพิ่ม"),
    ("p", ""),
    ("pb", "Use Case 4: Knowledge Query — RAG Flow (UC1 Detail)"),
    ("code", "  Guest                 LangGraph          Qdrant         Ollama/OpenRouter\n"
             "    |                      |                  |                  |\n"
             "    | \"WiFi password?\"     |                  |                  |\n"
             "    |--------------------->|                  |                  |\n"
             "    |                      | route→knowledge  |                  |\n"
             "    |                      |                  |                  |\n"
             "    |                      | embed query      |                  |\n"
             "    |                      |----------------->|                  |\n"
             "    |                      | top-3 chunks     |                  |\n"
             "    |                      |<-----------------|                  |\n"
             "    |                      |                  |                  |\n"
             "    |                      | prompt + context |                  |\n"
             "    |                      |---------------------------------->|\n"
             "    |                      | grounded answer  |                  |\n"
             "    |                      |<----------------------------------|\n"
             "    | \"WiFi: GrandHorizon  |                  |                  |\n"
             "    |  _Guest, Welcome2026\"|                  |                  |\n"
             "    |<---------------------|                  |                  |"),
    ("p", "Knowledge query ไม่ใช้ tool call แต่ invoke RAG pipeline โดยตรง "
          "context ถูกวางหลังคำถามผู้ใช้ (ไม่ใช่ก่อน) เพื่อให้โมเดลโฟกัสที่คำถาม"),
    ("p", ""),
    ("pb", "Use Case 5: Authentication Flow (UC7-UC12 prerequisite)"),
    ("code", "  Admin                  FastAPI                  PostgreSQL\n"
             "    |                      |                         |\n"
             "    | POST /auth/login     |                         |\n"
             "    | {user, password}     |                         |\n"
             "    |--------------------->|                         |\n"
             "    |                      | check IP rate limit     |\n"
             "    |                      | check user rate limit   |\n"
             "    |                      | check account lockout   |\n"
             "    |                      |                         |\n"
             "    |                      | SELECT * FROM users     |\n"
             "    |                      |------------------------>|\n"
             "    |                      | bcrypt.verify(password) |\n"
             "    |                      |                         |\n"
             "    |                      | JWT {sub, role=admin,   |\n"
             "    |                      |      iat, exp, jti}     |\n"
             "    |                      |                         |\n"
             "    | {access_token, user} |                         |\n"
             "    |<---------------------|                         |\n"
             "    |                      |                         |\n"
             "    | GET /admin/sessions  |                         |\n"
             "    | Authorization: Bearer|                         |\n"
             "    |--------------------->|                         |\n"
             "    |                      | decode JWT              |\n"
             "    |                      | check jti blocklist     |\n"
             "    |                      | check password_changed  |\n"
             "    |                      | verify role=admin       |\n"
             "    |                      |                         |\n"
             "    | [sessions list]      |                         |\n"
             "    |<---------------------|                         |"),
    ("p", ""),
    ("pb", "Use Case 6: Chat Scaling — Concurrent Request Flow"),
    ("code", "  User A     User B     FastAPI        Scaling         Ollama\n"
             "    |          |          |             Primitives        |\n"
             "    | /chat    |          |                |              |\n"
             "    |--------->|          |                |              |\n"
             "    |          | /chat    |                |              |\n"
             "    |          |--------->|                |              |\n"
             "    |          |          |                |              |\n"
             "    |          |          | rate_limit(A)  |              |\n"
             "    |          |          |--------------->| OK           |\n"
             "    |          |          | rate_limit(B)  |              |\n"
             "    |          |          |--------------->| OK           |\n"
             "    |          |          |                |              |\n"
             "    |          |          | session_lock(A)|              |\n"
             "    |          |          |--------------->| acquired     |\n"
             "    |          |          | session_lock(B)|              |\n"
             "    |          |          |--------------->| acquired     |\n"
             "    |          |          |                |              |\n"
             "    |          |          | llm_semaphore  |              |\n"
             "    |          |          |--------------->| slot 1 → A   |\n"
             "    |          |          |                | slot 2 → B   |\n"
             "    |          |          |                |              |\n"
             "    |          |          |                | LLM(A)------>|\n"
             "    |          |          |                | LLM(B)------>|\n"
             "    |          |          |                |    parallel  |\n"
             "    | resp A   |          |                |<-------------||\n"
             "    |<---------|----------|----------------|              |\n"
             "    |          | resp B   |                |<-------------|\n"
             "    |          |<---------|----------------|              |"),
    ("p", "เมื่อ User A และ B ส่งข้อความพร้อมกัน ทั้งคู่ผ่าน rate limiter และ session lock "
          "(ต่าง session) แล้วเข้า LLM semaphore (2 slots) ทำให้ Ollama ประมวลผลแบบ parallel "
          "ทั้งคู่ได้คำตอบภายใน ~9 วินาที แทนที่จะ serialized เป็น ~18 วินาที"),
]

# --- LangGraph Agent Loop explanation ---
LANGGRAPH_LOOP = [
    ("p", ""),
    ("pb", "LangGraph Agent Loop: ตัวอย่างการทำงานจริง"),
    ("p", "เมื่อแขกส่งข้อความ \"I want to book a Deluxe room for April 15-17\" "
          "LangGraph agent ทำงานดังนี้:"),
    ("code", "Loop Iteration 1:\n"
             "  primary_assistant → LLM → emit ToHotelBooking tool call\n"
             "  route_primary_assistant() → \"enter_booking\"\n"
             "  enter_booking → inject \"Booking Assistant\" system message\n"
             "\n"
             "Loop Iteration 2:\n"
             "  hotel_booking → LLM (with booking_tools) → emit check_room_availability\n"
             "  route_booking() → \"booking_tools\" (tool call detected)\n"
             "  booking_tools → execute check_room_availability() → return rooms list\n"
             "  booking_tools → hotel_booking (cyclic edge — loop back)\n"
             "\n"
             "Loop Iteration 3:\n"
             "  hotel_booking → LLM (sees room list) → emit create_reservation\n"
             "  route_booking() → \"booking_tools\"\n"
             "  booking_tools → execute create_reservation() → return confirmation\n"
             "  booking_tools → hotel_booking (loop back)\n"
             "\n"
             "Loop Iteration 4:\n"
             "  hotel_booking → LLM (sees confirmation) → no more tool calls → respond\n"
             "  route_booking() → END (no tool call detected)\n"
             "  → Return final response to guest"),
    ("p", "จำนวน loop iterations ขึ้นอยู่กับความซับซ้อนของคำขอ — "
          "knowledge query ใช้ 2 iterations (route + RAG) "
          "การจองที่สมบูรณ์ใช้ 4-6 iterations (route + check + create + confirm + upsell)"),
]

# ============================================================================
# 3. CH2 ADDITION — Python 3.10
# ============================================================================
PYTHON_ADDITION = [
    ("p", ""),
    ("h2", "2.5.6 Python"),
    ("p", "Python เป็นภาษาโปรแกรมที่ใช้กันแพร่หลายที่สุดในด้าน AI/ML และ web development "
          "โปรเจกต์นี้ใช้ Python 3.10+ (Van Rossum & Drake, 2023) สำหรับ backend ทั้งหมด "
          "ตั้งแต่ FastAPI server, LangGraph agent, RAG pipeline ไปจนถึง authentication "
          "Python 3.10 มีฟีเจอร์สำคัญ ได้แก่ structural pattern matching (match/case), "
          "improved error messages และ type union syntax (X | Y) ที่ใช้ใน type hints ตลอดโค้ด "
          "ตาม Python Enhancement Proposals (PEP 634-636)"),
]

# ============================================================================
# MAIN
# ============================================================================
def main():
    print("=" * 60)
    print("Patch v5: fonts + CH5 diagrams + CH2 Python + Fig 6.1")
    print("=" * 60)

    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    body = doc.element.body

    # --- 1. Fix fonts ---
    print("\n1. Fixing fonts (removing run-level overrides)...")
    fixed = fix_fonts(doc)
    print(f"  Fixed {fixed} runs")

    # --- 2. Find chapter positions ---
    chapters = find_chapter_indices(doc)
    for ch, idx in sorted(chapters.items(), key=lambda x: str(x[0])):
        print(f"  CH{ch}: P{idx}")

    ch5_start = chapters.get(5, 0)
    ch6_start = chapters.get(6, len(doc.paragraphs))
    ch2_start = chapters.get(2, 0)
    ch3_start = chapters.get(3, len(doc.paragraphs))

    # --- 3. Add tool calling sample after "ระบบมี hotel-specific tools จำนวน 15" ---
    print("\n2. Adding tool calling sample to CH5...")
    anchor = find_paragraph_containing(doc, "hotel-specific tools จำนวน 15", ch5_start, ch6_start)
    if anchor is not None:
        # Find the table after it (skip 1 paragraph)
        # Insert after the table placeholder/table
        insert_point = anchor + 1
        # Look for next heading or substantial paragraph
        for j in range(anchor + 1, min(anchor + 5, ch6_start)):
            text = doc.paragraphs[j].text.strip()
            if text.startswith("5.2.4") or text.startswith("Phase 4") or "Docker" in text:
                insert_point = j
                break
        ref_elem = doc.paragraphs[insert_point]._element
        nodes = build_nodes(doc, TOOL_CALLING_SAMPLE)
        parent = ref_elem.getparent()
        idx = list(parent).index(ref_elem)
        for j, n in enumerate(nodes):
            parent.insert(idx + j, n)
        print(f"  Inserted {len(nodes)} paragraphs at P{insert_point}")
    else:
        print("  [WARN] Tool anchor not found")

    # Re-scan after insertion
    chapters = find_chapter_indices(doc)
    ch5_start = chapters.get(5, 0)
    ch6_start = chapters.get(6, len(doc.paragraphs))

    # --- 4. Add Ollama/OpenRouter API sample after Phase 4 Docker section ---
    print("\n3. Adding Ollama/OpenRouter API samples...")
    anchor = find_paragraph_containing(doc, "RuntimeLLMConfig", ch5_start, ch6_start)
    if anchor is not None:
        # Find end of RuntimeLLMConfig code block
        insert_after_idx = anchor
        for j in range(anchor + 1, min(anchor + 20, ch6_start)):
            text = doc.paragraphs[j].text.strip()
            if text.startswith("5.2.5") or "Phase 5" in text or "Authentication" in text:
                insert_after_idx = j
                break
        ref_elem = doc.paragraphs[insert_after_idx]._element
        nodes = build_nodes(doc, OLLAMA_OPENROUTER_SAMPLE)
        parent = ref_elem.getparent()
        idx = list(parent).index(ref_elem)
        for j, n in enumerate(nodes):
            parent.insert(idx + j, n)
        print(f"  Inserted {len(nodes)} paragraphs at P{insert_after_idx}")
    else:
        print("  [WARN] RuntimeLLMConfig anchor not found")

    # Re-scan
    chapters = find_chapter_indices(doc)
    ch5_start = chapters.get(5, 0)
    ch6_start = chapters.get(6, len(doc.paragraphs))

    # --- 5. Add Use Case Diagrams + LangGraph Loop at end of CH5 section 5.2 ---
    print("\n4. Adding use case diagrams and LangGraph loop...")
    # Find last Phase 8 paragraph or end of section 5.2
    anchor = find_paragraph_containing(doc, "native thinking", ch5_start, ch6_start)
    if anchor is None:
        anchor = find_paragraph_containing(doc, "Phase 8", ch5_start, ch6_start)
    if anchor is not None:
        # Find end of Phase 8 section
        insert_idx = anchor
        for j in range(anchor + 1, min(anchor + 20, ch6_start)):
            text = doc.paragraphs[j].text.strip()
            if text.startswith("5.3") or "Prompt Engineering" in text:
                insert_idx = j
                break
        ref_elem = doc.paragraphs[insert_idx]._element
        all_nodes = build_nodes(doc, USE_CASE_DIAGRAMS + LANGGRAPH_LOOP)
        parent = ref_elem.getparent()
        idx = list(parent).index(ref_elem)
        for j, n in enumerate(all_nodes):
            parent.insert(idx + j, n)
        print(f"  Inserted {len(all_nodes)} paragraphs at P{insert_idx}")
    else:
        print("  [WARN] Phase 8 anchor not found")

    # --- 6. Add Python to CH2 (after Ant Design section) ---
    print("\n5. Adding Python 3.10 to CH2...")
    chapters = find_chapter_indices(doc)
    ch2_start = chapters.get(2, 0)
    ch3_start = chapters.get(3, len(doc.paragraphs))
    anchor = find_paragraph_containing(doc, "Ant Design", ch2_start, ch3_start)
    if anchor is not None:
        # Find end of Ant Design paragraph
        insert_idx = anchor + 1
        ref_elem = doc.paragraphs[insert_idx]._element
        nodes = build_nodes(doc, PYTHON_ADDITION)
        parent = ref_elem.getparent()
        idx = list(parent).index(ref_elem)
        for j, n in enumerate(nodes):
            parent.insert(idx + j, n)
        print(f"  Inserted {len(nodes)} paragraphs at P{insert_idx}")
    else:
        print("  [WARN] Ant Design anchor not found in CH2")

    # --- 7. Update Figure 6.1 caption ---
    print("\n6. Updating Figure 6.1 caption...")
    chapters = find_chapter_indices(doc)
    ch6_start = chapters.get(6, 0)
    ch7_start = chapters.get(7, len(doc.paragraphs))
    for i in range(ch6_start, ch7_start):
        text = doc.paragraphs[i].text.strip()
        if "Figure 6.1" in text and ("bar chart" in text.lower() or "accuracy" in text.lower()):
            for run in doc.paragraphs[i].runs:
                if "Figure 6.1" in run.text:
                    run.text = ("[Figure 6.1: ความแม่นยำรายหมวดหมู่ — Local 9B vs Cloud "
                                "ทั้งสองโมเดลแม่นยำ 100% ในหมวด Knowledge, Booking และ Language "
                                "ความแตกต่างอยู่ที่ Greeting (75% vs 100%) และ Edge Cases (75% vs 100%) "
                                "รวม: Local 92% vs Cloud 100%]")
                    print(f"  Updated at P{i}")
                    break
            break

    # --- 8. Re-attach Figure 6.1 PNG (new version) ---
    print("\n7. Re-attaching updated Figure 6.1...")
    fig_path = FIGURES_DIR / "Fig_6.1_Accuracy_Comparison.png"
    if fig_path.exists():
        for i in range(ch6_start, ch7_start):
            text = doc.paragraphs[i].text.strip()
            if "Figure 6.1" in text:
                try:
                    insert_image_paragraph(doc, body, doc.paragraphs[i]._element,
                                           fig_path, width_inches=5.5)
                    print(f"  Attached at P{i}")
                except Exception as e:
                    print(f"  [FAIL] {e}")
                break

    # --- 9. Save ---
    doc.save(str(OUT))
    sz = os.path.getsize(str(OUT)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUT} ({sz:.0f} KB)")
    print("=" * 60)

if __name__ == "__main__":
    main()
