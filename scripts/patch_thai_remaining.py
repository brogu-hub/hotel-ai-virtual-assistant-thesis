#!/usr/bin/env python3
"""
patch_thai_remaining.py
=======================
Translate remaining English sections in thesis_final_v3.docx:
  - CH1-CH4: full replacement with Thai
  - References: add new section between CH8 and Appendix
  - AP_A, AP_B: in-place translation of headings and prose

Input:  thesis/thesis_final_v3.docx
Output: thesis/thesis_final_v4.docx
"""
import os, sys, re, shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final_v3.docx")
OUT = Path("thesis/thesis_final_v4.docx")
FIGURES_DIR = Path("thesis/figures")

# === Helpers (same as patch_thai_translation.py) ===

def make_para(doc, text, style_name, bold=False, italic=False,
              font_name=None, font_size=None, alignment=None):
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
            rPr.append(OxmlElement('w:bCs'))
        if italic:
            rPr.append(OxmlElement('w:i'))
            rPr.append(OxmlElement('w:iCs'))
        if font_name:
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), font_name)
            rf.set(qn('w:hAnsi'), font_name)
            rf.set(qn('w:cs'), font_name)
            rPr.append(rf)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size * 2)))
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(sz)
            rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p

def make_code_para(doc, text):
    return make_para(doc, text, "TU_Paragraph_Normal", font_name="Consolas", font_size=10)

def build_nodes(doc, specs):
    nodes = []
    for spec in specs:
        kind, text = spec[0], spec[1] if len(spec) > 1 else ""
        if kind == "h1":
            nodes.append(make_para(doc, text, "TU_Sub-heading 1", bold=True))
        elif kind == "h2":
            nodes.append(make_para(doc, text, "TU_Sub-heading 2", bold=True))
        elif kind == "h3":
            nodes.append(make_para(doc, text, "TU_Sub-heading 3", bold=True))
        elif kind == "p":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal"))
        elif kind == "pb":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", bold=True))
        elif kind == "pi":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", italic=True))
        elif kind == "code":
            for line in text.split("\n"):
                nodes.append(make_code_para(doc, line if line else " "))
        elif kind == "figure":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", italic=True, alignment="center"))
    return nodes

def insert_after_element(body, ref_elem, nodes):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for j, node in enumerate(nodes):
        parent.insert(idx + 1 + j, node)

def remove_between_elements(body, start_elem, end_elem):
    children = list(body)
    start_pos = children.index(start_elem)
    end_pos = children.index(end_elem) if end_elem is not None else len(children)
    to_remove = children[start_pos + 1:end_pos]
    for elem in to_remove:
        body.remove(elem)
    return len(to_remove)

def create_word_table(doc, headers, rows, ref_element):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols, style='Table Grid')
    hdr = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                table.rows[i + 1].cells[j].text = str(cell_text)
    tbl = table._tbl
    tbl_parent = tbl.getparent()
    tbl_parent.remove(tbl)
    ref_parent = ref_element.getparent()
    idx = list(ref_parent).index(ref_element)
    ref_parent.insert(idx + 1, tbl)
    return tbl

def find_chapter_indices(doc):
    chapters = {}
    HEADING_MAP = {
        "introduction": 1, "literature review": 2, "methodology": 3,
        "system design": 4, "implementation": 5, "testing and evaluation": 6,
        "discussion": 7, "conclusion": 8,
    }
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            lower = text.lower()
            for ht, cn in HEADING_MAP.items():
                if ht in lower and cn not in chapters:
                    chapters[cn] = i
                    break
        if style == "TU_Chapter":
            if "บทคัดย่อ" in text:
                chapters["abstract_th"] = i
            elif text.strip() == "ABSTRACT":
                chapters["abstract_en"] = i
            elif "ภาคผนวก ก" in text:
                chapters["ap_a"] = i
            elif "ภาคผนวก ข" in text:
                chapters["ap_b"] = i
            elif "ภาคผนวก ค" in text:
                chapters["ap_c"] = i
            elif text.strip() == "ภาคผนวก":
                chapters["appendix"] = i
            elif "ประวัติผู้เขียน" in text:
                chapters["bio"] = i
    return chapters

# ============================================================================
# CH1: Introduction (บทนำ)
# ============================================================================
def get_ch1_specs():
    return [
        ("h1", "1.1 ภูมิหลังและแรงจูงใจ"),
        ("p", "อุตสาหกรรมการบริการทั่วโลกให้บริการนักท่องเที่ยวระหว่างประเทศมากกว่า 1.4 พันล้านคนต่อปี "
              "โดยประเทศไทยจัดอยู่ในอันดับสิบประเทศที่มีนักท่องเที่ยวมากที่สุดในโลก "
              "ขณะที่โรงแรมแข่งขันกันด้านความพึงพอใจของแขก คุณภาพและความรวดเร็วของการบริการลูกค้า "
              "จึงเป็นปัจจัยสร้างความแตกต่างที่สำคัญ การดำเนินงานแผนกต้อนรับแบบดั้งเดิม "
              "พึ่งพาพนักงานที่ทำงานเป็นกะ ทำให้เกิดปัญหาคอขวดในช่วงเวลาเร่งด่วน "
              "และแขกไม่ได้รับความช่วยเหลือในช่วงดึก"),
        ("p", "ปัญญาประดิษฐ์เริ่มเปลี่ยนแปลงการดำเนินงานของธุรกิจโรงแรม "
              "Buhalis and Moldavska (2022) แสดงให้เห็นว่า AI assistants ทั้งเสียงและข้อความ "
              "สามารถจัดการคำถามทั่วไปของแขก ตั้งแต่เวลาอาหารเช้าไปจนถึงรหัส WiFi "
              "ด้วยเวลาตอบสนองเป็นวินาที Buhalis, O'Connor, and Leung (2023) เสนอต่อไปว่า "
              "ระบบนิเวศ \"smart hospitality\" ที่ AI agents เชื่อมต่อกับระบบ property management systems (PMS) "
              "เป็นวิวัฒนาการขั้นต่อไปของเทคโนโลยีโรงแรม"),
        ("p", "อย่างไรก็ตาม chatbot โรงแรมส่วนใหญ่ที่ใช้งานอยู่ยังคงเป็นระบบ rule-based หรือ "
              "FAQ-retrieval ที่มี decision trees ตายตัว ไม่สามารถรับมือกับความหลากหลายของภาษาธรรมชาติ "
              "ขั้นตอนการจองแบบหลายขั้น หรือการสนทนาสองภาษา แขกที่ถามว่า "
              "\"มีห้องว่างวันจันทร์หน้าไหม\" จะไม่ได้คำตอบที่เป็นประโยชน์จากระบบ keyword-matching"),
        ("p", "ความก้าวหน้าล่าสุดของ Large Language Models (LLMs) และ agentic AI frameworks "
              "ทำให้สามารถสร้าง chatbot ที่เข้าใจภาษาธรรมชาติ วิเคราะห์งานหลายขั้นตอน "
              "ค้นหาข้อมูลจากฐานความรู้ และเชื่อมต่อกับฐานข้อมูล ทั้งหมดนี้โดยรักษา context "
              "ตลอดการสนทนา วิทยานิพนธ์นี้สำรวจโอกาสนี้โดยการสร้าง AI virtual assistant "
              "ระดับ production สำหรับโรงแรมหรู"),
        ("h1", "1.2 คำชี้แจงปัญหา"),
        ("p", "ปัญหาหลัก 3 ประการของ chatbot โรงแรมปัจจุบัน:"),
        ("p", "1. การจำแนก intent ที่ตายตัว — ระบบ rule-based ล้มเหลวเมื่อแขกใช้ภาษาต่างกัน "
              "คำว่า \"I'd like to cancel\" และ \"ยกเลิกการจอง\" ต้อง route ไปยัง handler เดียวกัน"),
        ("p", "2. ไม่มีการเชื่อมต่อฐานข้อมูล — chatbot ส่วนใหญ่ตอบจาก FAQ คงที่ "
              "ไม่สามารถตรวจสอบห้องว่าง สร้างการจอง หรือแก้ไข booking แบบ real-time"),
        ("p", "3. ไม่มีการพิจารณา scalability — chatbot แบบ single-threaded ไม่สามารถให้บริการ "
              "แขกหลายคนพร้อมกัน และขาดมาตรการ authentication, audit logging สำหรับพนักงาน"),
        ("h1", "1.3 วัตถุประสงค์"),
        ("p", "วิทยานิพนธ์นี้มีเป้าหมายในการออกแบบ พัฒนา และประเมินผล AI virtual assistant สำหรับโรงแรม:"),
        ("p", "1. ออกแบบระบบ multi-agent AI โดยใช้ LangGraph state machines เพื่อ route คำขอแขก "
              "ไปยัง sub-agents เฉพาะทาง (booking, service, knowledge, general conversation)"),
        ("p", "2. พัฒนา Retrieval-Augmented Generation (RAG) สำหรับฐานความรู้โรงแรม "
              "โดยใช้ vector embeddings และ Qdrant เพื่อตอบคำถามเกี่ยวกับสิ่งอำนวยความสะดวก นโยบาย และบริการ"),
        ("p", "3. สร้างระบบ fullstack ระดับ production ด้วย FastAPI backend, Next.js frontend, "
              "JWT authentication, admin dashboard, audit logging และ Docker deployment"),
        ("p", "4. เปรียบเทียบประสิทธิภาพ local vs cloud LLM โดยประเมิน local 9B model "
              "(Qwen3.5 Opus 9B บน Ollama) กับ cloud model (Qwen3 Max บน OpenRouter) "
              "ด้วย 25 test cases ในขอบเขตโรงแรม"),
        ("h1", "1.4 ขอบเขตและข้อจำกัด"),
        ("pb", "ในขอบเขต:"),
        ("p", "• การสนทนาสองภาษาทั้งไทยและอังกฤษ"),
        ("p", "• Sub-agents เฉพาะทาง 4 ตัว: booking, services, knowledge retrieval และ general conversation"),
        ("p", "• ตรวจสอบห้องว่าง, reservation CRUD, check-in/check-out workflows"),
        ("p", "• RAG สำหรับเอกสารความรู้โรงแรม 10 ฉบับ (dining, spa, policies, facilities, FAQ)"),
        ("p", "• JWT authentication แบ่ง user และ admin roles"),
        ("p", "• Admin dashboard สำหรับ session monitoring, chat intervention และ audit logging"),
        ("p", "• Concurrent-user scaling ด้วย LLM semaphore, session locks และ knowledge caching"),
        ("p", "• Runtime switching ระหว่าง local Ollama กับ cloud OpenRouter"),
        ("p", "• การประเมินอัตโนมัติด้วย 25 test cases และ 193 infrastructure tests"),
        ("pb", "นอกขอบเขต:"),
        ("p", "• การเชื่อมต่อระบบชำระเงินจริง (มีเฉพาะ mock payment links)"),
        ("p", "• Voice interaction (เฉพาะข้อความ)"),
        ("p", "• Multi-property support (โรงแรมเดียว)"),
        ("p", "• Production deployment ภายใต้ traffic จริง"),
        ("h1", "1.5 ผลงานที่คาดหวัง"),
        ("p", "1. Reference architecture ที่ทำซ้ำได้สำหรับการสร้าง hotel AI assistants ด้วย LangGraph และ RAG"),
        ("p", "2. การเปรียบเทียบเชิงประจักษ์ของ local 9B vs cloud LLM สำหรับงานโรงแรม (accuracy, latency, cost)"),
        ("p", "3. ชุด scaling primitives (LLM concurrency limiter, session lock manager, knowledge cache) "
              "ที่ใช้ได้กับทุก LLM-backed service"),
        ("p", "4. ชุดทดสอบสมบูรณ์ (193 infrastructure tests + 25 model evaluation cases) "
              "เป็น benchmark สำหรับงานวิจัย hotel chatbot ในอนาคต"),
        ("h1", "1.6 โครงสร้างวิทยานิพนธ์"),
        ("p", "TABLE_CH1_ORG_PLACEHOLDER"),
        ("figure", "[Figure 1.1: System context diagram — แขกใช้ chatbot ผ่านเว็บไซต์โรงแรม "
                   "chatbot เชื่อมต่อกับ LangGraph agent, PostgreSQL, Qdrant vector store "
                   "และ Ollama/OpenRouter LLM backends พนักงานเข้าถึง admin dashboard]"),
    ]

# ============================================================================
# CH2: Literature Review (การทบทวนวรรณกรรม)
# ============================================================================
def get_ch2_specs():
    return [
        ("h1", "2.1 ปัญญาประดิษฐ์ในอุตสาหกรรมการบริการ"),
        ("h2", "2.1.1 ระบบบริหารจัดการโรงแรมและการเปลี่ยนผ่านดิจิทัล"),
        ("p", "Hotel Property Management Systems (PMS) เป็นแกนหลักของการดำเนินงานโรงแรมสมัยใหม่ "
              "จัดการการจอง สถานะห้อง ข้อมูลแขก การเรียกเก็บเงิน และแม่บ้าน Oracle OPERA ซึ่งเป็น PMS "
              "ชั้นนำในตลาด กำหนด room status lifecycle เป็น: Vacant Clean → Occupied → Vacant Dirty "
              "→ Cleaning → Inspected → Vacant Clean (Oracle, 2023)"),
        ("p", "งานวิจัยทางวิชาการเกี่ยวกับระบบบริหารจัดการโรงแรม (Atlantis Press, ICETIS '13; JETIR, 2023) "
              "ได้กำหนดข้อกำหนดด้านฟังก์ชันหลัก: reservation management, guest registration, "
              "check-in/check-out และ billing ระบบ PMS สมัยใหม่กำลังเปลี่ยนจาก on-premise monoliths "
              "เป็น cloud-based microservices (AltexSoft, 2023; DDI Development, 2023) "
              "ซึ่งสอดคล้องกับสถาปัตยกรรม microservice ที่ใช้ในวิทยานิพนธ์นี้"),
        ("h2", "2.1.2 AI Chatbots ในโรงแรม: จาก Rule-Based สู่ Generative"),
        ("p", "วิวัฒนาการของ chatbot ในอุตสาหกรรมการบริการแบ่งเป็น 4 ยุค:"),
        ("figure", "[Figure 2.1: Chatbot generation taxonomy — (1) Rule-based: keyword matching; "
                   "(2) Retrieval-based: FAQ lookup; (3) Generative: LLM-powered; "
                   "(4) Agentic: multi-agent with tool calling โปรเจกต์นี้อยู่ใน Generation 4]"),
        ("p", "การวิเคราะห์ bibliometric ของบทความ chatbot 398 ฉบับ (2003-2023) (Taylor & Francis, 2025) "
              "เผยว่างานวิจัยเปลี่ยนทิศทางอย่างมากไปสู่ generative AI หลังปี 2022 "
              "การทบทวนบทความ 71 ฉบับเกี่ยวกับ conversational AI ในอุตสาหกรรมการบริการ "
              "(Taylor & Francis, 2026) ยืนยันว่าระบบที่ใช้ LLM เป็นทิศทางวิจัยหลักในปัจจุบัน"),
        ("p", "Buhalis and Moldavska (2022) ทำการศึกษาเชิงสำรวจเกี่ยวกับ AI assistants "
              "ในอุตสาหกรรมการบริการ แสดงให้เห็นว่า voice และ text assistants สามารถจัดการคำถามทั่วไป "
              "ด้วยคะแนนความพึงพอใจสูง พวกเขาระบุข้อกำหนดสำคัญ: natural language understanding, "
              "context retention across turns และ integration กับระบบปฏิบัติการโรงแรม "
              "Buhalis, O'Connor, and Leung (2023) ขยายแนวคิดนี้เป็น \"smart hospitality ecosystems\" "
              "ที่ AI agents โต้ตอบกับ PMS, housekeeping และ revenue management systems แบบ real-time "
              "ซึ่งเป็นสถาปัตยกรรมที่วิทยานิพนธ์นี้พัฒนา"),
        ("h1", "2.2 Large Language Models"),
        ("h2", "2.2.1 Transformer Architecture และ Instruction Tuning"),
        ("p", "Transformer architecture (Vaswani et al., 2017) นำเสนอ self-attention mechanism "
              "ที่ช่วยให้โมเดลประมวลผล sequences แบบ parallel แทน sequential "
              "LLMs สมัยใหม่สร้างบนพื้นฐานนี้ด้วย instruction tuning ซึ่งเป็นการ fine-tune "
              "บน instruction-response pairs ที่มนุษย์สร้างขึ้น ตระกูลโมเดล Qwen ที่ใช้ในโปรเจกต์นี้ "
              "ใช้ dense Transformer architectures พร้อม Rotary Position Embeddings (RoPE)"),
        ("h2", "2.2.2 Small vs. Large Models"),
        ("p", "คำถามสำคัญในการ deploy ระบบที่ใช้ LLM คือ trade-off ระหว่างขนาดโมเดลกับความสามารถ "
              "วิทยานิพนธ์นี้เปรียบเทียบ:"),
        ("p", "• Qwen3.5 Opus 9B — 9 พันล้าน parameters, รัน local บน GPU เดียว (RTX 5080, 16 GB VRAM)"),
        ("p", "• Qwen3 Max — cloud-hosted flagship model ผ่าน OpenRouter API"),
        ("p", "9B model เป็นเพดานที่เป็นไปได้สำหรับ single-GPU local deployment ในโรงแรม "
              "(ไม่พึ่งพา cloud, ไม่มีค่าใช้จ่ายต่อ query, ข้อมูลอยู่ on-premise) "
              "Cloud model เป็นความสามารถที่ไม่จำกัดแลกกับ latency, cost และ data privacy"),
        ("h2", "2.2.3 Thinking Models และ Reasoning"),
        ("p", "LLMs สมัยใหม่รวม explicit \"thinking\" mechanisms ที่โมเดลสร้าง internal reasoning tokens "
              "(ใน <think> tags) ก่อนสร้างคำตอบ Qwen3.5 models สร้าง tokens เหล่านี้โดยธรรมชาติ "
              "วิทยานิพนธ์นี้ประเมินผลกระทบของ reasoning capability ต่อ hotel task accuracy "
              "และพบว่าการปิด explicit thinking mode บน 9B local model ลด latency "
              "โดยไม่สูญเสียความแม่นยำที่วัดได้"),
        ("h1", "2.3 Agentic AI และ Orchestration"),
        ("h2", "2.3.1 LangGraph State Machines"),
        ("p", "LangGraph เป็นส่วนขยายของ LangChain ecosystem ที่ให้ framework สำหรับสร้าง "
              "stateful, multi-step AI agents เป็น directed graphs แต่ละ node แทนการกระทำของ agent "
              "(LLM call, tool execution, routing decision) และ edges กำหนดการไหลระหว่างการกระทำ "
              "พร้อม conditional branching"),
        ("p", "Liu et al. (2024) นำเสนอแหล่งอ้างอิงทางวิชาการ (arXiv:2411.18241) เกี่ยวกับ LangGraph "
              "สำหรับ multi-agent systems สำรวจว่า LangGraph ร่วมกับ frameworks อื่นสามารถ implement "
              "complex agent orchestration patterns ได้ ผลงานนี้ยืนยันสถาปัตยกรรม "
              "\"hybrid router → LangGraph → sub-agents\" ที่ใช้ในโปรเจกต์นี้"),
        ("figure", "[Figure 2.2: LangGraph concept — directed graph ที่ primary assistant node "
                   "route ไปยัง sub-agent nodes ตาม tool-call dispatch]"),
        ("h2", "2.3.2 Multi-Agent Routing Patterns"),
        ("p", "ReAct (Reasoning + Acting) pattern ที่ LLM สลับระหว่าง reasoning กับ actions (tool calls) "
              "เป็น foundational pattern สำหรับ agentic AI โปรเจกต์นี้ขยาย ReAct ด้วย "
              "sub-agent routing: primary assistant ไม่ execute tools โดยตรง "
              "แต่ dispatch ไปยัง sub-agents เฉพาะทางที่มี tool sets และ prompts ของตัวเอง"),
        ("h2", "2.3.3 Tool Calling และ Function Invocation"),
        ("p", "LLMs สมัยใหม่รองรับ structured tool calling ที่โมเดล output JSON-formatted function call "
              "LangGraph ToolNode ดักจับ calls เหล่านี้ execute Python function ที่เกี่ยวข้อง "
              "และส่งผลลัพธ์กลับ วิทยานิพนธ์นี้ implement hotel-specific tools 15 ตัว"),
        ("h1", "2.4 Retrieval-Augmented Generation"),
        ("h2", "2.4.1 RAG Pipeline"),
        ("p", "Retrieval-Augmented Generation (RAG) แก้ปัญหา knowledge cutoff และ hallucination "
              "ของ LLM โดยค้นหาเอกสารที่เกี่ยวข้องจากฐานความรู้ก่อนสร้างคำตอบ ขั้นตอน:"),
        ("p", "1. Chunking — แบ่งเอกสารเป็น segments ที่ overlap"),
        ("p", "2. Embedding — แปลง chunks เป็น dense vector representations"),
        ("p", "3. Indexing — เก็บ vectors ในฐานข้อมูล vector"),
        ("p", "4. Retrieval — ค้นหา chunks ที่ใกล้เคียงกับคำถามผู้ใช้มากที่สุด"),
        ("p", "5. Generation — ส่ง retrieved context ให้ LLM พร้อมคำถาม"),
        ("figure", "[Figure 2.3: RAG pipeline — hotel knowledge documents → chunked → embedded "
                   "via qwen3-embedding-8b (4096 dims) → Qdrant → query time: embed user message "
                   "→ top-k similar chunks → LLM generates grounded response]"),
        ("h2", "2.4.2 Vector Databases"),
        ("p", "Vector databases ถูกสร้างเฉพาะสำหรับ similarity search บน embeddings มิติสูง "
              "โปรเจกต์นี้ใช้ Qdrant ซึ่งเป็น open-source vector database ที่เขียนด้วย Rust "
              "เลือกจากประสิทธิภาพ cosine similarity search และ Docker-native deployment"),
        ("h2", "2.4.3 Reranking และการตัดสินใจลบออก"),
        ("p", "Cross-encoder rerankers สามารถปรับปรุงความแม่นยำของ retrieval "
              "แต่โปรเจกต์นี้พบว่า reranker เพิ่ม ~1-2 วินาทีของ CPU-bound latency ต่อ query "
              "และ block FastAPI async event loop ทำให้ concurrent requests ทั้งหมดหยุดชะงัก "
              "เนื่องจาก embedding-based search จาก Qdrant มีความแม่นยำ 100% อยู่แล้ว (8/8) "
              "reranker จึงถูกปิดโดย default"),
        ("h1", "2.5 เทคโนโลยี Web Application"),
        ("h2", "2.5.1 FastAPI และ Asynchronous Python"),
        ("p", "FastAPI เป็น Python web framework สมัยใหม่ที่สร้างบน Starlette และ Pydantic "
              "ให้ automatic OpenAPI documentation, request validation และ async support "
              "ความสามารถ async มีความสำคัญ: LLM inference ใช้เวลา 3-50 วินาที "
              "server ต้องจัดการ requests อื่น (health checks, admin queries, concurrent chats) โดยไม่ blocking"),
        ("h2", "2.5.2 Next.js App Router และ React Server Components"),
        ("p", "Frontend ใช้ Next.js 15 กับ App Router ที่ใช้ React Server Components (RSC) "
              "สำหรับ server-side rendering ด้วย client JavaScript bundles ที่น้อยที่สุด "
              "(Next.js, 2024; Osmani, 2024) Chat interface ใช้ Server-Sent Events (SSE) "
              "สำหรับ real-time token streaming"),
        ("h2", "2.5.3 State Management: Zustand และ SWR"),
        ("p", "Zustand ถูกเลือกแทน Redux สำหรับ global state management "
              "จากการวิเคราะห์เปรียบเทียบโดย Salah (2024) ที่แสดงว่า Zustand มี bundle size เล็กกว่า "
              "API ง่ายกว่า (ไม่ต้อง reducers/actions boilerplate) SWR implement caching strategy "
              "จาก RFC 5861 (Nottingham, 2010): แสดงข้อมูล cached ทันทีขณะ revalidate ใน background"),
        ("h2", "2.5.4 TypeScript"),
        ("p", "Gao, Bird, and Barr (2017) พบว่า TypeScript สามารถป้องกัน 15% ของ public bugs "
              "ใน JavaScript projects บน GitHub Fischer and Hanenberg (2015) แสดงผลเชิงบวก "
              "ของ static typing ต่อ API usability โปรเจกต์นี้ใช้ TypeScript ตลอด frontend"),
        ("h2", "2.5.5 Ant Design"),
        ("p", "UI component library คือ Ant Design 5 ซึ่งเป็น enterprise-grade React component library "
              "ที่ให้ 60+ components (Table, Form, Modal, Descriptions ฯลฯ) "
              "พร้อม design language ที่สม่ำเสมอ (Ant Design, 2024)"),
        ("h1", "2.6 ความปลอดภัยและการปฏิบัติตามกฎหมาย"),
        ("h2", "2.6.1 JWT Authentication และ RBAC"),
        ("p", "JSON Web Tokens (JWT) ให้ stateless authentication โดย encode user identity "
              "และ role claims ใน cryptographically signed token โปรเจกต์นี้ implement "
              "role-based access control (RBAC) ด้วย 2 roles: user (registered guests) "
              "และ admin (hotel staff)"),
        ("h2", "2.6.2 PII Redaction"),
        ("p", "การสนทนาของแขกอาจมีข้อมูลส่วนบุคคล (PII) เช่น หมายเลขบัตรเครดิต "
              "หมายเลขหนังสือเดินทาง และเบอร์โทรศัพท์ ก่อนส่งข้อความให้ LLM "
              "ระบบ PII redactor ใช้ regex scrub ข้อมูลอ่อนไหว แทนที่ด้วย category tokens "
              "เช่น [CREDIT_CARD]"),
        ("h1", "2.7 วิธีการประเมินโมเดล"),
        ("h2", "2.7.1 RAG Evaluation Metrics"),
        ("p", "DeepEval ให้ automated evaluation metrics สำหรับระบบ RAG ได้แก่:"),
        ("p", "• Faithfulness — คำตอบมีเฉพาะข้อมูลจาก retrieved context หรือไม่"),
        ("p", "• Answer Relevancy — คำตอบเกี่ยวข้องกับคำถามผู้ใช้หรือไม่"),
        ("p", "• Context Recall — retrieved context มีข้อมูลที่ต้องการสำหรับตอบหรือไม่"),
        ("h2", "2.7.2 Cohen's Kappa สำหรับ Inter-Model Agreement"),
        ("p", "Cohen's Kappa (κ) วัดความเห็นพ้องระหว่าง 2 raters (ในกรณีนี้คือ 2 โมเดล) "
              "เกินกว่าที่คาดหวังโดยบังเอิญ สูตร: κ = (p_o - p_e) / (1 - p_e) "
              "โดย p_o คืออัตรา observed agreement และ p_e คือ expected agreement โดยบังเอิญ"),
    ]

# ============================================================================
# CH3: Methodology (ระเบียบวิธีวิจัย)
# ============================================================================
def get_ch3_specs():
    return [
        ("h1", "3.1 แนวทางการวิจัย"),
        ("p", "วิทยานิพนธ์นี้ใช้ระเบียบวิธี Design Science Research (DSR) ซึ่งเหมาะสำหรับ "
              "งานวิจัยระบบสารสนเทศที่มีเป้าหมายสร้างและประเมิน artifact — ในกรณีนี้คือ "
              "hotel AI virtual assistant กรอบ DSR ประกอบด้วย 3 วงจร: "
              "relevance (ระบุปัญหาจริง), design (สร้าง artifact) และ rigor (อ้างอิงองค์ความรู้เดิมและประเมิน)"),
        ("p", "กระบวนการพัฒนาผสมผสาน:"),
        ("p", "1. Literature-driven design — การเลือกเทคโนโลยีและ architecture patterns "
              "อ้างอิงจากงานวิจัยที่ตีพิมพ์ (บทที่ 2)"),
        ("p", "2. Iterative prototyping — ระบบถูกสร้างแบบ incremental ใน 16 สัปดาห์ "
              "มีการทดสอบต่อเนื่องในแต่ละขั้นตอน"),
        ("p", "3. Empirical evaluation — ระบบสุดท้ายถูกประเมินเทียบกับ golden dataset "
              "ด้วย quantitative metrics (บทที่ 6)"),
        ("h1", "3.2 กระบวนการพัฒนา"),
        ("h2", "3.2.1 Agile Iteration Cycle"),
        ("p", "การพัฒนาใช้กระบวนการ agile ที่ปรับแต่งด้วย sprints 2 สัปดาห์:"),
        ("p", "TABLE_CH3_SPRINT_PLACEHOLDER"),
        ("h2", "3.2.2 Version Control และ Documentation"),
        ("p", "การเปลี่ยนแปลงโค้ดทั้งหมดถูกติดตามใน Git ด้วย conventional commit messages "
              "(feat:, fix:, perf:, docs:) การตัดสินใจด้านสถาปัตยกรรมถูกบันทึกใน "
              "docs/WORKFLOW.md ตามที่เกิดขึ้น ไม่ใช่ย้อนหลัง"),
        ("h1", "3.3 การเลือกเทคโนโลยี"),
        ("h2", "3.3.1 เกณฑ์การเลือก"),
        ("p", "TABLE_CH3_CRITERIA_PLACEHOLDER"),
        ("h2", "3.3.2 การเลือก LLM Framework"),
        ("figure", "[Figure 3.1: LLM orchestration framework comparison matrix]"),
        ("p", "TABLE_CH3_FRAMEWORK_PLACEHOLDER"),
        ("pb", "การตัดสินใจ:"),
        ("p", "LangGraph ถูกเลือกจาก checkpointed state persistence (conversation memory รอดจาก server restarts), "
              "native cyclic tool loops (booking sub-agent เรียก tools ซ้ำจนเสร็จ) "
              "และ time-travel debugging (admin ย้อนกลับ conversations ได้)"),
        ("h2", "3.3.3 การเลือก LLM Model"),
        ("p", "TABLE_CH3_MODEL_PLACEHOLDER"),
        ("pb", "การตัดสินใจ:"),
        ("p", "ทั้งสองโมเดลถูก deploy โดย local 9B model เป็น primary backend "
              "(ต้นทุนศูนย์, ข้อมูลอยู่ on-premise) และ cloud Qwen3 Max เป็น runtime-switchable fallback"),
        ("h2", "3.3.4 การเลือก Vector Database"),
        ("p", "TABLE_CH3_VECTOR_PLACEHOLDER"),
        ("pb", "การตัดสินใจ:"),
        ("p", "Qdrant ถูกเลือกจาก Docker-native single-container deployment, Rust-based performance "
              "และ REST API ที่ตรงไปตรงมา NVIDIA blueprint เดิมใช้ Milvus แต่ Qdrant "
              "มี deployment model ที่ง่ายกว่า"),
        ("h2", "3.3.5 การเลือก Frontend Framework"),
        ("pb", "การตัดสินใจ:"),
        ("p", "Next.js 15 กับ App Router ถูกเลือกจาก built-in API route proxy "
              "(ขจัดปัญหา CORS), React Server Components สำหรับ fast initial page loads "
              "และ TypeScript support Ant Design 5 เป็น UI component library "
              "สำหรับ enterprise-grade components (60+) และ Thai locale support"),
        ("h2", "3.3.6 Authentication Approach"),
        ("pb", "การตัดสินใจ:"),
        ("p", "Stateless JWT กับ HS256 ถูกเลือกจากความง่ายและ horizontal scalability "
              "ข้อจำกัดของ JWT (ไม่สามารถ revoke ทันที) ถูกแก้โดย jti-based in-memory blocklist "
              "และ persistent password_changed_at invalidation"),
        ("h1", "3.4 ระเบียบวิธีการประเมิน"),
        ("h2", "3.4.1 กรอบการประเมิน"),
        ("p", "ระบบถูกประเมิน 2 ระดับ:"),
        ("p", "1. Functional correctness — chatbot สร้างคำตอบที่ถูกต้อง เกี่ยวข้อง และสองภาษาหรือไม่"),
        ("p", "2. Non-functional quality — ระบบตรง latency, concurrency, security, reliability requirements หรือไม่"),
        ("h2", "3.4.2 การออกแบบ Golden Dataset"),
        ("p", "Golden dataset จำนวน 25 test cases ถูกออกแบบให้ครอบคลุมการใช้งาน hotel chatbot ทั้งหมด "
              "(รายละเอียดในบทที่ 6)"),
        ("h2", "3.4.3 วิธีการให้คะแนน"),
        ("p", "คำตอบถือว่า PASS เมื่อเข้าเงื่อนไข 3 ข้อ: "
              "(1) keyword accuracy ≥ 50%, (2) ภาษาถูกต้อง, (3) คำตอบไม่ว่างเปล่า "
              "เกณฑ์ 50% เป็น conservative: คำตอบที่ดีอาจ fail หากใช้ synonyms "
              "ที่ไม่อยู่ใน keyword list"),
        ("h2", "3.4.4 Inter-Model Agreement (Cohen's Kappa)"),
        ("p", "เพื่อวัดว่า local และ cloud models จัดการ test cases เดียวกันได้สม่ำเสมอเพียงใด "
              "ใช้ Cohen's Kappa (κ): κ = (p_o - p_e) / (1 - p_e)"),
        ("h2", "3.4.5 กลยุทธ์การทดสอบโครงสร้างพื้นฐาน"),
        ("p", "นอกเหนือจากการประเมินโมเดล คุณสมบัติ non-functional ของระบบถูกตรวจสอบด้วย "
              "193 automated assertions ใน 4 test suites (รายละเอียดในบทที่ 6)"),
        ("h2", "3.4.6 Performance Benchmarking"),
        ("p", "ประสิทธิภาพถูกวัดใน 3 มิติ: latency (per-request response time), "
              "throughput (concurrent sessions) และ resource efficiency (GPU VRAM, DB pool, cache hit rate) "
              "Benchmarks รันบน target hardware (RTX 5080, 16 GB VRAM)"),
        ("h1", "3.5 จริยธรรมการวิจัย"),
        ("h2", "3.5.1 ความเป็นส่วนตัวของข้อมูลแขก"),
        ("p", "การสนทนาของแขกอาจมีข้อมูลส่วนบุคคล ระบบลดความเสี่ยงด้าน privacy ด้วย: "
              "PII redaction, local LLM deployment (ข้อมูลไม่ออกจากเครือข่ายโรงแรม) "
              "และ audit logging (ทุกการเข้าถึงของ admin ถูกบันทึก)"),
        ("h2", "3.5.2 ความโปร่งใสของ AI"),
        ("p", "Chatbot ไม่แอบอ้างเป็นมนุษย์ System messages แจ้งเมื่อพนักงานเข้ามา "
              "และ AI แนะนำตัวเป็น virtual assistant คำตอบ grounded ใน RAG knowledge base"),
        ("h2", "3.5.3 การลดอคติ"),
        ("p", "การออกแบบสองภาษาทำให้คุณภาพบริการเท่าเทียมกันสำหรับผู้พูดไทยและอังกฤษ "
              "Language detection ใช้ข้อความล่าสุดของแขก (ไม่สมมติจากสัญชาติ) "
              "และ evaluation dataset มี test cases ไทยและอังกฤษอย่างสมดุล"),
    ]

# ============================================================================
# CH4: System Design (การออกแบบระบบ)
# ============================================================================
def get_ch4_specs():
    return [
        ("h1", "4.1 การวิเคราะห์ความต้องการ"),
        ("h2", "4.1.1 Functional Requirements"),
        ("p", "TABLE_CH4_FR_PLACEHOLDER"),
        ("h2", "4.1.2 Non-Functional Requirements"),
        ("p", "TABLE_CH4_NFR_PLACEHOLDER"),
        ("h1", "4.2 สถาปัตยกรรมระบบ"),
        ("h2", "4.2.1 Microservice Topology"),
        ("p", "ระบบประกอบด้วย Docker services จำนวน 5 ตัวที่สื่อสารผ่าน dedicated bridge network:"),
        ("figure", "[Figure 4.1: System architecture diagram — Five Docker containers "
                   "(hotel-api, hotel-ollama, hotel-db, hotel-qdrant, hotel-redis)]"),
        ("p", "TABLE_CH4_SERVICES_PLACEHOLDER"),
        ("h2", "4.2.2 Request Flow"),
        ("p", "ขั้นตอนการประมวลผลคำขอ chat ของแขก:"),
        ("code", "POST /chat {message, session_id}\n"
                 "  ├─ PII redactor (regex scrub)\n"
                 "  ├─ Chat rate limiter (per-session, 30/min)\n"
                 "  ├─ Session lock (per-session asyncio.Lock)\n"
                 "  ├─ Safety router (input validation)\n"
                 "  ├─ LLM concurrency semaphore (acquire slot)\n"
                 "  │   └─ LangGraph Agent\n"
                 "  │       ├─ Primary Assistant (routing)\n"
                 "  │       │   ├─ ToHotelBooking → Booking sub-agent\n"
                 "  │       │   ├─ ToHotelService → Service sub-agent\n"
                 "  │       │   ├─ ToHotelKnowledge → Knowledge sub-agent (RAG)\n"
                 "  │       │   └─ HandleOtherTalk → General conversation\n"
                 "  │       └─ Response\n"
                 "  ├─ Escalation check\n"
                 "  └─ Return ChatResponse"),
        ("h1", "4.3 การออกแบบ LangGraph Agent"),
        ("h2", "4.3.1 State Definition"),
        ("p", "State ของ agent ถูกกำหนดเป็น TypedDict ที่ไหลผ่านทุก node ใน graph:"),
        ("code", "class HotelState(TypedDict):\n"
                 "    messages: Annotated[List[AnyMessage], add_messages]\n"
                 "    session_id: str\n"
                 "    user_id: str\n"
                 "    language: str       # 'th', 'en', or 'auto'\n"
                 "    current_intent: str  # booking, service, knowledge, other\n"
                 "    tool_calls_made: List[Dict[str, Any]]"),
        ("h2", "4.3.2 Primary Router และ Sub-Agent Dispatch"),
        ("p", "Primary assistant ทำหน้าที่เป็น router ไม่ใช่ responder "
              "รับข้อความของแขกและตัดสินใจว่า sub-agent ใดควรจัดการ:"),
        ("code", "def route_primary_assistant(state: HotelState):\n"
                 "    tool_calls = state[\"messages\"][-1].tool_calls\n"
                 "    if tool_calls:\n"
                 "        tool_name = tool_calls[0][\"name\"]\n"
                 "        if tool_name == ToHotelBooking.__name__:\n"
                 "            return \"enter_booking\"\n"
                 "        elif tool_name == ToHotelService.__name__:\n"
                 "            return \"enter_service\"\n"
                 "        elif tool_name == ToHotelKnowledge.__name__:\n"
                 "            return \"enter_knowledge\"\n"
                 "        elif tool_name == HandleOtherTalk.__name__:\n"
                 "            return \"other_talk\"\n"
                 "    return END"),
        ("figure", "[Figure 4.2: LangGraph state machine diagram — START → primary_assistant "
                   "→ conditional edges → sub-agents → tool loops → response]"),
        ("h2", "4.3.3 สถาปัตยกรรม Sub-Agent"),
        ("p", "แต่ละ sub-agent มี: specialized system prompt, restricted tool set "
              "และ independent LLM call ด้วย max_tokens ที่เหมาะสม"),
        ("p", "TABLE_CH4_SUBAGENTS_PLACEHOLDER"),
        ("h1", "4.4 การออกแบบฐานข้อมูล"),
        ("h2", "4.4.1 Entity-Relationship Diagram"),
        ("figure", "[Figure 4.3: ER diagram — room_types (1:N rooms), rooms (1:N reservations), "
                   "guests (1:N reservations), users (1:N audit_log), conversation_history]"),
        ("h2", "4.4.2 ตารางหลัก"),
        ("p", "PostgreSQL schema กำหนด 10 ตาราง:"),
        ("p", "TABLE_CH4_TABLES_PLACEHOLDER"),
        ("h1", "4.5 การออกแบบ RAG Pipeline"),
        ("figure", "[Figure 4.4: RAG pipeline — 10 hotel knowledge files → chunk → embed "
                   "via qwen3-embedding-8b (4096 dims) → Qdrant → query time: top-k search → LLM]"),
        ("h2", "4.5.1 Embedding Configuration"),
        ("p", "TABLE_CH4_EMBED_PLACEHOLDER"),
        ("h1", "4.6 การออกแบบ Authentication และ Authorization"),
        ("h2", "4.6.1 JWT Authentication Flow"),
        ("figure", "[Figure 4.5: JWT authentication sequence diagram — register/login → bcrypt verify "
                   "→ generate JWT {sub, role, iat, exp, jti} → subsequent requests: decode + verify "
                   "→ check jti blocklist → check password_changed_at → return user or 401]"),
        ("h2", "4.6.2 Access Control Matrix"),
        ("figure", "[Figure 4.6: Access control matrix]"),
        ("p", "การออกแบบนี้ทำให้ guest chat flow ไม่ต้อง authentication (ใช้ email-only) "
              "ขณะที่ staff operations ทั้งหมดถูกป้องกันด้วย admin JWT"),
        ("h1", "4.7 การออกแบบ Frontend"),
        ("h2", "4.7.1 Next.js 15 App Router"),
        ("p", "Frontend ใช้ App Router กับ React Server Components สำหรับ server-rendered pages "
              "และ 'use client' directives สำหรับ interactive components:"),
        ("p", "• Server Components สำหรับหน้า static (landing, about) — ไม่มี client JavaScript"),
        ("p", "• Client Components สำหรับ interactive features (chat SSE, admin dashboard)"),
        ("p", "• API Routes (/api/hotel/[...path]) เป็น proxy ไปยัง backend — ขจัด CORS"),
        ("h2", "4.7.2 สถาปัตยกรรม State Management"),
        ("p", "TABLE_CH4_STATE_PLACEHOLDER"),
        ("h1", "4.8 สถาปัตยกรรมการ Deploy"),
        ("figure", "[Figure 4.7: Docker Compose topology — 5 services on hotel-ai-network bridge "
                   "volume mounts, health checks, depends_on with service_healthy]"),
        ("code", "# deploy/compose/docker-compose.hotel.yaml (simplified)\n"
                 "services:\n"
                 "  hotel-ollama:   # GPU, NUM_PARALLEL=2, FLASH_ATTENTION=1\n"
                 "  hotel-db:       # PostgreSQL 16, init-hotel.sql\n"
                 "  hotel-redis:    # Session cache\n"
                 "  hotel-qdrant:   # Vector store\n"
                 "  hotel-api:      # FastAPI, depends_on all above"),
    ]

# ============================================================================
# Table definitions for CH1-CH4
# ============================================================================
TABLES_CH1_4 = {
    "TABLE_CH1_ORG": (
        ["บทที่", "หัวข้อ", "เนื้อหา"],
        [
            ["1", "บทนำ", "ภูมิหลัง วัตถุประสงค์ ขอบเขต"],
            ["2", "การทบทวนวรรณกรรม", "AI ในโรงแรม, LLMs, LangGraph, RAG, web technologies"],
            ["3", "ระเบียบวิธีวิจัย", "แนวทางวิจัย การเลือกเทคโนโลยี การประเมิน"],
            ["4", "การออกแบบระบบ", "สถาปัตยกรรม agent design ฐานข้อมูล RAG auth deployment"],
            ["5", "การพัฒนาระบบ", "รายละเอียดระดับโค้ด backend frontend security scaling"],
            ["6", "การทดสอบและประเมินผล", "เปรียบเทียบโมเดล infrastructure tests benchmarks"],
            ["7", "อภิปรายผล", "Trade-offs ข้อจำกัด การตัดสินใจด้านสถาปัตยกรรม"],
            ["8", "สรุปผลการวิจัย", "ผลงาน งานในอนาคต"],
        ],
    ),
    "TABLE_CH3_SPRINT": (
        ["Sprint", "ช่วงเวลา", "ผลลัพธ์", "Tests เพิ่ม"],
        [
            ["1", "ม.ค. W1-2", "Literature review, technology selection", "—"],
            ["2", "ม.ค. W3-ก.พ. W2", "FastAPI server, PostgreSQL, Qdrant RAG", "RAG accuracy"],
            ["3", "ก.พ. W3-มี.ค. W1", "LangGraph multi-agent, hotel tools", "Integration"],
            ["4", "มี.ค. W2-3", "Next.js frontend, chat SSE, admin dashboard", "—"],
            ["5", "มี.ค. W4-เม.ย. W1", "JWT auth, RBAC, rate limiting, audit log", "72+38 auth"],
            ["6", "เม.ย. W1", "Chat scaling (LLM semaphore, session locks, cache)", "46+37 scaling"],
            ["7", "เม.ย. W1-2", "Ollama GPU tuning, reranker removal, prompt optimization", "Benchmarks"],
            ["8", "เม.ย. W2", "Model evaluation (25 cases), thesis writing", "Evaluation"],
        ],
    ),
    "TABLE_CH3_CRITERIA": (
        ["เกณฑ์", "น้ำหนัก", "เหตุผล"],
        [
            ["ความเหมาะสมกับ domain โรงแรม", "สูง", "รองรับสองภาษา booking แบบ real-time multi-turn"],
            ["ความสามารถ deploy local", "สูง", "ข้อมูลแขก PII ต้องอยู่ on-premise"],
            ["Open-source", "ปานกลาง", "Reproducibility สำหรับการประเมิน"],
            ["ความสมบูรณ์ของ community", "ปานกลาง", "เสถียร documentation ดี"],
            ["ความเข้ากันได้", "ปานกลาง", "ทำงานร่วมกันใน Docker Compose stack"],
        ],
    ),
    "TABLE_CH3_FRAMEWORK": (
        ["เกณฑ์", "LangGraph", "CrewAI", "AutoGen", "LangChain (plain)"],
        [
            ["State persistence", "✅ AsyncPostgresSaver", "❌ Manual", "❌ Manual", "❌ Manual"],
            ["Cyclic tool loops", "✅ Native graph edges", "⚠️ Limited", "⚠️ Limited", "❌ Sequential"],
            ["Time-travel debug", "✅ Checkpoint replay", "❌", "❌", "❌"],
            ["Sub-agent routing", "✅ Conditional edges", "✅ Role-based", "✅ Conversation", "⚠️ Router"],
            ["Production maturity", "✅ LangChain ecosystem", "⚠️ Growing", "⚠️ Research", "✅ Mature"],
        ],
    ),
    "TABLE_CH3_MODEL": (
        ["เกณฑ์", "Qwen3.5 Opus 9B (Local)", "Qwen3 Max (Cloud)"],
        [
            ["Deployment", "On-premise (Ollama, single GPU)", "OpenRouter API"],
            ["ต้นทุนต่อ query", "$0 (amortized)", "~$0.001-0.01"],
            ["สองภาษา", "✅ Native", "✅ Native"],
            ["Tool calling", "✅ Structured JSON", "✅ Structured JSON"],
            ["Context window", "4,096 tokens", "262,144 tokens"],
            ["Data privacy", "✅ On-premise", "❌ Third-party API"],
            ["Concurrent", "2 (GPU-bound)", "Unlimited"],
        ],
    ),
    "TABLE_CH3_VECTOR": (
        ["เกณฑ์", "Qdrant", "Milvus", "Pinecone", "ChromaDB"],
        [
            ["Open-source", "✅", "✅", "❌ (cloud only)", "✅"],
            ["Docker-native", "✅", "⚠️ (complex)", "❌", "✅"],
            ["Performance", "✅ Rust-based", "✅", "✅", "⚠️ Python"],
            ["Bilingual", "✅", "✅", "✅", "✅"],
        ],
    ),
    "TABLE_CH4_FR": (
        ["ID", "ความต้องการ", "ลำดับ"],
        [
            ["FR1", "สนทนาสองภาษา (ไทย/อังกฤษ) พร้อม language detection อัตโนมัติ", "Must"],
            ["FR2", "ตรวจสอบห้องว่างด้วย real-time database queries", "Must"],
            ["FR3", "Reservation CRUD ผ่าน natural language", "Must"],
            ["FR4", "Q&A ข้อมูลโรงแรม (facilities, policies, dining, spa, transport)", "Must"],
            ["FR5", "Check-in / check-out operations", "Must"],
            ["FR6", "Admin dashboard พร้อม session monitoring และ chat intervention", "Must"],
            ["FR7", "JWT authentication แบ่ง user/admin", "Must"],
            ["FR8", "Runtime switching ระหว่าง local กับ cloud LLM", "Should"],
            ["FR9", "Dynamic pricing (early-bird / last-minute)", "Should"],
            ["FR10", "PII redaction ก่อน LLM processing", "Should"],
            ["FR11", "Auto escalation เมื่อแขก frustrated", "Could"],
            ["FR12", "Mock payment link generation", "Could"],
        ],
    ),
    "TABLE_CH4_NFR": (
        ["ID", "ความต้องการ", "เป้าหมาย"],
        [
            ["NFR1", "Response latency (warm, single user)", "< 10 วินาที"],
            ["NFR2", "Concurrent users ไม่เสื่อมคุณภาพ", "≥ 2 simultaneous chats"],
            ["NFR3", "Knowledge retrieval accuracy", "≥ 90%"],
            ["NFR4", "System availability", "Health check + graceful degradation"],
            ["NFR5", "Security", "bcrypt, JWT+jti, rate limiting, audit log"],
            ["NFR6", "Deployment", "Docker Compose, single-command startup"],
        ],
    ),
    "TABLE_CH4_SERVICES": (
        ["Container", "Service", "Port", "วัตถุประสงค์"],
        [
            ["hotel-api", "FastAPI + LangGraph", "8088", "Application server"],
            ["hotel-ollama", "Ollama", "11435", "Local LLM inference (GPU)"],
            ["hotel-db", "PostgreSQL 16", "5433", "Hotel database"],
            ["hotel-qdrant", "Qdrant", "6334", "Vector store"],
            ["hotel-redis", "Redis 7", "6380", "Session cache"],
        ],
    ),
    "TABLE_CH4_SUBAGENTS": (
        ["Sub-Agent", "Tools", "max_tokens", "วัตถุประสงค์"],
        [
            ["Booking", "12 tools", "2048", "Full booking lifecycle"],
            ["Service", "2 tools", "1024", "Amenity/service requests"],
            ["Knowledge", "RAG search (direct)", "1024", "Hotel information Q&A"],
            ["Other Talk", "None", "512", "Greetings, thanks, off-topic"],
        ],
    ),
    "TABLE_CH4_TABLES": (
        ["ตาราง", "Rows (seeded)", "วัตถุประสงค์"],
        [
            ["room_types", "4", "Standard, Deluxe, Suite, Penthouse"],
            ["rooms", "~50", "Individual rooms with floor, status, view"],
            ["guests", "Dynamic", "Guest profiles (email as unique ID)"],
            ["reservations", "Dynamic", "Bookings with confirmation number (HTL...)"],
            ["users", "Dynamic", "Auth accounts"],
            ["audit_log", "Dynamic", "Admin action trail (JSONB)"],
            ["conversation_history", "Dynamic", "Chat messages"],
            ["service_requests", "Dynamic", "Amenity requests"],
            ["payment_links", "Dynamic", "Mock payment tokens"],
            ["hotel_services", "~10", "Available services catalog"],
        ],
    ),
    "TABLE_CH4_EMBED": (
        ["พารามิเตอร์", "ค่า", "เหตุผล"],
        [
            ["Model", "qwen/qwen3-embedding-8b", "รองรับสองภาษา ไทย/อังกฤษ"],
            ["Dimensions", "4096", "Native output dimensionality"],
            ["Chunk size", "Auto (~3,200 chars)", "80% ของ token limit × 4 chars/token"],
            ["Chunk overlap", "20%", "รักษา context ระหว่าง chunk boundaries"],
            ["Distance metric", "Cosine similarity", "มาตรฐานสำหรับ text embeddings"],
        ],
    ),
    "TABLE_CH4_STATE": (
        ["Concern", "Solution", "เหตุผล"],
        [
            ["Global state (auth, theme)", "Zustand stores", "น้อย boilerplate กว่า Redux"],
            ["Server data (rooms, bookings)", "SWR hooks", "stale-while-revalidate (RFC 5861)"],
            ["Form state", "React local state", "ไม่ต้อง global store"],
            ["Chat messages", "Zustand + SSE", "Real-time streaming + persistent client state"],
        ],
    ),
}

# ============================================================================
# References section (บรรณานุกรม)
# ============================================================================
def get_references_specs():
    """References stay in English (standard for Thai academic papers). Thai heading only."""
    refs = [
        ("pb", "AI in Hospitality"),
        ("p", "Buhalis, D. (2020). Technology in tourism — from ICTs to eTourism and smart tourism towards ambient intelligence tourism. Tourism Review, 75(1)."),
        ("p", "Buhalis, D., & Leung, R. (2018). Smart hospitality — Interconnectivity and interoperability towards an ecosystem. International Journal of Hospitality Management, 71, 41–50."),
        ("p", "Buhalis, D., & Moldavska, I. (2022). Voice assistants in hospitality: using artificial intelligence for customer service. Journal of Hospitality and Tourism Technology, 13."),
        ("p", "Buhalis, D., O'Connor, P., & Leung, R. (2023). Smart hospitality: from smart cities and smart tourism towards agile business ecosystems. International Journal of Contemporary Hospitality Management, 35."),
        ("p", "MDPI Tourism & Hospitality. (2023). Artificial intelligence in tourism through chatbot support in the booking process. Tourism & Hospitality, 6(1), 36."),
        ("p", "Taylor & Francis. (2024). Personalizing guest experience with generative AI in the hotel industry. Current Issues in Tourism."),
        ("p", "Taylor & Francis. (2025). Two decades of chatbot research in tourism and hospitality: Bibliometric analysis and future directions."),
        ("p", "Taylor & Francis. (2026). Conversational AI in hospitality and tourism: a bibliometric–systematic review."),
        ("pb", "Hotel Property Management Systems"),
        ("p", "AltexSoft. (2023). What Is a Hotel Property Management System (PMS): Complete Guide."),
        ("p", "Atlantis Press. (2013). Research and Design of Hotel Management System Model. Proceedings of ICETIS '13."),
        ("p", "Cloudbeds. (2023). Housekeeping room conditions."),
        ("p", "Oracle. (2023). OPERA — Housekeeping Room Status."),
        ("pb", "LLM and Agent Frameworks"),
        ("p", "Liu, Y., et al. (2024). Exploration of LLM Multi-Agent Application Implementation Based on LangGraph+CrewAI. arXiv:2411.18241."),
        ("p", "Vaswani, A., et al. (2017). Attention Is All You Need. Advances in Neural Information Processing Systems, 30."),
        ("pb", "Web Technologies"),
        ("p", "Ant Design. (2024). Ant Design — A UI Design Language and React UI Library."),
        ("p", "Gao, Z., Bird, C., & Barr, E. T. (2017). To Type or Not to Type: Quantifying Detectable Bugs in JavaScript. ICSE 2017."),
        ("p", "Next.js. (2024). Official Documentation — App Router."),
        ("p", "Nottingham, M. (2010). HTTP Cache-Control Extensions for Stale Content. RFC 5861, IETF."),
        ("p", "Osmani, A. (2024). React Server Components, Next.js App Router and examples."),
        ("p", "Salah, M. A. (2024). State Management in React: Redux vs Zustand. ResearchGate."),
        ("pb", "Security"),
        ("p", "OWASP. (2023). OWASP Top 10 — 2021."),
        ("p", "RFC 7519. (2015). JSON Web Token (JWT). IETF."),
        ("pb", "Evaluation Methods"),
        ("p", "Cohen, J. (1960). A coefficient of agreement for nominal scales. Educational and Psychological Measurement, 20(1), 37–46."),
        ("p", "DeepEval. (2024). RAG Evaluation Metrics. https://docs.confident-ai.com/docs/metrics-introduction"),
        ("pb", "Software and Libraries"),
        ("p", "Docker. (2024). Docker Documentation. FastAPI. (2024). FastAPI Documentation. LangGraph. (2024). LangGraph Documentation. Ollama. (2024). Ollama Documentation. OpenRouter. (2024). OpenRouter API Documentation. PostgreSQL. (2024). PostgreSQL 16 Documentation. Qdrant. (2024). Qdrant Vector Database Documentation. Redis. (2024). Redis Documentation. TypeScript. (2024). TypeScript Documentation."),
    ]
    return refs

# ============================================================================
# In-place translation for Appendix headings/prose
# ============================================================================
APPENDIX_TRANSLATIONS = {
    # AP_A headings
    "Test Results and Development Timeline": "ผลการทดสอบและลำดับเวลาการพัฒนา",
    "A.1 Model Evaluation — Full 25-Case Results": "A.1 การประเมินโมเดล — ผลลัพธ์ 25 กรณีทดสอบ",
    "Test Configuration": "การตั้งค่าการทดสอบ",
    "A.1.1 Knowledge Retrieval (K01–K08)": "A.1.1 การค้นหาข้อมูล (K01–K08)",
    "A.1.2 Booking Operations (B01–B06)": "A.1.2 การจองห้องพัก (B01–B06)",
    "A.1.3 Greetings & Small Talk (G01–G04)": "A.1.3 การทักทายและสนทนาทั่วไป (G01–G04)",
    "A.1.4 Language Handling (L01–L03)": "A.1.4 การจัดการภาษา (L01–L03)",
    "A.1.5 Edge Cases (E01–E04)": "A.1.5 กรณีพิเศษ (E01–E04)",
    "A.2 Model Evaluation Summary": "A.2 สรุปผลการประเมินโมเดล",
    "A.3 Infrastructure Test Results (193/193)": "A.3 ผลการทดสอบโครงสร้างพื้นฐาน (193/193)",
    "A.4 Performance Benchmarks": "A.4 ผลการวัดประสิทธิภาพ",
    "A.5 Development Timeline (Gantt Chart)": "A.5 ลำดับเวลาการพัฒนา (Gantt Chart)",
    "A.6 Source Code Repository Structure": "A.6 โครงสร้าง Repository ของซอร์สโค้ด",
    # AP_B headings
    "User Manual": "คู่มือผู้ใช้",
    "B.1 System Requirements": "B.1 ข้อกำหนดของระบบ",
    "Hardware": "ฮาร์ดแวร์",
    "Software": "ซอฟต์แวร์",
    "B.2 Installation": "B.2 การติดตั้ง",
    "Step 1: Clone the Repository": "ขั้นตอนที่ 1: Clone Repository",
    "Step 2: Configure Environment": "ขั้นตอนที่ 2: ตั้งค่า Environment",
    "Step 3: Start the Docker Stack": "ขั้นตอนที่ 3: เริ่ม Docker Stack",
    "Step 4: Pull the Local Model": "ขั้นตอนที่ 4: ดาวน์โหลด Local Model",
    "Step 5: Ingest Hotel Knowledge": "ขั้นตอนที่ 5: นำเข้าฐานความรู้โรงแรม",
    "Step 6: Verify": "ขั้นตอนที่ 6: ตรวจสอบ",
    "B.2b Frontend Installation (Next.js)": "B.2b การติดตั้ง Frontend (Next.js)",
    "Prerequisites": "ข้อกำหนดเบื้องต้น",
    "Step 1: Clone and Install": "ขั้นตอนที่ 1: Clone และติดตั้ง",
    "Step 3: Start Development Server": "ขั้นตอนที่ 3: เริ่ม Development Server",
    "Step 4: Verify": "ขั้นตอนที่ 4: ตรวจสอบ",
    "Frontend Pages Overview": "ภาพรวมหน้า Frontend",
    "Frontend Technology Stack": "Frontend Technology Stack",
    "B.3 Guest Usage (Chat)": "B.3 การใช้งานสำหรับแขก (Chat)",
    "Starting a Conversation": "เริ่มการสนทนา",
    "Multi-Turn Conversation": "การสนทนาหลายรอบ",
    "B.4 Admin Dashboard": "B.4 Admin Dashboard",
    "B.5 Runtime Model Switching": "B.5 การสลับโมเดลขณะ Runtime",
    "B.6 Running Tests": "B.6 การรันการทดสอบ",
    "B.7 Troubleshooting": "B.7 การแก้ไขปัญหา",
    "B.8 API Reference": "B.8 เอกสาร API",
}

PROSE_TRANSLATIONS = {
    "This downloads ~6.5 GB of model weights. First run takes 5–10 minutes.":
        "การดาวน์โหลดน้ำหนักโมเดลประมาณ 6.5 GB ครั้งแรกใช้เวลา 5-10 นาที",
    "This embeds the 10 hotel knowledge documents into Qdrant. Takes ~2 minutes.":
        "การ embed เอกสารความรู้โรงแรม 10 ฉบับเข้า Qdrant ใช้เวลาประมาณ 2 นาที",
    "This starts 5 services:":
        "คำสั่งนี้จะเริ่ม 5 services:",
    "Key settings to configure:":
        "การตั้งค่าสำคัญ:",
    "Copy the example environment file and edit:":
        "คัดลอกไฟล์ environment ตัวอย่างและแก้ไข:",
    "Frontend will be available at `http://localhost:3000`.":
        "Frontend จะพร้อมใช้งานที่ http://localhost:3000",
    "The system automatically:":
        "ระบบจะทำงานอัตโนมัติ:",
    "Send a POST request to `/chat` with a message. No authentication required:":
        "ส่ง POST request ไปยัง /chat พร้อมข้อความ ไม่ต้อง authentication:",
}

def translate_in_place(doc, start_idx, end_idx):
    """Translate known headings and prose in place."""
    count = 0
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text:
            continue
        # Check heading translations
        if text in APPENDIX_TRANSLATIONS:
            new_text = APPENDIX_TRANSLATIONS[text]
            for run in p.runs:
                if run.text.strip():
                    run.text = new_text
                    count += 1
                    break
        # Check prose translations
        elif text in PROSE_TRANSLATIONS:
            new_text = PROSE_TRANSLATIONS[text]
            for run in p.runs:
                if run.text.strip():
                    run.text = new_text
                    count += 1
                    break
    return count


# ============================================================================
# MAIN
# ============================================================================
def main():
    print("=" * 60)
    print("Patching remaining English sections with Thai")
    print("  CH1-CH4 + References + Appendix headings")
    print("=" * 60)

    if not SRC.exists():
        print(f"ERROR: {SRC} not found")
        sys.exit(1)

    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    body = doc.element.body

    # Find chapter boundaries
    print("\n1. Scanning chapters...")
    chapters = find_chapter_indices(doc)
    for ch, idx in sorted(chapters.items(), key=lambda x: str(x[0])):
        print(f"  {ch}: P{idx} — {doc.paragraphs[idx].text.strip()[:50]}")

    int_chs = sorted([(k, v) for k, v in chapters.items() if isinstance(k, int)], key=lambda x: x[1])

    def get_next_elem(ch_num):
        found = False
        for k, v in int_chs:
            if found:
                return doc.paragraphs[v]._element
            if k == ch_num:
                found = True
        if "appendix" in chapters:
            return doc.paragraphs[chapters["appendix"]]._element
        return None

    # Replace CH4, CH3, CH2, CH1 (reverse order)
    for ch_num, specs_fn, label in [
        (4, get_ch4_specs, "CH4 System Design"),
        (3, get_ch3_specs, "CH3 Methodology"),
        (2, get_ch2_specs, "CH2 Literature Review"),
        (1, get_ch1_specs, "CH1 Introduction"),
    ]:
        if ch_num not in chapters:
            print(f"\n  [SKIP] {label}")
            continue
        start_elem = doc.paragraphs[chapters[ch_num]]._element
        end_elem = get_next_elem(ch_num)
        print(f"\n2.{ch_num}. Replacing {label}...")
        removed = remove_between_elements(body, start_elem, end_elem)
        print(f"  Removed {removed} old elements")
        nodes = build_nodes(doc, specs_fn())
        insert_after_element(body, start_elem, nodes)
        print(f"  Inserted {len(nodes)} Thai paragraphs")

    # Insert References between CH8 and Appendix
    print("\n3. Adding References (บรรณานุกรม)...")
    if 8 in chapters and "appendix" in chapters:
        # Find last element of CH8 (before appendix)
        appendix_elem = doc.paragraphs[chapters["appendix"]]._element
        ref_heading = make_para(doc, "บรรณานุกรม", "TU_Chapter", bold=True)
        ref_nodes = [ref_heading] + build_nodes(doc, get_references_specs())
        parent = appendix_elem.getparent()
        idx = list(parent).index(appendix_elem)
        for j, node in enumerate(ref_nodes):
            parent.insert(idx + j, node)
        print(f"  Inserted {len(ref_nodes)} reference paragraphs")

    # Insert tables at placeholders
    print("\n4. Inserting tables...")
    table_count = 0
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for table_id, (headers, rows) in TABLES_CH1_4.items():
            placeholder = f"{table_id}_PLACEHOLDER"
            if text == placeholder:
                try:
                    create_word_table(doc, headers, rows, p._element)
                    body.remove(p._element)
                    table_count += 1
                    print(f"  [OK] {table_id}")
                except Exception as e:
                    print(f"  [FAIL] {table_id}: {e}")
                break
    print(f"  Total: {table_count}")

    # Translate appendix headings in-place
    print("\n5. Translating appendix headings in-place...")
    # Re-scan after changes
    chapters2 = find_chapter_indices(doc)
    if "ap_a" in chapters2:
        ap_a_end = chapters2.get("ap_b", chapters2.get("ap_c", chapters2.get("bio", len(doc.paragraphs))))
        count_a = translate_in_place(doc, chapters2["ap_a"], ap_a_end)
        print(f"  AP_A: translated {count_a} paragraphs")
    if "ap_b" in chapters2:
        ap_b_end = chapters2.get("ap_c", chapters2.get("bio", len(doc.paragraphs)))
        count_b = translate_in_place(doc, chapters2["ap_b"], ap_b_end)
        print(f"  AP_B: translated {count_b} paragraphs")

    # Save
    doc.save(str(OUT))
    sz = os.path.getsize(str(OUT)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUT} ({sz:.0f} KB)")
    print("=" * 60)

if __name__ == "__main__":
    main()
