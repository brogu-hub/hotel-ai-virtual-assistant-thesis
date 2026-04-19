#!/usr/bin/env python3
"""
patch_v6_final.py
=================
1. Fix ALL heading/paragraph fonts: remove entire rPr on non-code runs so TU_ styles take full control
2. Add WORKFLOW.md diagrams as ASCII figures with Thai explanations
3. Add NVIDIA AI Blueprint references to CH5 phases
4. Add NVIDIA Blueprint to References

Input:  thesis/thesis_final_v5.docx
Output: thesis/thesis_final_v6.docx
"""
import os, sys, re, shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final_v5.docx")
OUT = Path("thesis/thesis_final_v6.docx")
FIGURES_DIR = Path("thesis/figures")


# ============================================================================
# 1. AGGRESSIVE FONT FIX — remove entire rPr on non-code runs
# ============================================================================
def fix_fonts_aggressive(doc):
    """Remove ALL run-level rPr children (bold, italic, fonts, size) on non-Consolas runs.
    The TU_ styles define everything. Run-level overrides cause font mismatch."""
    fixed = 0
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        # Only fix TU_ styled paragraphs
        if not style_name.startswith("TU_"):
            continue
        for run in p.runs:
            rpr = run._element.find(qn('w:rPr'))
            if rpr is None:
                continue
            # Check if this is a Consolas (code) run — keep those
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is not None:
                ascii_font = rfonts.get(qn('w:ascii'), '')
                if ascii_font == 'Consolas':
                    continue
            # Remove the entire rPr — style will take over
            run._element.remove(rpr)
            fixed += 1
    return fixed


# ============================================================================
# Helpers
# ============================================================================
def make_para(doc, text, style_name, bold=False, italic=False,
              font_name=None, font_size=None, alignment=None):
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        # Only add rPr for code blocks (Consolas)
        if font_name:
            rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), font_name)
            rf.set(qn('w:hAnsi'), font_name)
            rf.set(qn('w:cs'), font_name)
            rPr.append(rf)
            if font_size:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(font_size * 2)))
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), str(int(font_size * 2)))
                rPr.append(sz)
                rPr.append(szCs)
            r.append(rPr)
        # NO rPr for non-code — let style handle bold/italic/font
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p

def make_code_para(doc, text):
    return make_para(doc, text, "TU_Paragraph_Normal", font_name="Consolas", font_size=10)

def build_nodes(doc, specs):
    nodes = []
    for spec in specs:
        kind, text = spec[0], spec[1] if len(spec) > 1 else ""
        if kind == "h1":
            nodes.append(make_para(doc, text, "TU_Sub-heading 1"))
        elif kind == "h2":
            nodes.append(make_para(doc, text, "TU_Sub-heading 2"))
        elif kind == "h3":
            nodes.append(make_para(doc, text, "TU_Sub-heading 3"))
        elif kind == "p":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal"))
        elif kind == "pb":
            # Bold paragraph — use TU_Sub-heading 3 (smallest heading) for bold
            # since we can't add rPr bold without overriding font
            nodes.append(make_para(doc, text, "TU_Sub-heading 3"))
        elif kind == "pi":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal"))
        elif kind == "code":
            for line in text.split("\n"):
                nodes.append(make_code_para(doc, line if line else " "))
        elif kind == "figure":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", alignment="center"))
    return nodes

def insert_before_element(body, ref_elem, nodes):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for j, node in enumerate(nodes):
        parent.insert(idx + j, node)

def find_chapter_indices(doc):
    chapters = {}
    HEADING_MAP = {
        "introduction": 1, "literature review": 2, "methodology": 3,
        "system design": 4, "implementation": 5, "testing and evaluation": 6,
        "discussion": 7, "conclusion": 8,
    }
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            lower = text.lower()
            for ht, cn in HEADING_MAP.items():
                if ht in lower and cn not in chapters:
                    chapters[cn] = i
                    break
        if style == "TU_Chapter":
            if "ภาคผนวก" == text.strip():
                chapters["appendix"] = i
            elif "บรรณานุกรม" in text:
                chapters["references"] = i
    return chapters

def find_para(doc, text_fragment, start=0, end=None):
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, min(end, len(doc.paragraphs))):
        if text_fragment in doc.paragraphs[i].text:
            return i
    return None


# ============================================================================
# WORKFLOW.md DIAGRAMS — to inject into CH4 and CH5
# ============================================================================

# --- For CH4 (System Design) — after section 4.2 Request Flow ---
WORKFLOW_CH4 = [
    ("p", ""),
    ("h2", "4.2.3 Reservation Status Lifecycle"),
    ("p", "วงจรสถานะของการจองห้องพักแสดงการเปลี่ยนสถานะผ่าน API tools:"),
    ("code", "  create_reservation()        confirm_reservation()\n"
             "         |                           |\n"
             "         v                           v\n"
             "     +--------+              +-----------+\n"
             "     | pending | ----------> | confirmed |\n"
             "     +--------+              +-----------+\n"
             "         |                        |    |\n"
             "         |   cancel_reservation() |    |  check_in_guest()\n"
             "         |         |              |    |\n"
             "         v         v              v    v\n"
             "   +-----------+              +------------+\n"
             "   | cancelled |              | checked_in |\n"
             "   +-----------+              +------------+\n"
             "                                     |\n"
             "                                     |  check_out_guest()\n"
             "                                     v\n"
             "                              +--------------+\n"
             "                              | checked_out  |\n"
             "                              +--------------+"),
    ("p", "ทุกการเปลี่ยนสถานะใช้ UPDATE (ไม่มี hard delete) "
          "ระบบไม่มีการชำระเงินจริง — payment_status คงอยู่ที่ pending ตลอด"),

    ("h2", "4.2.4 Guest Identification"),
    ("p", "ระบบ identification ของแขกไม่ต้องสร้างบัญชี:"),
    ("code", "  +-------------------+     +---------------------------+\n"
             "  | Email address     |     | Confirmation number       |\n"
             "  | (primary key)     |     | (HTL + YYMMDD + seq)      |\n"
             "  |                   |     |                           |\n"
             "  | ใช้สำหรับ:        |     | ใช้สำหรับ:                |\n"
             "  | - จองห้องใหม่     |     | - ค้นหาการจอง             |\n"
             "  | - ค้นหาประวัติ    |     | - แก้ไข / ยกเลิก         |\n"
             "  | - Auto-register   |     | - Check-in ที่ front desk |\n"
             "  +-------------------+     +---------------------------+"),
    ("p", "เมื่อ email ใหม่ถูกใช้จอง → Guest record ถูกสร้างอัตโนมัติ (ไม่ต้องกรอกฟอร์ม) "
          "→ Loyalty tier = Standard, points = 0"),

    ("h2", "4.2.5 Authentication: สองระบบ Identity"),
    ("code", "  +----------------------+              +--------------------------+\n"
             "  |  GUESTS table        |              |  USERS table             |\n"
             "  |  (hotel profile)     |              |  (login credentials)     |\n"
             "  +----------------------+              +--------------------------+\n"
             "  | guest_id (PK)        |<------+------| guest_id (FK, nullable)  |\n"
             "  | email (unique)       |       |      | user_id (PK)             |\n"
             "  | first/last name      |       |      | username (unique)        |\n"
             "  | phone, nationality   |       |      | password_hash (bcrypt)   |\n"
             "  | loyalty_tier/points  |       |      | role ('user' | 'admin')  |\n"
             "  +----------------------+       |      +--------------------------+\n"
             "         ^                       |\n"
             "         | ใช้โดย                |\n"
             "  +-----------------------+      |      +---------------------------+\n"
             "  | Guest chat flow       |      |      | Registered users          |\n"
             "  | (ไม่ต้อง login)       |      |      | (JWT bearer token)        |\n"
             "  | email-only identity   |      +------| users.guest_id -> guests  |\n"
             "  +-----------------------+             +---------------------------+"),
    ("p", "Guest chat flow ไม่ต้อง authentication — ใช้ email-only identification "
          "ส่วน admin operations ทั้งหมดต้อง JWT token"),
]

# --- For CH5 — Admin Monitoring, Dynamic Pricing, Escalation, Time-travel ---
WORKFLOW_CH5 = [
    ("p", ""),
    ("h1", "5.8 Admin Monitoring และ Chat Intervention"),
    ("p", "ระบบ admin monitoring ให้พนักงานโรงแรมติดตามและเข้าแทรกการสนทนาได้แบบ real-time:"),
    ("code", "  Admin Dashboard                    Server                     Guest Chat\n"
             "       |                               |                            |\n"
             "       |  GET /admin/sessions           |                            |\n"
             "       |------------------------------>|                            |\n"
             "       |  [sessions with previews,     |                            |\n"
             "       |   status: bot_active]         |                            |\n"
             "       |<------------------------------|                            |\n"
             "       |                               |                            |\n"
             "       |  POST /admin/sessions/        |                            |\n"
             "       |       {id}/takeover           |                            |\n"
             "       |------------------------------>|                            |\n"
             "       |  status -> admin_controlled   |  [System] Staff joined     |\n"
             "       |<------------------------------|--------------------------->|\n"
             "       |                               |                            |\n"
             "       |  POST /admin/chat/override    |                            |\n"
             "       |  \"Let me help you with...\"    |  [Admin] \"Let me help...\"  |\n"
             "       |------------------------------>|--------------------------->|\n"
             "       |                               |                            |\n"
             "       |  POST /admin/sessions/        |                            |\n"
             "       |       {id}/release            |                            |\n"
             "       |------------------------------>|  [System] AI resumed       |\n"
             "       |  status -> bot_active         |--------------------------->|\n"
             "       |<------------------------------|                            |"),
    ("p", "เมื่อ admin takeover แล้ว ข้อความของแขกยังถูกบันทึก แต่ bot จะหยุดตอบ "
          "admin ส่งข้อความผ่าน /admin/chat/override จนกว่าจะ release กลับให้ bot"),

    ("h1", "5.9 Dynamic Pricing"),
    ("p", "ระบบคำนวณราคาอัตโนมัติตามจำนวนวันก่อน check-in:"),
    ("code", "  วันก่อน check-in        ตัวคูณ          ป้ายราคา\n"
             "  ====================    ==========    ==================\n"
             "  30+ วัน                   0.85x       Early Bird ลด 15%\n"
             "  14-29 วัน                 0.90x       Advance ลด 10%\n"
             "  7-13 วัน                  1.00x       อัตราปกติ\n"
             "  1-6 วัน                   1.20x       Last-Minute +20%\n"
             "  วันเดียวกัน               1.30x       Same-Day +30%"),
    ("p", "ราคาถูกคำนวณอัตโนมัติใน create_reservation() "
          "tool calculate_dynamic_price แสดงราคาจริงก่อนจอง"),

    ("h1", "5.10 Time-Travel / Checkpoint Replay"),
    ("p", "LangGraph checkpointer บันทึกทุก state transition ของการสนทนา "
          "admin สามารถย้อนเวลาและ replay จาก checkpoint ใดก็ได้:"),
    ("code", "  Admin                          Server                      PostgreSQL\n"
             "       |                            |                            |\n"
             "       |  GET /admin/sessions/      |                            |\n"
             "       |       {id}/states          |                            |\n"
             "       |--------------------------->|  aget_state_history()      |\n"
             "       |  [Step 0: guest msg]       |<--------------------------||\n"
             "       |  [Step 1: router]          |                            |\n"
             "       |  [Step 3: availability]    |                            |\n"
             "       |  [Step 5: booking done]    |                            |\n"
             "       |<---------------------------|                            |\n"
             "       |                            |                            |\n"
             "       |  POST .../rollback         |  aupdate_state()           |\n"
             "       |  checkpoint_id=Step3       |--------------------------->|\n"
             "       |--------------------------->|  (fork from checkpoint)    |\n"
             "       |  \"rolled back to step 3\"   |<--------------------------||\n"
             "       |<---------------------------|                            |\n"
             "       |                            |                            |\n"
             "       |  POST .../replay           |  ainvoke(new_msg, config)  |\n"
             "       |  message=\"Try Suite\"       |--------------------------->|\n"
             "       |--------------------------->|  (new branch from step 3)  |\n"
             "       |  \"Suite available at...\"   |<--------------------------||\n"
             "       |<---------------------------|                            |"),
    ("p", "คุณสมบัติ time-travel ช่วยให้ admin debug ปัญหาการ routing ของ agent "
          "และทดสอบ \"what-if\" scenarios โดย branch จาก checkpoint เดิม"),

    ("h1", "5.11 Mock Payment Flow"),
    ("code", "  Guest                     Bot                    Server              DB\n"
             "    |                        |                        |                 |\n"
             "    |  \"Confirm booking\"     |                        |                 |\n"
             "    |---------------------->|                        |                 |\n"
             "    |                        |  generate_payment_link |                 |\n"
             "    |                        |---------------------->|  INSERT token   |\n"
             "    |                        |                        |---------------->|\n"
             "    |  \"Pay here:            |                        |  expires 30min  |\n"
             "    |   /checkout/{token}\"   |<-----------------------|                 |\n"
             "    |<-----------------------|                        |                 |\n"
             "    |                        |                        |                 |\n"
             "    |  GET /payment/{token}  |                        |                 |\n"
             "    |----------------------------------------------->|  booking details|\n"
             "    |  {amount, room, dates} |                        |<----------------||\n"
             "    |<-----------------------------------------------|                 |\n"
             "    |                        |                        |                 |\n"
             "    |  POST /payment/        |                        |                 |\n"
             "    |       {token}/complete |                        |                 |\n"
             "    |----------------------------------------------->|  status=paid    |\n"
             "    |  \"Payment successful\"  |                        |---------------->|\n"
             "    |<-----------------------------------------------|                 |"),
    ("p", "Payment links เป็น mock (UUID token, หมดอายุ 30 นาที) "
          "ใน production จะแทนที่ด้วย Stripe หรือ PromptPay"),

    ("h1", "5.12 Audit Log"),
    ("p", "ทุกการกระทำของ admin, เหตุการณ์ authentication และการเข้าถึงข้อมูลแขก "
          "ถูกบันทึกในตาราง audit_log:"),
    ("code", "  Admin Dashboard              /admin/audit API            DB\n"
             "       |                             |                     |\n"
             "       | GET /admin/audit?           |                     |\n"
             "       |   action_prefix=admin.      |                     |\n"
             "       |   &actor_username=john_doe  |                     |\n"
             "       |   &limit=50&offset=0        |                     |\n"
             "       |---------------------------->|                     |\n"
             "       |                             | require_admin       |\n"
             "       |                             | list_audit_entries()|\n"
             "       |                             |-------------------->|\n"
             "       |                             | COUNT + LIMIT/OFFSET|\n"
             "       |                             |<--------------------||\n"
             "       | {entries, total, has_more}  |                     |\n"
             "       |<----------------------------|                     |\n"
             "       |                             | audit: AUDIT_VIEWED |  <- meta-audit\n"
             "       |                             |-------------------->|"),
    ("p", "Meta-audit: ทุกการเรียก GET /admin/audit ถูกบันทึกเป็น admin.audit.viewed "
          "ทำให้ตรวจสอบได้ว่าใครค้นหา audit log เมื่อไหร่ "
          "เมื่อ admin อ่านประวัติสนทนาของแขก จะสร้าง admin.session.viewed "
          "เพื่อ compliance กับกฎหมายคุ้มครองข้อมูลส่วนบุคคล"),

    ("h1", "5.13 Knowledge Cache Hot Path"),
    ("code", "  search_hotel_knowledge(\"what time is breakfast?\")\n"
             "       |\n"
             "       v\n"
             "  KnowledgeCache.get(\"what time is breakfast?\")\n"
             "       |\n"
             "       +-- HIT --> return cached (content, sources)   [~1ms]\n"
             "       |\n"
             "       +-- MISS -> Qdrant vector search              [~500ms]\n"
             "                       |\n"
             "                       v\n"
             "                  trim to top_k (no reranker)         [~1ms]\n"
             "                       |\n"
             "                       v\n"
             "                  KnowledgeCache.set(...)\n"
             "                       |\n"
             "                       v\n"
             "                  return (content, sources)"),
    ("p", "คำถามที่พบบ่อยจะ hit cache หลังจากถามครั้งแรก ตอบภายใน ~1ms "
          "แทนที่จะผ่าน Qdrant (~500ms) เก็บได้ 500 distinct queries, TTL 5 นาที "
          "hit rate เฉลี่ย 76%"),
]

# --- NVIDIA AI Blueprint references for CH5 phases ---
NVIDIA_PHASE_NOTES = {
    "Phase 1": (
        "Phase นี้ปรับเปลี่ยนจาก NVIDIA AI Blueprint for AI Virtual Assistant (NVIDIA, 2024) "
        "ซึ่ง blueprint เดิมใช้ NVIDIA NIM endpoints สำหรับ LLM inference "
        "โปรเจกต์นี้แทนที่ด้วย OpenRouter API และ Ollama สำหรับ local inference "
        "โดยยังคงโครงสร้าง FastAPI server ตาม blueprint แต่เพิ่ม "
        "OpenAPI documentation, hotel-specific endpoints และ PostgreSQL schema ใหม่ทั้งหมด"
    ),
    "Phase 2": (
        "Blueprint เดิมใช้ Milvus เป็น vector store และ NVIDIA embeddings "
        "โปรเจกต์นี้แทนที่ด้วย Qdrant (Docker-native single container) "
        "และ OpenRouter qwen3-embedding-8b (4096 dims) สำหรับ bilingual embeddings "
        "ฐานความรู้ถูกสร้างใหม่ทั้งหมดด้วย 10 hotel-specific markdown documents "
        "แทนที่ retail/e-commerce data ของ blueprint เดิม"
    ),
    "Phase 3": (
        "Blueprint เดิมใช้ LangGraph กับ sub-agents สำหรับ ProductQA, OrderStatus "
        "และ ReturnProcessing โปรเจกต์นี้คง LangGraph state machine architecture "
        "แต่เปลี่ยน sub-agents เป็น hotel-specific: booking, service, knowledge, other_talk "
        "พร้อม hotel tools 15 ตัวที่เชื่อมต่อกับ PostgreSQL "
        "แทนที่ retail tools ของ blueprint เดิม"
    ),
    "Phase 4": (
        "Blueprint เดิมออกแบบสำหรับ cloud deployment ด้วย NVIDIA GPU Cloud "
        "โปรเจกต์นี้สร้าง Docker Compose stack ใหม่สำหรับ local deployment "
        "ด้วย Ollama (local GPU inference) พร้อม runtime model switching "
        "ที่ไม่มีใน blueprint เดิม ทำให้โรงแรมสลับระหว่าง local/cloud ได้ทันที"
    ),
    "Phase 5": (
        "ระบบ authentication, rate limiting, account lockout และ token management "
        "เป็นส่วนที่โปรเจกต์นี้พัฒนาเพิ่มทั้งหมด — blueprint เดิมไม่มีระบบ auth "
        "ทุก endpoint เป็น public"
    ),
    "Phase 6": (
        "Audit logging จำนวน 26 action types และ DB connection pool "
        "เป็นส่วนที่พัฒนาเพิ่มทั้งหมดใน production hardening — "
        "blueprint เดิมไม่มี audit trail หรือ database scaling"
    ),
    "Phase 7": (
        "Chat scaling primitives ทั้ง 5 ตัว (LLM semaphore, session locks, "
        "chat rate limiter, stream cap, knowledge cache) เป็นนวัตกรรมของโปรเจกต์นี้ "
        "— blueprint เดิมออกแบบสำหรับ single-user demo ไม่มี concurrent user handling"
    ),
}


# ============================================================================
# MAIN
# ============================================================================
def main():
    print("=" * 60)
    print("Patch v6: font fix + WORKFLOW diagrams + NVIDIA references")
    print("=" * 60)

    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    body = doc.element.body

    # --- 1. Aggressive font fix ---
    print("\n1. Aggressive font fix (removing ALL rPr on TU_ styled runs)...")
    fixed = fix_fonts_aggressive(doc)
    print(f"  Fixed {fixed} runs")

    # --- 2. Find chapters ---
    chapters = find_chapter_indices(doc)
    for ch, idx in sorted(chapters.items(), key=lambda x: str(x[0])):
        print(f"  {ch}: P{idx}")

    ch4_start = chapters.get(4, 0)
    ch5_start = chapters.get(5, 0)
    ch6_start = chapters.get(6, len(doc.paragraphs))

    # --- 3. Add NVIDIA Blueprint phase notes to CH5 ---
    print("\n2. Adding NVIDIA Blueprint references to CH5 phases...")
    added_nvidia = 0
    for phase_key, note_text in NVIDIA_PHASE_NOTES.items():
        # Find "Phase X:" or "5.2.X Phase X:" heading
        anchor = find_para(doc, phase_key, ch5_start, ch6_start)
        if anchor is not None:
            # Find the end of this phase section (next heading or significant paragraph)
            # Insert the note as a new paragraph after the first paragraph following the heading
            insert_idx = anchor + 2  # after heading + first paragraph
            if insert_idx < len(doc.paragraphs):
                ref_elem = doc.paragraphs[insert_idx]._element
                note_node = make_para(doc, note_text, "TU_Paragraph_Normal")
                parent = ref_elem.getparent()
                idx = list(parent).index(ref_elem)
                parent.insert(idx, note_node)
                added_nvidia += 1
                # Re-scan since indices shifted
                chapters = find_chapter_indices(doc)
                ch5_start = chapters.get(5, 0)
                ch6_start = chapters.get(6, len(doc.paragraphs))
    print(f"  Added {added_nvidia} NVIDIA Blueprint notes")

    # --- 4. Add WORKFLOW diagrams to CH4 ---
    print("\n3. Adding WORKFLOW diagrams to CH4...")
    chapters = find_chapter_indices(doc)
    ch4_start = chapters.get(4, 0)
    ch5_start = chapters.get(5, 0)
    # Insert before CH5 heading (at end of CH4)
    ch5_elem = doc.paragraphs[ch5_start]._element
    nodes = build_nodes(doc, WORKFLOW_CH4)
    insert_before_element(body, ch5_elem, nodes)
    print(f"  Inserted {len(nodes)} CH4 diagram paragraphs")

    # --- 5. Add WORKFLOW diagrams to CH5 (end of chapter) ---
    print("\n4. Adding WORKFLOW diagrams to CH5...")
    chapters = find_chapter_indices(doc)
    ch6_start = chapters.get(6, len(doc.paragraphs))
    ch6_elem = doc.paragraphs[ch6_start]._element
    nodes = build_nodes(doc, WORKFLOW_CH5)
    insert_before_element(body, ch6_elem, nodes)
    print(f"  Inserted {len(nodes)} CH5 workflow paragraphs")

    # --- 6. Add NVIDIA Blueprint reference to bibliography ---
    print("\n5. Adding NVIDIA Blueprint reference...")
    chapters = find_chapter_indices(doc)
    ref_idx = chapters.get("references")
    if ref_idx:
        # Find "LLM and Agent Frameworks" in references
        anchor = find_para(doc, "LLM and Agent Frameworks", ref_idx, ref_idx + 50)
        if anchor:
            ref_elem = doc.paragraphs[anchor]._element
            nvidia_ref = make_para(doc,
                "NVIDIA. (2024). AI Blueprint for AI Virtual Assistant. "
                "NVIDIA Developer. https://developer.nvidia.com/ai-blueprints",
                "TU_Paragraph_Normal")
            parent = ref_elem.getparent()
            idx = list(parent).index(ref_elem)
            parent.insert(idx, nvidia_ref)
            print("  Added NVIDIA reference")

    # --- 7. Save ---
    doc.save(str(OUT))
    sz = os.path.getsize(str(OUT)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUT} ({sz:.0f} KB)")
    print("=" * 60)


if __name__ == "__main__":
    main()
