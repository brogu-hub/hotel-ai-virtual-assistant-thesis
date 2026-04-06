#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""
Inject thesis content into TULIBS .docx template at correct positions.

Instead of appending to the end, this script:
1. Finds each Heading 1 / TU_Chapter in the template
2. Removes the placeholder content between headings
3. Inserts real thesis content at the correct position using XML manipulation

Usage:
    python scripts/inject_into_template.py
"""
import os
import sys
import re
import copy
from pathlib import Path
from typing import Dict, List, Optional
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

THESIS_DIR = Path(__file__).parent.parent / "thesis"
TEMPLATE_PATH = THESIS_DIR / "TULIBS_Thesis-template-Thai_rev_2024.docx"
OUTPUT_PATH = THESIS_DIR / "thesis_final.docx"

CHAPTER_FILES = [
    "CH1_Introduction.md",
    "CH2_Literature_Review.md",
    "CH3_Methodology.md",
    "CH4_System_Design.md",
    "CH5_Implementation.md",
    "CH6_Testing_and_Evaluation.md",
    "CH7_Discussion.md",
    "CH8_Conclusion.md",
]
APPENDIX_FILES = ["AP_A_Test_Results.md", "AP_B_User_Manual.md"]
REFERENCES_FILE = "REFERENCES.md"

ABSTRACT_EN = """This thesis presents the design, implementation, and evaluation of an AI-powered virtual assistant for The Grand Horizon Hotel, a luxury 5-star hotel in Thailand. The system employs a LangGraph-based multi-agent architecture with four specialized sub-agents for booking operations, hotel services, knowledge retrieval, and general conversation, all supporting bilingual Thai and English interaction.

The backend is built with FastAPI and integrates Retrieval-Augmented Generation (RAG) using Qdrant vector embeddings over a hotel knowledge base of 10 documents. The system supports runtime switching between a local 9-billion-parameter LLM (Qwen3.5 Opus 9B on Ollama) and a cloud-hosted model (Qwen3 Max on OpenRouter).

Production-grade features include JWT authentication with role-based access control, login rate limiting, account lockout, audit logging, PII redaction, and five chat scaling primitives. The frontend is built with Next.js 15 and Ant Design 5, featuring a real-time chat interface with SSE streaming and an admin dashboard.

Evaluation across 25 hotel-domain test cases shows the local 9B model achieves 92% accuracy (23/25) compared to 100% (25/25) for the cloud model. Infrastructure testing covers 193 automated assertions, all passing. Performance optimization reduced warm chat latency from 18 seconds to 5 seconds."""

KEYWORDS_EN = "hotel AI, virtual assistant, LangGraph, multi-agent system, RAG, Qdrant, LLM, Ollama, FastAPI, Next.js, bilingual chatbot"
KEYWORDS_TH = "ปัญญาประดิษฐ์โรงแรม, ผู้ช่วยเสมือน, LangGraph, ระบบหลายตัวแทน, RAG, Qdrant, LLM, Ollama, FastAPI, Next.js, แชทบอทสองภาษา"

# Chapter titles for Heading 1 replacement
CHAPTER_TITLES = {
    1: "Introduction",
    2: "Literature Review",
    3: "Methodology",
    4: "System Design",
    5: "Implementation",
    6: "Testing and Evaluation",
    7: "Discussion",
    8: "Conclusion",
}


# =============================================================================
# XML-level paragraph insertion
# =============================================================================


def make_paragraph_element(doc, text, style_name, bold=False, italic=False,
                           font_name=None, font_size=None, color=None,
                           alignment=None, left_indent=None,
                           space_before=None, space_after=None):
    """Create a new w:p element with the given style and text."""
    # Find style ID from name
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name

    p = OxmlElement('w:p')

    # Paragraph properties
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)

    if alignment is not None:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)

    if left_indent is not None:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(left_indent * 567)))  # cm to twips
        pPr.append(ind)

    if space_before is not None or space_after is not None:
        spacing = OxmlElement('w:spacing')
        if space_before is not None:
            spacing.set(qn('w:before'), str(int(space_before * 20)))
        if space_after is not None:
            spacing.set(qn('w:after'), str(int(space_after * 20)))
        pPr.append(spacing)

    p.append(pPr)

    # Run with text
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)
        if font_name:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)
            rPr.append(rFonts)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size * 2)))  # half-points
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(font_size * 2)))
            rPr.append(szCs)
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)

    return p


def make_rich_paragraph(doc, text, style_name):
    """Create a paragraph with **bold**, *italic*, `code` runs."""
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)

    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            b = OxmlElement('w:b')
            rPr.append(b)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            part = part[1:-1]
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)
        elif part.startswith("`") and part.endswith("`"):
            part = part[1:-1]
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Consolas')
            rFonts.set(qn('w:hAnsi'), 'Consolas')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')  # 12pt
            rPr.append(sz)

        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = part
        r.append(t)
        p.append(r)

    return p


def make_table_element(doc, rows):
    """Create a w:tbl element from row data."""
    if not rows:
        return None
    max_cols = max(len(r) for r in rows)

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')
    tblPr.append(tblW)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    for ri, row in enumerate(rows):
        tr = OxmlElement('w:tr')
        for ci in range(max_cols):
            tc = OxmlElement('w:tc')
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if ri == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t_elem = OxmlElement('w:t')
            t_elem.set(qn('xml:space'), 'preserve')
            t_elem.text = row[ci] if ci < len(row) else ""
            r.append(t_elem)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


# =============================================================================
# Markdown parser (same as before)
# =============================================================================


def parse_markdown(md_text, chapter_num):
    elements = []
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_lines = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                elements.append({"type": "code", "content": "\n".join(code_lines), "lang": code_lang})
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                in_code = True
                code_lang = stripped[3:].strip()
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("$$"):
            latex = stripped[2:]
            if stripped.endswith("$$") and len(stripped) > 4:
                elements.append({"type": "latex", "content": latex[:-2]})
            else:
                eq = [latex]
                i += 1
                while i < len(lines):
                    if lines[i].strip().endswith("$$"):
                        eq.append(lines[i].strip()[:-2])
                        break
                    eq.append(lines[i].strip())
                    i += 1
                elements.append({"type": "latex", "content": " ".join(eq)})
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip().replace("**", "")
            elements.append({"type": "heading", "level": level, "content": text, "chapter": chapter_num})
            i += 1
            continue

        if stripped.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r'^\|[\s\-:|]+\|$', row):
                    cells = [c.strip().replace("**", "") for c in row.split("|")[1:-1]]
                    table_rows.append(cells)
                i += 1
            if table_rows:
                elements.append({"type": "table", "rows": table_rows})
            continue

        if stripped.startswith("[Figure"):
            elements.append({"type": "figure", "content": stripped})
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            elements.append({"type": "bullet", "content": stripped[2:]})
            i += 1
            continue
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            elements.append({"type": "numbered", "content": m.group(2)})
            i += 1
            continue
        if not stripped:
            i += 1
            continue

        elements.append({"type": "paragraph", "content": stripped})
        i += 1
    return elements


# =============================================================================
# Get style name for main heading per chapter
# =============================================================================

def get_main_heading_style(ch):
    if ch == 1:
        return "TU_Main Heading _Chapter1"
    return f"TU_Main Heading_Chapter{ch}"


# =============================================================================
# Convert elements to XML nodes and insert after a reference element
# =============================================================================

def elements_to_xml(doc, elements):
    """Convert parsed markdown elements to a list of XML nodes (w:p or w:tbl)."""
    nodes = []
    for elem in elements:
        et = elem["type"]
        ch = elem.get("chapter", 1)

        if et == "heading":
            level = elem["level"]
            if level == 1:
                style = "Heading 1"
            elif level == 2:
                style = get_main_heading_style(ch)
            elif level == 3:
                style = "TU_Sub-heading 1"
            elif level == 4:
                style = "TU_Sub-heading 2"
            else:
                style = "TU_Sub-heading 3"
            nodes.append(make_paragraph_element(doc, elem["content"], style, bold=True))

        elif et == "paragraph":
            nodes.append(make_rich_paragraph(doc, elem["content"], "TU_Paragraph_Normal"))

        elif et == "bullet":
            nodes.append(make_rich_paragraph(doc, "• " + elem["content"], "TU_Paragraph_Normal"))

        elif et == "numbered":
            nodes.append(make_rich_paragraph(doc, elem["content"], "TU_Paragraph_Normal"))

        elif et == "code":
            content = elem["content"]
            code_lines = content.split("\n")
            # Source ref
            if code_lines and code_lines[0].strip().startswith("#"):
                fm = re.match(r'^#\s*(src/\S+|deploy/\S+)', code_lines[0].strip())
                if fm:
                    nodes.append(make_paragraph_element(
                        doc, f"[Source: {fm.group(1)}]", "TU_Paragraph_Normal",
                        italic=True, font_size=10, color="666666"))
                    code_lines = code_lines[1:]
            for cl in code_lines:
                nodes.append(make_paragraph_element(
                    doc, cl, "TU_Paragraph_Normal",
                    font_name="Consolas", font_size=10, left_indent=1,
                    space_before=0, space_after=0))

        elif et == "table":
            tbl = make_table_element(doc, elem["rows"])
            if tbl is not None:
                nodes.append(tbl)

        elif et == "figure":
            # Placeholder with red marker
            p = make_paragraph_element(
                doc, elem["content"] + "\n[Insert diagram here / แทรกภาพที่นี่]",
                "TU_Paragraph_Normal", italic=True, font_size=11,
                alignment="center", color="CC0000")
            nodes.append(p)

        elif et == "latex":
            latex = elem["content"]
            um = {r'\kappa': 'κ', r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
                  r'\sigma': 'σ', r'\mu': 'μ', r'\pi': 'π', r'\theta': 'θ',
                  r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\times': '×',
                  r'\cdot': '·', r'\infty': '∞'}
            rendered = latex
            for tex, uni in um.items():
                rendered = rendered.replace(tex, uni)
            rendered = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1) / (\2)', rendered)
            rendered = re.sub(r'\\(\w+)', r'\1', rendered)
            rendered = rendered.replace('{', '').replace('}', '')
            nodes.append(make_paragraph_element(
                doc, rendered, "TU_Paragraph_Normal", italic=True,
                font_name="Cambria Math", font_size=14, alignment="center"))

    return nodes


def insert_nodes_after(body, ref_element, nodes):
    """Insert a list of XML nodes right after ref_element in the document body."""
    parent = ref_element.getparent()
    idx = list(parent).index(ref_element)
    for j, node in enumerate(nodes):
        parent.insert(idx + 1 + j, node)


def remove_placeholder_paragraphs(doc, start_para_idx, end_para_idx):
    """Remove placeholder paragraphs between two indices."""
    body = doc.element.body
    placeholders = ("เริ่มพิมพ์เนื้อหา", "เริ่มพิมพ์ย่อหน้าใหม่", "Insert text here",
                    "หัวข้อใหญ่", "1.1.1 หัวข้อย่อยระดับที่ 1", "1.2.1 หัวข้อย่อยระดับที่ 1",
                    "หัวข้อย่อยระดับที่ 1", "หัวข้อย่อยระดับที่ 2", "หัวข้อย่อยระดับที่ 3",
                    "(1) หัวข้อย่อยระดับที่ 3")

    to_remove = []
    for i in range(start_para_idx, min(end_para_idx, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = p.text.strip()
        style = p.style.name if p.style else ""
        # Remove if it's placeholder text or a template sub-heading/main-heading with placeholder content
        if text in placeholders or (
            any(text.endswith(ph) for ph in placeholders) and
            style.startswith("TU_")
        ):
            to_remove.append(p._element)

    for elem in to_remove:
        elem.getparent().remove(elem)

    return len(to_remove)


# =============================================================================
# Main
# =============================================================================

def main():
    print("=" * 60)
    print("Thesis Template Injector v2 — Insert at correct positions")
    print("=" * 60)

    doc = Document(str(TEMPLATE_PATH))
    body = doc.element.body
    print(f"Template: {len(doc.paragraphs)} paragraphs")

    # Find all Heading 1 positions (chapter boundaries)
    heading1_indices = []
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1":
            heading1_indices.append(i)
            print(f"  Heading 1 at P{i}: {p.text.strip()[:50]}")

    print(f"\n  Found {len(heading1_indices)} chapter slots in template")

    # Fill abstracts + keywords
    print("\nFilling abstract + keywords...")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "Insert text here":
            for j in range(max(0, i - 5), i):
                if "ABSTRACT" in doc.paragraphs[j].text:
                    for run in p.runs:
                        run.text = ""
                    p.add_run(ABSTRACT_EN)
                    break
        if text.startswith("Keywords: Insert"):
            for run in p.runs:
                run.text = ""
            p.add_run(f"Keywords: {KEYWORDS_EN}")
        if text.startswith("คำสำคัญ: พิมพ์"):
            for run in p.runs:
                run.text = ""
            p.add_run(f"คำสำคัญ: {KEYWORDS_TH}")

    # Process chapters in REVERSE order so insertions don't shift
    # the paragraph indices of earlier chapters
    print("\nInjecting chapters (reverse order to preserve indices)...")
    for ch_idx in reversed(range(len(CHAPTER_FILES))):
        ch_file = CHAPTER_FILES[ch_idx]
        ch_num = ch_idx + 1
        ch_path = THESIS_DIR / ch_file
        if not ch_path.exists():
            print(f"  [SKIP] {ch_file}")
            continue

        md_text = ch_path.read_text(encoding="utf-8")
        elements = parse_markdown(md_text, ch_num)

        # Skip the chapter title heading (template already has Heading 1)
        if elements and elements[0]["type"] == "heading" and elements[0]["level"] == 1:
            elements = elements[1:]

        if ch_idx >= len(heading1_indices):
            print(f"  [WARN] No template slot for CH{ch_num}, appending")
            continue

        # Find the Heading 1 paragraph element in the template
        h1_para_idx = heading1_indices[ch_idx]
        h1_para = doc.paragraphs[h1_para_idx]

        # Update Heading 1 text
        title = CHAPTER_TITLES.get(ch_num, ch_file)
        for run in h1_para.runs:
            run.text = ""
        h1_para.add_run(title)

        # Remove placeholder content between this Heading 1 and the next
        next_boundary = heading1_indices[ch_idx + 1] if ch_idx + 1 < len(heading1_indices) else len(doc.paragraphs)
        removed = remove_placeholder_paragraphs(doc, h1_para_idx + 1, next_boundary)

        # Convert elements to XML nodes
        xml_nodes = elements_to_xml(doc, elements)

        # Insert after the Heading 1 element
        insert_nodes_after(body, h1_para._element, xml_nodes)

        print(f"  [CH{ch_num}] {ch_file}: {len(elements)} elements, {len(xml_nodes)} nodes inserted, {removed} placeholders removed")

    # Process appendices and references BEFORE chapters (they come later in doc,
    # so we do them first, then chapters in reverse — this way no insertion
    # shifts the indices of any later insertion target)

    # References
    print("\nInjecting references...")
    ref_path = THESIS_DIR / REFERENCES_FILE
    if ref_path.exists():
        md_text = ref_path.read_text(encoding="utf-8")
        elements = parse_markdown(md_text, 0)
        if elements and elements[0]["type"] == "heading":
            elements = elements[1:]
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == "TU_Chapter" and "รายการอ้างอิง" in p.text:
                xml_nodes = elements_to_xml(doc, elements)
                insert_nodes_after(body, p._element, xml_nodes)
                print(f"  [REF] {len(xml_nodes)} nodes inserted after รายการอ้างอิง")
                break

    # Appendices
    print("\nInjecting appendices...")
    appendix_labels = ["ภาคผนวก ก", "ภาคผนวก ข"]
    for ap_idx, ap_file in enumerate(APPENDIX_FILES):
        ap_path = THESIS_DIR / ap_file
        if not ap_path.exists():
            continue
        md_text = ap_path.read_text(encoding="utf-8")
        elements = parse_markdown(md_text, 0)
        if elements and elements[0]["type"] == "heading":
            elements = elements[1:]

        label = appendix_labels[ap_idx] if ap_idx < len(appendix_labels) else None
        if label:
            for i, p in enumerate(doc.paragraphs):
                if p.style.name == "TU_Chapter" and label in p.text:
                    # Find the "ชื่อภาคผนวก" after it and insert after that
                    for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                        if "ชื่อภาคผนวก" in doc.paragraphs[j].text:
                            # Update appendix title
                            ap_title = "Test Results and Development Timeline" if ap_idx == 0 else "User Manual"
                            for run in doc.paragraphs[j].runs:
                                run.text = ""
                            doc.paragraphs[j].add_run(ap_title)
                            # Insert content
                            xml_nodes = elements_to_xml(doc, elements)
                            insert_nodes_after(body, doc.paragraphs[j]._element, xml_nodes)
                            print(f"  [AP{ap_idx + 1}] {ap_file}: {len(xml_nodes)} nodes inserted")
                            break
                    break

    # Save
    doc.save(str(OUTPUT_PATH))
    sz = os.path.getsize(str(OUTPUT_PATH)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUTPUT_PATH} ({sz:.0f} KB)")
    print("=" * 60)
    print("\nIn Word:")
    print("  1. Right-click Table of Contents → Update Field → Update Entire Table")
    print("  2. Ctrl+F → 'แทรกภาพที่นี่' to find figure placeholders")
    print("  3. Fill cover page: university, author, advisor")
    print("  4. Delete any remaining template placeholder text")


if __name__ == "__main__":
    main()
