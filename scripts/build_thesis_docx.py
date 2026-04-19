#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""
Build thesis .docx — combines all chapters, translates to Thai via local
Ollama, and generates a Word document with proper formatting.

Features:
  - Combines all thesis/*.md chapters in order
  - Extracts Table of Contents, List of Figures, List of Tables
  - Translates each section to Thai via Ollama qwen3.5-opus:9b
  - Generates .docx with heading styles, tables, code blocks
  - Preserves English technical terms (LangGraph, FastAPI, JWT, etc.)

Usage:
    python scripts/build_thesis_docx.py
    python scripts/build_thesis_docx.py --skip-translation   # English only
    python scripts/build_thesis_docx.py --model qwen/qwen3-max --api openrouter  # use cloud model
"""
import os
import sys
import re
import json
import argparse
import time
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

# =============================================================================
# Configuration
# =============================================================================

THESIS_DIR = Path(__file__).parent.parent / "thesis"
OUTPUT_DIR = Path(__file__).parent.parent / "thesis"

CHAPTER_ORDER = [
    "CH1_Introduction.md",
    "CH2_Literature_Review.md",
    "CH3_Methodology.md",
    "CH4_System_Design.md",
    "CH5_Implementation.md",
    "CH6_Testing_and_Evaluation.md",
    "CH7_Discussion.md",
    "CH8_Conclusion.md",
    "AP_A_Test_Results.md",
    "AP_B_User_Manual.md",
    "REFERENCES.md",
]

OLLAMA_URL = "http://localhost:11435"
OLLAMA_MODEL = "fredrezones55/qwen3.5-opus:9b"

ABSTRACT_EN = """
This thesis presents the design, implementation, and evaluation of an AI-powered
virtual assistant for The Grand Horizon Hotel, a luxury 5-star hotel in Thailand.
The system employs a LangGraph-based multi-agent architecture with four specialized
sub-agents for booking operations, hotel services, knowledge retrieval, and general
conversation, all supporting bilingual Thai and English interaction.

The backend is built with FastAPI and integrates Retrieval-Augmented Generation (RAG)
using Qdrant vector embeddings over a hotel knowledge base of 10 documents. The system
supports runtime switching between a local 9-billion-parameter LLM (Qwen3.5 Opus 9B
on Ollama) and a cloud-hosted model (Qwen3 Max on OpenRouter), enabling hotels to
balance cost, privacy, and capability based on traffic demands.

Production-grade features include JWT authentication with role-based access control,
login rate limiting, account lockout, audit logging of all admin actions, PII redaction,
and five chat scaling primitives (LLM concurrency semaphore, per-session locks, knowledge
cache, chat rate limiter, and SSE stream cap). The frontend is built with Next.js 15
and Ant Design 5, featuring a real-time chat interface with SSE streaming, an admin
dashboard for session monitoring and intervention, and runtime LLM model switching.

Evaluation across 25 hotel-domain test cases shows the local 9B model achieves 92%
accuracy (23/25) compared to 100% (25/25) for the cloud model. Both models achieve
100% accuracy on knowledge retrieval (8/8) and booking operations (6/6). Infrastructure
testing covers 193 automated assertions across authentication, security hardening,
audit logging, and concurrent-user scaling, all passing. Performance optimization
reduced warm chat latency from 18 seconds to 5 seconds through reranker removal,
prompt trimming, and GPU parallelism tuning.

Keywords: hotel AI, virtual assistant, LangGraph, multi-agent system, RAG, Qdrant,
LLM, Ollama, FastAPI, Next.js, bilingual chatbot
""".strip()

UNIVERSITY_PLACEHOLDER = "[University Name / ชื่อมหาวิทยาลัย]"
DEPARTMENT_PLACEHOLDER = "[Department of Computer Science / ภาควิชาวิทยาการคอมพิวเตอร์]"
AUTHOR_PLACEHOLDER = "[Student Name / ชื่อนักศึกษา]"
ADVISOR_PLACEHOLDER = "[Advisor Name / ชื่ออาจารย์ที่ปรึกษา]"
YEAR_PLACEHOLDER = "2026"


# =============================================================================
# Translation via Ollama
# =============================================================================


def translate_to_thai(text: str, model: str = OLLAMA_MODEL, api: str = "ollama") -> str:
    """Translate English text to Thai using the local Ollama model."""
    if not text.strip():
        return text

    prompt = f"""Translate the following English academic thesis text to Thai.

RULES:
1. Keep ALL technical terms in English: LangGraph, FastAPI, JWT, RAG, Qdrant, Ollama, OpenRouter, Docker, PostgreSQL, Redis, Next.js, React, TypeScript, Ant Design, Zustand, SWR, API, SSE, PII, RBAC, LLM, GPU, VRAM, etc.
2. Keep ALL code blocks, commands, URLs, file paths, and variable names exactly as-is
3. Keep ALL table data (numbers, percentages, model names) as-is — only translate the header labels
4. Keep ALL figure/table references like [Figure X.Y: ...] — translate only the description part
5. Use formal academic Thai (ภาษาเขียน not ภาษาพูด)
6. Use ครับ/ค่ะ-free academic style (no polite particles — this is a thesis, not a conversation)
7. Translate section headings but keep the numbering (e.g., "## 3.1 Research Approach" → "## 3.1 แนวทางการวิจัย")
8. For markdown tables: translate header text but keep the markdown pipe syntax intact
9. DO NOT add any commentary or notes — output ONLY the translated text

TEXT TO TRANSLATE:
{text}"""

    if api == "ollama":
        try:
            r = requests.post(
                f"{OLLAMA_URL}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,
                        "num_predict": 4096,
                    },
                },
                timeout=300,
            )
            if r.status_code == 200:
                result = r.json().get("response", "")
                # Strip any <think>...</think> tags from Qwen3.5
                result = re.sub(r'<think>.*?</think>', '', result, flags=re.DOTALL).strip()
                return result
            else:
                print(f"  [WARN] Ollama returned {r.status_code}")
                return text
        except Exception as e:
            print(f"  [WARN] Translation failed: {e}")
            return text
    elif api == "openrouter":
        api_key = os.getenv("OPENROUTER_API_KEY", "")
        try:
            r = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "max_tokens": 4096,
                },
                timeout=120,
            )
            if r.status_code == 200:
                result = r.json()["choices"][0]["message"]["content"]
                result = re.sub(r'<think>.*?</think>', '', result, flags=re.DOTALL).strip()
                return result
            else:
                print(f"  [WARN] OpenRouter returned {r.status_code}: {r.text[:200]}")
                return text
        except Exception as e:
            print(f"  [WARN] Translation failed: {e}")
            return text
    return text


# =============================================================================
# Markdown parsing
# =============================================================================


def split_into_sections(md_text: str) -> List[Tuple[str, str]]:
    """Split markdown into (heading, content) pairs for section-by-section translation."""
    sections = []
    current_heading = ""
    current_content = []

    for line in md_text.split("\n"):
        if line.startswith("#"):
            if current_heading or current_content:
                sections.append((current_heading, "\n".join(current_content)))
            current_heading = line
            current_content = []
        else:
            current_content.append(line)

    if current_heading or current_content:
        sections.append((current_heading, "\n".join(current_content)))

    return sections


def extract_figures_tables(all_text: str) -> Tuple[List[str], List[str]]:
    """Extract all [Figure X.Y: ...] and table captions for LoF and LoT."""
    figures = re.findall(r'\[Figure \d+\.\d+:.*?\]', all_text)
    # Tables: lines starting with | that have a header row
    tables = []
    table_count = 0
    for line in all_text.split("\n"):
        if line.strip().startswith("|") and "---" not in line:
            # Check if this looks like a table header (has | separators)
            if line.count("|") >= 3 and table_count < 50:
                # Only count unique table starts
                pass
    # Simpler: find all markdown table sections
    in_table = False
    for line in all_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("|") and not in_table:
            in_table = True
            table_count += 1
        elif not stripped.startswith("|") and in_table:
            in_table = False

    return figures, [f"Table {i+1}" for i in range(table_count)]


# =============================================================================
# DOCX generation
# =============================================================================


def setup_styles(doc: Document):
    """Configure document styles for Thai/English thesis."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cordia New'
    font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Cordia New')

    # Heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Cordia New'
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(20)
        elif i == 2:
            heading_style.font.size = Pt(18)
        else:
            heading_style.font.size = Pt(16)


def add_cover_page(doc: Document):
    """Add thesis cover page."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ระบบผู้ช่วยเสมือนอัจฉริยะสำหรับโรงแรม\n"
                     "โดยใช้ปัญญาประดิษฐ์แบบ Multi-Agent\n\n"
                     "Hotel AI Virtual Assistant\n"
                     "Using Multi-Agent Artificial Intelligence")
    run.font.size = Pt(20)
    run.bold = True
    run.font.name = 'Cordia New'

    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{AUTHOR_PLACEHOLDER}")
    run.font.size = Pt(16)
    run.font.name = 'Cordia New'

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรวิทยาศาสตรบัณฑิต\n"
        f"{DEPARTMENT_PLACEHOLDER}\n"
        f"{UNIVERSITY_PLACEHOLDER}\n"
        f"ปีการศึกษา {YEAR_PLACEHOLDER}"
    )
    run.font.size = Pt(16)
    run.font.name = 'Cordia New'

    doc.add_page_break()


def add_abstract_page(doc: Document, thai_abstract: str):
    """Add abstract pages (Thai + English)."""
    # Thai abstract
    doc.add_heading('บทคัดย่อ', level=1)
    p = doc.add_paragraph(thai_abstract)
    p.style.font.name = 'Cordia New'
    doc.add_page_break()

    # English abstract
    doc.add_heading('Abstract', level=1)
    p = doc.add_paragraph(ABSTRACT_EN)
    p.style.font.name = 'Cordia New'
    doc.add_page_break()


def add_toc_placeholder(doc: Document):
    """Add Table of Contents placeholder."""
    doc.add_heading('สารบัญ (Table of Contents)', level=1)
    doc.add_paragraph(
        "[สารบัญจะถูกสร้างอัตโนมัติใน Microsoft Word: "
        "References → Table of Contents → Automatic Table 1]\n\n"
        "Insert → Table of Contents in MS Word to auto-generate from Heading styles."
    )
    doc.add_page_break()


def add_list_of_figures(doc: Document, figures: List[str]):
    """Add List of Figures."""
    doc.add_heading('สารบัญภาพ (List of Figures)', level=1)
    if figures:
        for fig in figures:
            doc.add_paragraph(fig, style='List Bullet')
    else:
        doc.add_paragraph("[รายการภาพจะถูกสร้างอัตโนมัติใน MS Word]")
    doc.add_page_break()


def add_list_of_tables(doc: Document, table_count: int):
    """Add List of Tables."""
    doc.add_heading('สารบัญตาราง (List of Tables)', level=1)
    doc.add_paragraph(
        f"[เอกสารนี้มีตารางทั้งหมด {table_count} ตาราง — "
        f"สร้างรายการอัตโนมัติใน MS Word: References → Insert Table of Figures → Table]"
    )
    doc.add_page_break()


def add_rich_paragraph(doc: Document, text: str):
    """Add a paragraph with **bold**, *italic*, and `code` rendered as Word formatting."""
    p = doc.add_paragraph()
    # Split on bold, italic, and code patterns
    # Process in order: **bold** first, then *italic*, then `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(12)
        else:
            p.add_run(part)
    return p


def add_latex_equation(doc: Document, latex: str):
    """Add a LaTeX equation as a centered Word paragraph with Unicode rendering."""
    # Map common LaTeX symbols to Unicode
    unicode_map = {
        r'\kappa': 'κ', r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
        r'\delta': 'δ', r'\sigma': 'σ', r'\mu': 'μ', r'\lambda': 'λ',
        r'\pi': 'π', r'\theta': 'θ', r'\epsilon': 'ε', r'\phi': 'φ',
        r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\approx': '≈',
        r'\times': '×', r'\cdot': '·', r'\infty': '∞', r'\sum': 'Σ',
    }
    # Handle \frac{a}{b} → a/b or (a)/(b)
    rendered = latex.strip().strip('$')
    for tex, uni in unicode_map.items():
        rendered = rendered.replace(tex, uni)
    # Convert \frac{num}{den} to (num)/(den)
    rendered = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1) / (\2)', rendered)
    # Remove remaining backslashes from unknown commands
    rendered = re.sub(r'\\(\w+)', r'\1', rendered)
    # Clean up braces
    rendered = rendered.replace('{', '').replace('}', '')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(rendered)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(14)
    run.italic = True
    return p


def add_code_block(doc: Document, code_lines: List[str], lang: str = "",
                   source_file: str = "", source_lines: str = ""):
    """Add a code block as a bordered, monospace paragraph with source reference."""
    # Source file reference header
    if source_file:
        ref_text = f"📁 {source_file}"
        if source_lines:
            ref_text += f" (lines {source_lines})"
        p = doc.add_paragraph()
        run = p.add_run(ref_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True

    # Code content
    for line in code_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(1)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    # Blank line after code block
    doc.add_paragraph("")


def add_figure_placeholder(doc: Document, figure_text: str, ascii_art: str = ""):
    """Add a figure placeholder as a bordered centered box with optional ASCII art."""
    # Extract figure number and description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)

    if ascii_art:
        # Render ASCII diagram as monospace
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for art_line in ascii_art.split("\n"):
            run = p2.add_run(art_line + "\n")
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    # Figure caption
    run = p.add_run("━" * 40 + "\n")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run = p.add_run(figure_text)
    run.italic = True
    run.font.size = Pt(11)
    run = p.add_run("\n[แทรกภาพที่นี่ / Insert diagram here]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run = p.add_run("\n" + "━" * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def md_to_docx_section(doc: Document, md_text: str, workflow_diagrams: Optional[Dict] = None):
    """Convert a markdown section to docx paragraphs with rich formatting."""
    if workflow_diagrams is None:
        workflow_diagrams = {}

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_block_lines: List[str] = []
    code_block_lang = ""
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # LaTeX display equations: $$..$$
        if stripped.startswith("$$"):
            latex_content = stripped[2:]
            if stripped.endswith("$$") and len(stripped) > 4:
                # Single-line equation
                add_latex_equation(doc, latex_content[:-2])
            else:
                # Multi-line: collect until closing $$
                eq_lines = [latex_content]
                i += 1
                while i < len(lines):
                    if lines[i].strip().endswith("$$"):
                        eq_lines.append(lines[i].strip()[:-2])
                        break
                    eq_lines.append(lines[i].strip())
                    i += 1
                add_latex_equation(doc, " ".join(eq_lines))
            i += 1
            continue

        # Code blocks — collect and render with source reference
        if stripped.startswith("```"):
            if in_code_block:
                # Detect source file reference from the code or preceding comment
                source_file = ""
                source_lines_ref = ""
                # Check if the first line of code is a comment with file path
                if code_block_lines and code_block_lines[0].strip().startswith("#"):
                    comment = code_block_lines[0].strip()
                    # Match "# src/hotel_guardrails/auth.py" or similar
                    fm = re.match(r'^#\s*(src/\S+|deploy/\S+)', comment)
                    if fm:
                        source_file = fm.group(1)
                        code_block_lines = code_block_lines[1:]  # remove the comment

                add_code_block(doc, code_block_lines, code_block_lang,
                              source_file, source_lines_ref)
                in_code_block = False
                code_block_lines = []
                code_block_lang = ""
            else:
                in_code_block = True
                code_block_lang = stripped[3:].strip()  # e.g., "python", "yaml"
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Headings — strip any ** from heading text
        if stripped.startswith("# "):
            heading = stripped[2:].replace("**", "")
            doc.add_heading(heading, level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:].replace("**", "")
            doc.add_heading(heading, level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            heading = stripped[4:].replace("**", "")
            doc.add_heading(heading, level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            heading = stripped[5:].replace("**", "")
            doc.add_heading(heading, level=4)
            i += 1
            continue

        # Tables — collect rows then add as Word table
        if stripped.startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # Clean ** from cell content
            cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
            table_rows.append(cells)
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith("|"):
                if table_rows:
                    try:
                        max_cols = max(len(r) for r in table_rows)
                        tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                        tbl.style = 'Table Grid'
                        for ri, row in enumerate(table_rows):
                            for ci, cell in enumerate(row):
                                if ci < max_cols:
                                    tbl.rows[ri].cells[ci].text = cell
                                    if ri == 0:
                                        for paragraph in tbl.rows[ri].cells[ci].paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
                    except Exception as e:
                        doc.add_paragraph(f"[Table rendering error: {e}]")
                in_table = False
                table_rows = []
            i += 1
            continue

        # Figure placeholders — with optional ASCII art from WORKFLOW.md
        if stripped.startswith("[Figure"):
            # Check if there's a matching ASCII diagram
            fig_match = re.match(r'\[Figure (\d+\.\d+)', stripped)
            ascii_art = ""
            if fig_match:
                fig_id = fig_match.group(1)
                ascii_art = workflow_diagrams.get(fig_id, "")
            add_figure_placeholder(doc, stripped, ascii_art)
            i += 1
            continue

        # Horizontal rules
        if stripped == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # Bullet points — with rich text
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            add_rich_paragraph(doc, content).style = doc.styles['List Bullet']
            i += 1
            continue

        # Numbered items
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            add_rich_paragraph(doc, m.group(2)).style = doc.styles['List Number']
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph — with rich formatting (bold, italic, code)
        add_rich_paragraph(doc, stripped)
        i += 1


# =============================================================================
# Main
# =============================================================================


def main():
    parser = argparse.ArgumentParser(description="Build thesis .docx with Thai translation")
    parser.add_argument("--skip-translation", action="store_true", help="Skip Thai translation (English only)")
    parser.add_argument("--model", default=OLLAMA_MODEL, help="Model for translation")
    parser.add_argument("--api", default="ollama", choices=["ollama", "openrouter"], help="API backend")
    parser.add_argument("--output", default=str(OUTPUT_DIR / "thesis_combined_thai.docx"), help="Output path")
    args = parser.parse_args()

    print("=" * 60)
    print("Hotel AI Virtual Assistant — Thesis Builder")
    print("=" * 60)
    print(f"Translation: {'SKIP' if args.skip_translation else f'{args.model} via {args.api}'}")
    print(f"Output: {args.output}")
    print()

    # Step 1: Read all chapters
    print("Reading chapters...")
    all_sections = []  # (chapter_file, heading, content)
    all_text = ""
    for fname in CHAPTER_ORDER:
        fpath = THESIS_DIR / fname
        if not fpath.exists():
            print(f"  [SKIP] {fname} not found")
            continue
        text = fpath.read_text(encoding="utf-8")
        all_text += text + "\n\n"
        sections = split_into_sections(text)
        for heading, content in sections:
            all_sections.append((fname, heading, content))
        print(f"  [OK] {fname} ({len(sections)} sections)")

    # Step 2: Extract figures and tables for front matter
    figures, tables = extract_figures_tables(all_text)
    table_count = all_text.count("| ") // 3  # rough estimate
    print(f"\nFound {len(figures)} figure references, ~{table_count} tables")

    # Step 2b: Load ASCII diagrams from WORKFLOW.md for figure placeholders
    workflow_diagrams: Dict[str, str] = {}
    workflow_path = THESIS_DIR.parent / "docs" / "WORKFLOW.md"
    if workflow_path.exists():
        wf_text = workflow_path.read_text(encoding="utf-8")
        # Extract all ``` code blocks (ASCII diagrams)
        code_blocks = re.findall(r'```\n(.*?)```', wf_text, re.DOTALL)
        # Map figure IDs to the nearest preceding code block
        # For simplicity, store all code blocks and let figure placeholders pick them up
        for idx_b, block in enumerate(code_blocks):
            if any(c in block for c in ['│', '├', '→', '┌', '└', '+--', '|']):
                workflow_diagrams[f"workflow_{idx_b}"] = block.strip()
        print(f"  Loaded {len(workflow_diagrams)} ASCII diagrams from WORKFLOW.md")

    # Step 3: Translate abstract
    thai_abstract = ""
    if not args.skip_translation:
        print("\nTranslating abstract to Thai...")
        start = time.time()
        thai_abstract = translate_to_thai(ABSTRACT_EN, model=args.model, api=args.api)
        print(f"  Done ({time.time()-start:.1f}s, {len(thai_abstract)} chars)")
    else:
        thai_abstract = "[บทคัดย่อภาษาไทยจะถูกแปลโดย LLM — ใช้ --skip-translation=false เพื่อแปล]"

    # Step 4: Build .docx
    print("\nBuilding .docx...")
    doc = Document()
    setup_styles(doc)

    # Front matter
    add_cover_page(doc)
    add_abstract_page(doc, thai_abstract)
    add_toc_placeholder(doc)
    add_list_of_figures(doc, figures)
    add_list_of_tables(doc, table_count)

    # Step 5: Process chapters — translate if requested, then add to docx
    total_sections = len(all_sections)
    for idx, (fname, heading, content) in enumerate(all_sections):
        section_text = heading + "\n" + content if heading else content
        section_text = section_text.strip()
        if not section_text:
            continue

        label = heading[:60] if heading else fname
        print(f"  [{idx+1}/{total_sections}] {label}", end="", flush=True)

        if not args.skip_translation and len(section_text) > 50:
            start = time.time()
            translated = translate_to_thai(section_text, model=args.model, api=args.api)
            elapsed = time.time() - start
            print(f" → translated ({elapsed:.1f}s)")
            md_to_docx_section(doc, translated, workflow_diagrams)
        else:
            print(" → English")
            md_to_docx_section(doc, section_text, workflow_diagrams)

    # Step 6: Save
    os.makedirs(os.path.dirname(args.output) or ".", exist_ok=True)
    doc.save(args.output)
    file_size = os.path.getsize(args.output) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {args.output} ({file_size:.0f} KB)")
    print(f"Total sections translated: {total_sections}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
