#!/usr/bin/env python3
"""
patch_v8: prompts in 5.3 + PNG figures + page breaks
=====================================================
1. Expand 5.3 Prompt Engineering with all prompts and changes
2. Insert PNG figures at diagram sections (5.2.9, 5.8, 5.10, etc.)
3. Add page breaks before every chapter heading
"""
import os, sys, re, shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SRC = Path("thesis/thesis_final_v7.docx")
OUT = Path("thesis/thesis_final_v8.docx")
FIGURES_DIR = Path("thesis/figures")


# ============================================================================
# Helpers
# ============================================================================
def make_para(doc, text, style_name, alignment=None):
    style_id = None
    for s in doc.styles:
        if s.name == style_name:
            style_id = s.style_id
            break
    if style_id is None:
        style_id = style_name
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p


def make_code_para_styled(doc, text, tokens_func):
    """Create a code paragraph with One Dark Pro styling."""
    C_BG = "282C34"
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'TUParagraphNormal')
    pPr.append(pStyle)
    # Background
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), C_BG)
    pPr.append(shd)
    # Left accent
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '528BFF')
    pBdr.append(left)
    pPr.append(pBdr)
    # Spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    # Indent
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '284')
    ind.set(qn('w:right'), '284')
    pPr.append(ind)
    p.append(pPr)
    # Single run with Consolas + default color
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Consolas')
    rf.set(qn('w:hAnsi'), 'Consolas')
    rf.set(qn('w:cs'), 'Consolas')
    rPr.append(rf)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '18')
    rPr.append(sz)
    rPr.append(szCs)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), 'ABB2BF')
    rPr.append(c)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def build_nodes(doc, specs):
    nodes = []
    for spec in specs:
        kind, text = spec[0], spec[1] if len(spec) > 1 else ""
        if kind == "h1":
            nodes.append(make_para(doc, text, "TU_Sub-heading 1"))
        elif kind == "h2":
            nodes.append(make_para(doc, text, "TU_Sub-heading 2"))
        elif kind == "h3":
            nodes.append(make_para(doc, text, "TU_Sub-heading 3"))
        elif kind == "p":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal"))
        elif kind == "code":
            for line in text.split("\n"):
                nodes.append(make_code_para_styled(doc, line if line else " ", None))
        elif kind == "figure":
            nodes.append(make_para(doc, text, "TU_Paragraph_Normal", alignment="center"))
    return nodes


def insert_before(body, ref_elem, nodes):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for j, node in enumerate(nodes):
        parent.insert(idx + j, node)


def insert_after(body, ref_elem, nodes):
    parent = ref_elem.getparent()
    idx = list(parent).index(ref_elem)
    for j, node in enumerate(nodes):
        parent.insert(idx + 1 + j, node)


def remove_between(body, start_elem, end_elem):
    children = list(body)
    s = children.index(start_elem)
    e = children.index(end_elem) if end_elem is not None else len(children)
    to_remove = children[s + 1:e]
    for elem in to_remove:
        body.remove(elem)
    return len(to_remove)


def find_para(doc, text_fragment, start=0, end=None):
    if end is None:
        end = len(doc.paragraphs)
    for i in range(start, min(end, len(doc.paragraphs))):
        if text_fragment in doc.paragraphs[i].text:
            return i
    return None


def find_heading1_indices(doc):
    indices = {}
    HEADING_MAP = {
        "introduction": 1, "literature review": 2, "methodology": 3,
        "system design": 4, "implementation": 5, "testing and evaluation": 6,
        "discussion": 7, "conclusion": 8,
    }
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        text = p.text.strip().lower()
        if style == "Heading 1":
            for ht, cn in HEADING_MAP.items():
                if ht in text and cn not in indices:
                    indices[cn] = i
                    break
        if style == "TU_Chapter":
            if "ภาคผนวก" == p.text.strip():
                indices["appendix"] = i
            elif "บรรณานุกรม" in p.text:
                indices["references"] = i
    return indices


def insert_image(doc, body, ref_element, image_path, width_inches=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    elem = p._element
    parent = elem.getparent()
    parent.remove(elem)
    ref_parent = ref_element.getparent()
    idx = list(ref_parent).index(ref_element)
    ref_parent.insert(idx + 1, elem)
    return elem


# ============================================================================
# 5.3 Prompt Engineering — expanded content
# ============================================================================
PROMPT_ENGINEERING_SPECS = [
    ("h2", "5.3.1 โครงสร้าง Prompt System"),
    ("p", "ระบบ prompt ถูกจัดเก็บในไฟล์ YAML (src/agent/hotel_prompt.yaml) "
          "แยกจากโค้ด ทำให้แก้ไขได้โดยไม่ต้อง deploy ใหม่ "
          "โครงสร้างประกอบด้วย 5 ส่วน:"),
    ("p", "1. main_prompt — กฎหลัก, การตรวจจับภาษา, รายการ tools, กฎการตอบ"),
    ("p", "2. booking_flow — ขั้นตอนการจอง, การเก็บข้อมูล, dynamic pricing"),
    ("p", "3. service_prompt — การจัดการ service requests"),
    ("p", "4. error_handling — การจัดการข้อผิดพลาด"),
    ("p", "5. greeting_templates — เทมเพลตทักทายภาษาไทยและอังกฤษ"),

    ("h2", "5.3.2 Main Prompt"),
    ("p", "Main prompt เป็นหัวใจของระบบ กำหนดบทบาท กฎภาษา และรายการ tools:"),
    ("code", "# src/agent/hotel_prompt.yaml — main_prompt\n"
             "main_prompt: |\n"
             "  You are a professional bilingual (Thai/English) hotel\n"
             "  assistant for The Grand Horizon Hotel, a luxury 5-star\n"
             "  hotel in Thailand.\n"
             "\n"
             "  **CRITICAL LANGUAGE RULE**: Detect the guest's language\n"
             "  from their LATEST message only.\n"
             "  - English message → respond ENTIRELY in English\n"
             "  - Thai message → respond ENTIRELY in Thai\n"
             "  - NEVER mix languages.\n"
             "\n"
             "  ## Date & Time\n"
             "  Current: {current_date} {current_time} (Bangkok, GMT+7)\n"
             "\n"
             "  ## Tools (ALWAYS use tools — never answer from memory)\n"
             "  - search_hotel_knowledge → hotel info\n"
             "  - check_room_availability → rooms, pricing\n"
             "  - create_reservation / confirm / cancel → booking ops\n"
             "  - get_reservation_details → lookup by HTL number\n"
             "  - calculate_dynamic_price → actual pricing\n"
             "  - check_upsell_opportunity → room upgrade\n"
             "  - generate_payment_link → payment after confirm\n"
             "\n"
             "  ## Response rules\n"
             "  1. Give COMPLETE answers with details (floor, hours, prices)\n"
             "  2. Use Thai honorifics when responding in Thai\n"
             "  3. Confirm details before making changes\n"
             "  4. Offer alternatives when unavailable\n"
             "  5. Never quote prices from memory — always use tools"),
    ("p", "ตัวแปร {current_date} และ {current_time} ถูก inject อัตโนมัติตอน runtime "
          "ผ่านฟังก์ชัน load_hotel_prompts() เพื่อให้ agent ทราบวันเวลาปัจจุบัน"),

    ("h2", "5.3.3 Booking Flow Prompt"),
    ("p", "Booking prompt กำหนดขั้นตอนการจองที่ agent ต้องปฏิบัติตาม:"),
    ("code", "# src/agent/hotel_prompt.yaml — booking_flow\n"
             "booking_flow: |\n"
             "  ## Booking Guide\n"
             "  Collect naturally:\n"
             "  1) Dates\n"
             "  2) Room type preference → use check_room_availability\n"
             "  3) Number of guests (default 1)\n"
             "  4) Email (required before booking)\n"
             "\n"
             "  Use calculate_dynamic_price to show actual price.\n"
             "  Pick first available room — don't ask for room numbers.\n"
             "\n"
             "  After booking: call check_upsell_opportunity,\n"
             "  suggest if available. Offer generate_payment_link."),

    ("h2", "5.3.4 NeMo Guardrails Safety Prompts"),
    ("p", "NeMo Guardrails ใช้ self-check prompts สำหรับกรอง input/output:"),
    ("code", "# src/hotel_guardrails/config/prompts.yml\n"
             "prompts:\n"
             "  - task: self_check_input\n"
             "    content: |\n"
             "      Your task is to check if the user message\n"
             "      complies with the hotel assistant policy.\n"
             "\n"
             "      Hotel Policy:\n"
             "      - Should be helpful about the hotel\n"
             "      - Should not engage with harmful content\n"
             "      - Should maintain professional conversation\n"
             "\n"
             "      User message: \"{{ user_input }}\"\n"
             "      Question: Should the assistant respond?\n"
             "      Answer (yes/no):\n"
             "\n"
             "  - task: self_check_output\n"
             "    content: |\n"
             "      Check if the bot response complies with policy.\n"
             "      - Should not contain harmful content\n"
             "      - Should not reveal API keys or passwords\n"
             "      - Should maintain professional tone\n"
             "\n"
             "      Bot response: \"{{ bot_response }}\"\n"
             "      Question: Is this response appropriate?\n"
             "      Answer (yes/no):"),

    ("h2", "5.3.5 การเปลี่ยนแปลง Prompt ที่สำคัญ"),
    ("p", "ตลอดการพัฒนา prompt ถูกปรับปรุงหลายรอบเพื่อแก้ปัญหาที่พบ:"),
    ("p", ""),
    ("h3", "การเปลี่ยนแปลงที่ 1: ลดขนาด Prompt จาก 5,500 เป็น 2,800 ตัวอักษร"),
    ("p", "Prompt เดิมมีทุกคำสั่งซ้ำเป็นภาษาไทย เช่น:"),
    ("code", "# ก่อน (5,500 chars) — ทุกคำสั่งซ้ำ 2 ภาษา\n"
             "\"DO NOT answer from memory. Search the knowledge base.\"\n"
             "\"ห้ามตอบคำถามข้อมูลโรงแรมจากความจำ ค้นหาจากฐานความรู้ก่อนเสมอ\"\n"
             "\n"
             "# หลัง (2,800 chars) — ลบ Thai duplicate\n"
             "# เหตุผล: โมเดลเข้าใจทั้ง 2 ภาษาอยู่แล้ว\n"
             "# ไม่จำเป็นต้องสั่งซ้ำ"),
    ("p", "ผลกระทบ: ลด token ต่อ request 50% ทำให้ 9B model มี context window "
          "เหลือสำหรับ conversation history มากขึ้น"),

    ("h3", "การเปลี่ยนแปลงที่ 2: เพิ่ม Routing Examples ใน Router Prompt"),
    ("p", "Router prompt เดิมไม่มีตัวอย่างที่ชัดเจน ทำให้ \"ยกเลิกการจอง\" "
          "ถูก route ไป HandleOtherTalk แทน ToHotelBooking:"),
    ("code", "# เพิ่มตัวอย่างเฉพาะใน primary_prompt\n"
             "IMPORTANT routing rules:\n"
             "- \"cancel my booking\" / \"ยกเลิกการจอง\"\n"
             "  → ToHotelBooking (NOT HandleOtherTalk)\n"
             "- \"what services do you have?\"\n"
             "  → ToHotelKnowledge (general info)\n"
             "- \"I need a spa booking\"\n"
             "  → ToHotelService (specific request)\n"
             "- When in doubt → prefer ToHotelKnowledge"),
    ("p", "ผลกระทบ: cancel routing fixed จาก 0% เป็น 100% ใน test cases"),

    ("h3", "การเปลี่ยนแปลงที่ 3: Knowledge Context Placement"),
    ("p", "ตำแหน่งของ RAG context ส่งผลต่อคุณภาพคำตอบของ 9B model อย่างมาก:"),
    ("code", "# ผิด: context ก่อนคำถาม → โมเดลสรุป context\n"
             "messages = [\n"
             "    (\"system\", main_prompt),\n"
             "    (\"system\", f\"Context: {knowledge}\"),\n"
             "    (\"human\", user_message),       # ← ถูกบดบัง\n"
             "]\n"
             "\n"
             "# ถูกต้อง: คำถามก่อน context → โมเดลตอบคำถาม\n"
             "messages = [\n"
             "    (\"system\", main_prompt),\n"
             "    (\"human\", user_message),        # ← โฟกัสตรงนี้\n"
             "    (\"system\", f\"Hotel info:\\n{knowledge}\"),\n"
             "]"),
    ("p", "ผลกระทบ: Knowledge accuracy เพิ่มจาก 5/8 เป็น 8/8 บน 9B model"),

    ("h3", "การเปลี่ยนแปลงที่ 4: Greeting Templates"),
    ("p", "เทมเพลตทักทายแยกภาษาเพื่อให้คำตอบแรกเป็นธรรมชาติ:"),
    ("code", "# src/agent/hotel_prompt.yaml\n"
             "greeting_templates:\n"
             "  thai: |\n"
             "    สวัสดีค่ะ ยินดีต้อนรับสู่ The Grand Horizon Hotel\n"
             "    มีอะไรให้ช่วยเหลือไหมคะ?\n"
             "\n"
             "  english: |\n"
             "    Welcome to The Grand Horizon Hotel!\n"
             "    How may I assist you today?"),

    ("h3", "การเปลี่ยนแปลงที่ 5: Dynamic Date Injection"),
    ("p", "Prompt ถูก inject วันเวลาปัจจุบันทุกครั้งที่โหลด:"),
    ("code", "# src/hotel_guardrails/hotel_langgraph.py — load_hotel_prompts()\n"
             "bangkok_tz = timezone(timedelta(hours=7))\n"
             "now = datetime.now(bangkok_tz)\n"
             "prompts[\"main_prompt\"] = prompts[\"main_prompt\"].format(\n"
             "    current_date=now.strftime(\"%Y-%m-%d\"),\n"
             "    current_time=now.strftime(\"%H:%M\"),\n"
             ")"),
    ("p", "ทำให้ agent ทราบวันที่ปัจจุบันเพื่อคำนวณ \"วันจันทร์หน้า\" "
          "หรือ \"เดือนหน้า\" ได้ถูกต้อง"),
]


# ============================================================================
# Figure insertion map: text_fragment → PNG file
# ============================================================================
FIGURE_INSERTIONS = {
    # 5.2.9 Use Case Diagrams
    "Use Case 1: Guest": "Fig_5.25_UC_Guest.png",
    "Use Case 2: Admin": "Fig_5.26_UC_Admin.png",
    "Use Case 3:": "Fig_5.13_UC_Booking_Sequence.png",
    "Use Case 4:": "Fig_5.14_UC_Knowledge_RAG.png",
    "Use Case 5:": "Fig_5.15_UC_Auth_Flow.png",
    "Use Case 6: Chat Scaling": "Fig_5.16_UC_Concurrent_Chat.png",
    "LangGraph Agent Loop": "Fig_5.27_LangGraph_Loop.png",
    # 5.5 Escalation
    "5.5 ระบบ Human Escalation": "Fig_5.28_Auto_Escalation.png",
    # 5.8-5.13 workflows
    "5.8 Admin Monitoring": "Fig_5.17_Admin_Takeover.png",
    "5.9 Dynamic Pricing": "Fig_5.23_Dynamic_Pricing.png",
    "5.10 Time-Travel": "Fig_5.18_Timetravel.png",
    "5.11 Mock Payment": "Fig_5.19_Payment_Flow.png",
    "5.12 Audit Log": "Fig_5.20_Audit_Flow.png",
    "5.13 Knowledge Cache": "Fig_5.21_Knowledge_Cache.png",
    # CH4 diagrams
    "4.2.3 Reservation": "Fig_4.8_Reservation_Lifecycle.png",
    "4.2.4 Guest Identification": "Fig_4.9_Guest_Identification.png",
    "4.2.5 Authentication": "Fig_4.10_Dual_Identity.png",
    # Manage booking
    "Flow 3:": "Fig_5.22_Manage_Booking.png",
}


# ============================================================================
# MAIN
# ============================================================================
def main():
    print("=" * 60)
    print("Patch v8: Prompts + Figures + Page Breaks")
    print("=" * 60)

    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    body = doc.element.body

    chapters = find_heading1_indices(doc)
    ch5_start = chapters.get(5, 0)
    ch6_start = chapters.get(6, len(doc.paragraphs))

    # --- 1. Replace section 5.3 with expanded prompt content ---
    print("\n1. Expanding 5.3 Prompt Engineering...")
    anchor_53 = find_para(doc, "5.3 Prompt Engineering", ch5_start, ch6_start)
    if anchor_53 is not None:
        # Find 5.4 as the end boundary
        anchor_54 = find_para(doc, "5.4 PII Redaction", ch5_start, ch6_start)
        if anchor_54 is not None:
            start_elem = doc.paragraphs[anchor_53]._element
            end_elem = doc.paragraphs[anchor_54]._element
            removed = remove_between(body, start_elem, end_elem)
            print(f"  Removed {removed} old 5.3 elements")
            nodes = build_nodes(doc, PROMPT_ENGINEERING_SPECS)
            insert_after(body, start_elem, nodes)
            print(f"  Inserted {len(nodes)} new prompt paragraphs")
        else:
            print("  [WARN] 5.4 anchor not found")
    else:
        print("  [WARN] 5.3 anchor not found")

    # --- 2. Insert PNG figures ---
    print("\n2. Inserting PNG figures at diagram locations...")
    # Re-scan after insertion
    chapters = find_heading1_indices(doc)
    ch4_start = chapters.get(4, 0)
    ch6_start = chapters.get(6, len(doc.paragraphs))
    inserted_figs = 0
    for text_frag, fig_file in FIGURE_INSERTIONS.items():
        fig_path = FIGURES_DIR / fig_file
        if not fig_path.exists():
            print(f"  [SKIP] {fig_file} not found")
            continue
        anchor = find_para(doc, text_frag, ch4_start, ch6_start)
        if anchor is not None:
            try:
                insert_image(doc, body, doc.paragraphs[anchor]._element,
                            fig_path, width_inches=5.5)
                inserted_figs += 1
                print(f"  [OK] {fig_file} at P{anchor}")
            except Exception as e:
                print(f"  [FAIL] {fig_file}: {e}")
        else:
            pass  # Some fragments may not exist
    print(f"  Total figures: {inserted_figs}")

    # --- 3. Add page breaks before chapters ---
    print("\n3. Adding page breaks before chapters...")
    chapters = find_heading1_indices(doc)
    pagebreaks_added = 0
    for ch_num in sorted(k for k in chapters.keys() if isinstance(k, int)):
        if not isinstance(ch_num, int):
            continue
        idx = chapters[ch_num]
        p_elem = doc.paragraphs[idx]._element
        # Check if there's already a page break
        existing_br = p_elem.find(f'.//{qn("w:pPr")}/{qn("w:pageBreakBefore")}')
        if existing_br is not None:
            continue
        # Add pageBreakBefore to paragraph properties
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)
        pb = OxmlElement('w:pageBreakBefore')
        pPr.append(pb)
        pagebreaks_added += 1
    # Also add to ภาคผนวก and บรรณานุกรม
    for key in ["appendix", "references"]:
        if key in chapters:
            idx = chapters[key]
            p_elem = doc.paragraphs[idx]._element
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_elem.insert(0, pPr)
            existing_br = pPr.find(qn('w:pageBreakBefore'))
            if existing_br is None:
                pb = OxmlElement('w:pageBreakBefore')
                pPr.append(pb)
                pagebreaks_added += 1
    print(f"  Added {pagebreaks_added} page breaks")

    # --- 4. Save ---
    doc.save(str(OUT))
    sz = os.path.getsize(str(OUT)) / 1024
    print(f"\n{'=' * 60}")
    print(f"Saved: {OUT} ({sz:.0f} KB)")
    print("=" * 60)


if __name__ == "__main__":
    main()
